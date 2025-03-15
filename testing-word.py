from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns

def add_header_footer(doc):
    """Adds header and footer with page numbers in the format 'Page X of Y'"""
    section = doc.sections[0]

    # Header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "This is the Header"
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Footer
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add "Page X of Y" format
    run = footer_para.add_run("Page ")

    # PAGE field (Current Page Number)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(ns.qn("w:fldCharType"), "begin")

    instrText1 = OxmlElement("w:instrText")
    instrText1.set(ns.qn("xml:space"), "preserve")
    instrText1.text = "PAGE"

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(ns.qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText1)
    run._r.append(fldChar2)

    run.add_text(" of ")

    # NUMPAGES field (Total Number of Pages)
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(ns.qn("w:fldCharType"), "begin")

    instrText2 = OxmlElement("w:instrText")
    instrText2.set(ns.qn("xml:space"), "preserve")
    instrText2.text = "NUMPAGES"

    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(ns.qn("w:fldCharType"), "end")

    run._r.append(fldChar3)
    run._r.append(instrText2)
    run._r.append(fldChar4)


def add_colored_heading(doc, text, level, color):
    """Adds a heading with color"""
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    run.bold = True
    run.font.size = Pt(14) if level == 1 else Pt(12)
    run.font.color.rgb = color
    heading.style = f"Heading {level}"




def add_table(doc):
    """Creates a table with 'Field: Value' format and applies color to the header"""
    data = [
        ("Field", "Value"),
        ("1", "AA"),
        ("2", "AAA"),
        ("3", "QQ"),
        ("4", "rr")
    ]

    table = doc.add_table(rows=len(data), cols=2)
    table.style = "Table Grid"

    for i, row in enumerate(data):
        for j, text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = text

            # Apply background color only to the header row (first row)
            if i == 0:
                shading_elm = OxmlElement("w:shd")
                shading_elm.set(ns.qn("w:fill"), "D9D9D9")  # Light Gray color
                cell._tc.get_or_add_tcPr().append(shading_elm)


def create_word_document():
    doc = Document()

    # Add header and footer with page numbers
    add_header_footer(doc)
    


    # Add headings with colors
    add_colored_heading(doc, "Section 1", 1, RGBColor(0, 0, 255))  # Blue
    doc.add_paragraph("This is some text under Section 1.")

    add_colored_heading(doc, "Section 1.1", 2, RGBColor(0, 128, 0))  # Green
    doc.add_paragraph("This is some text under Section 1.1.")
    
    
    doc.add_page_break()  # Ensure TOC is separate from the content
    

    add_colored_heading(doc, "Section 1.2", 2, RGBColor(255, 0, 0))  # Red
    doc.add_paragraph("This is some text under Section 1.2.")

    add_colored_heading(doc, "Section 2", 3, RGBColor(128, 0, 128))  # Purple
    doc.add_paragraph("This is some text under Section 2.")

    # Add table
    add_table(doc)

    # Save document
    doc.save("generated_doc.docx")


create_word_document()
print("Word document created successfully! Open it and update the TOC & Page Numbers.")
