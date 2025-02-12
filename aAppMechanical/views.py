from django.shortcuts import render, redirect
from .forms import formCalcMS, formCalcBC, formCalcGR, formCalcPS, formCalcTH, formCalcMX, formCalcRT, formCalcCT, formCalcSC, formCalcBS
from datetime import datetime
from django.contrib.auth.models import User
from aApp1.models import UserRole, RoleAutho, Autho
import requests

from .models import FormFieldConfig
from django.http import HttpResponse
from docx import Document
import os
import ezdxf
from django.conf import settings


# Function to load the MS page
def load_ms_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to login page if user is not authenticated

    result = check_user_autho(request.user.username, 'MS')
    print('#####')
    print(result)
    print('######')

    # Fetch initial values from DB
    form_name = 'formCalcMS'
    field_configs = FormFieldConfig.objects.filter(form_name=form_name)
    initial_values = {config.field_name: config.initial_value for config in field_configs}

    form1 = formCalcMS(initial=initial_values)  # Pass DB values

    return render(request, 'MS.html', {'form1': form1})

# Function to handle form submission
def handle_ms_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated


    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcMS(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')
            oSec01Field08_value = form1.cleaned_data.get('oSec01Field08')
            oSec01Field09_value = form1.cleaned_data.get('oSec01Field09')
            oSec01Field10_value = form1.cleaned_data.get('oSec01Field10')
            oSec01Field11_value = form1.cleaned_data.get('oSec01Field11')

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "MS"  
            input_data = {
                "MS_ChannelHeight":     oSec01Field01_value,
                "MS_ScreenWidth":       oSec01Field02_value,
                "MS_BeltHeight":        oSec01Field03_value,
                "MS_WaterLevel":        oSec01Field04_value,
                "MS_BarSpacing":        oSec01Field05_value,
                "MS_BarThickness":      oSec01Field06_value,
                "MS_BarWidth":          oSec01Field07_value,
                "MS_InclinationDegree": oSec01Field08_value,
                "MS_SprocketDiameter":  oSec01Field09_value,
                "MS_Velocity":          oSec01Field10_value,
                "MS_FOS":               oSec01Field11_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["MS_w"]
            oSec02Field02_result = response["MS_p"]
            oSec02Field03_result = response["MS_s"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field2
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field3
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "MS"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcMS(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,
                'oSec01Field07': oSec01Field07_value,
                'oSec01Field08': oSec01Field08_value,
                'oSec01Field09': oSec01Field09_value,
                'oSec01Field10': oSec01Field10_value,
                'oSec01Field11': oSec01Field11_value,
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
            })

            return render(request, 'MS.html', {'form1': form1})

    return redirect('ms_load')  # Redirect to the page if the request is invalid


def generate_ms_report(request):
    
    if request.method == "POST":
        form1 = formCalcMS(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Mechanical Screen Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
                (form1.fields["oSec01Field05"].label, request.POST.get("oSec01Field05", "N/A")),
                (form1.fields["oSec01Field06"].label, request.POST.get("oSec01Field06", "N/A")),
                (form1.fields["oSec01Field07"].label, request.POST.get("oSec01Field07", "N/A")),
                (form1.fields["oSec01Field08"].label, request.POST.get("oSec01Field08", "N/A")),
                (form1.fields["oSec01Field09"].label, request.POST.get("oSec01Field09", "N/A")),
                (form1.fields["oSec01Field10"].label, request.POST.get("oSec01Field10", "N/A")),
                (form1.fields["oSec01Field11"].label, request.POST.get("oSec01Field11", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Mechanical_Screen_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)


def modify_ms_dxf(request):
    if request.method == "POST":
        # Define the path to the DXF file in the static directory
        static_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "MS_General_Drawing.dxf")
        modified_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "modified_fileNew.dxf")

        # Load the DXF file
        doc = ezdxf.readfile(static_path)

        # Iterate over the modelspace to find all DIMENSION entities
        for entity in doc.modelspace().query("DIMENSION"):
            if entity.dxf.text == "MS_chHeight":
                entity.dxf.text = request.POST.get("oSec01Field01", "000")
            elif entity.dxf.text == "MS_chWidth":
                entity.dxf.text = request.POST.get("oSec01Field02", "000")
            elif entity.dxf.text == "BeltHeight":
                entity.dxf.text = request.POST.get("oSec01Field03", "000")
            elif entity.dxf.text == "MS_angle":
                entity.dxf.text = request.POST.get("oSec01Field08", "000")
            elif entity.dxf.text == "MS_barSpace":
                entity.dxf.text = request.POST.get("oSec01Field05", "000")
            elif entity.dxf.text == "MS_barTh":
                entity.dxf.text = request.POST.get("oSec01Field06", "000")

            # Render the dimension to apply changes
            entity.render()

        # Save the modified DXF file
        doc.saveas(modified_path)

        # Serve the modified file for download
        with open(modified_path, "rb") as dxf_file:
            response = HttpResponse(dxf_file.read(), content_type="application/dxf")
            response["Content-Disposition"] = 'attachment; filename="modified_fileNew.dxf"'
            return response

    return HttpResponse("Invalid request", status=400)




# Function to load the MS page
def load_bc_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'BC')
    print('#####')
    print(result)
    print('######')
    
    # Fetch initial values from DB
    form_name = 'formCalcBC'
    field_configs = FormFieldConfig.objects.filter(form_name=form_name)
    initial_values = {config.field_name: config.initial_value for config in field_configs}

    form1 = formCalcBC(initial=initial_values)  # Pass DB values

    return render(request, 'BC.html', {'form1': form1})

# Function to handle form submission
def handle_bc_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcBC(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')

            # Calculate new values for fields

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "BC"  
            input_data = {
                "BC_Length":     oSec01Field01_value,
                "BC_Width":       oSec01Field02_value,
                "BC_DrumDia":        oSec01Field03_value,
                "BC_Friction":        oSec01Field04_value,
                "BC_Velocity":        oSec01Field05_value,
                "BC_FOS":      oSec01Field06_value,
                "BC_Belt_weight_per_meter":      oSec01Field07_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["BC_w"]
            oSec02Field02_result = response["BC_p"]
            oSec02Field03_result = response["BC_s"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field2
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field3
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "BC"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcBC(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,
                
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
            })

            return render(request, 'BC.html', {'form1': form1})
    
    
    return redirect('bc_load')  # Redirect to the page if the request is invalid


def generate_bc_report(request):
    
    if request.method == "POST":
        form1 = formCalcBC(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Belt Conveyor Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
                (form1.fields["oSec01Field05"].label, request.POST.get("oSec01Field05", "N/A")),
                (form1.fields["oSec01Field06"].label, request.POST.get("oSec01Field06", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Belt_Conveyor_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)


def modify_bc_dxf(request):
    if request.method == "POST":
        # Define the path to the DXF file in the static directory
        static_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "BC_General_Drawing.dxf")
        modified_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "BC_General_Drawing_new.dxf")

        # Load the DXF file
        doc = ezdxf.readfile(static_path)

        # Iterate over the modelspace to find all DIMENSION entities
        for entity in doc.modelspace().query("DIMENSION"):
            if entity.dxf.text == "BC_length":
                entity.dxf.text = request.POST.get("oSec01Field01", "000")
            elif entity.dxf.text == "BC_width":
                entity.dxf.text = request.POST.get("oSec01Field02", "000")

            # Render the dimension to apply changes
            entity.render()

        # Save the modified DXF file
        doc.saveas(modified_path)

        # Serve the modified file for download
        with open(modified_path, "rb") as dxf_file:
            response = HttpResponse(dxf_file.read(), content_type="application/dxf")
            response["Content-Disposition"] = 'attachment; filename="BC_General_Drawing_new.dxf"'
            return response

    return HttpResponse("Invalid request", status=400)





# Function to load the MS page
def load_gr_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'GR')
    print('#####')
    print(result)
    print('######')

    # Fetch initial values from DB
    form_name = 'formCalcGR'
    field_configs = FormFieldConfig.objects.filter(form_name=form_name)
    initial_values = {config.field_name: config.initial_value for config in field_configs}

    form1 = formCalcGR(initial=initial_values)  # Pass DB values

    return render(request, 'GR.html', {'form1': form1})

# Function to handle form submission
def handle_gr_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcGR(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')
            oSec01Field08_value = form1.cleaned_data.get('oSec01Field08')

            # Calculate new values for fields

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "GR"  
            input_data = {
                "GR_n_channel":     oSec01Field01_value,
                "GR_channel_width":       oSec01Field02_value,
                "GR_civil_width":        oSec01Field03_value,
                "GR_bridge_length":        oSec01Field04_value,
                "GR_wheel_diameter":        oSec01Field05_value,
                "GR_Friction":      oSec01Field06_value,
                "GR_Velocity":      oSec01Field07_value,
                "GR_FOS":      oSec01Field08_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["GR_out3"]
            oSec02Field02_result = response["GR_out1"]
            oSec02Field03_result = response["GR_out2"]
            oSec02Field04_result = response["GR_out4"]
            oSec02Field05_result = response["GR_out5"]
            oSec02Field06_result = response["GR_out6"]
            
            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field2
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field3
            instance.oSec02Field04 = oSec02Field04_result  # Update S2field1
            instance.oSec02Field05 = oSec02Field05_result  # Update S2field2
            instance.oSec02Field06 = oSec02Field06_result  # Update S2field3
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "GR"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcGR(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,
                'oSec01Field07': oSec01Field07_value,
                'oSec01Field08': oSec01Field08_value,
                
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
                'oSec02Field04': oSec02Field04_result,
                'oSec02Field05': oSec02Field05_result,
                'oSec02Field06': oSec02Field06_result,
            })

            return render(request, 'GR.html', {'form1': form1})
    
    
    return redirect('gr_load')  # Redirect to the page if the request is invalid


def generate_gr_report(request):
    
    if request.method == "POST":
        form1 = formCalcGR(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Gritremoval Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
                (form1.fields["oSec01Field05"].label, request.POST.get("oSec01Field05", "N/A")),
                (form1.fields["oSec01Field06"].label, request.POST.get("oSec01Field06", "N/A")),
                (form1.fields["oSec01Field07"].label, request.POST.get("oSec01Field07", "N/A")),
                (form1.fields["oSec01Field08"].label, request.POST.get("oSec01Field08", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
                (form1.fields["oSec02Field04"].label, request.POST.get("oSec02Field04", "N/A")),
                (form1.fields["oSec02Field05"].label, request.POST.get("oSec02Field05", "N/A")),
                (form1.fields["oSec02Field06"].label, request.POST.get("oSec02Field06", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Gritremoval_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)


def modify_gr_dxf(request):
    if request.method == "POST":
        # Define the path to the DXF file in the static directory
        static_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "MS_DXF_B.dxf")
        modified_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "modified_file.dxf")

        # Load the DXF file
        doc = ezdxf.readfile(static_path)

        # Iterate over the modelspace to find all DIMENSION entities
        for entity in doc.modelspace().query("DIMENSION"):
            if entity.dxf.text == "CHH111":
                entity.dxf.text = request.POST.get("oSec01Field01", "000")
            elif entity.dxf.text == "CHW111":
                entity.dxf.text = request.POST.get("oSec01Field02", "000")
            elif entity.dxf.text == "BeltHeight":
                entity.dxf.text = request.POST.get("oSec01Field03", "000")
            elif entity.dxf.text == "DEG111":
                entity.dxf.text = request.POST.get("oSec01Field08", "000")

            # Render the dimension to apply changes
            entity.render()

        # Save the modified DXF file
        doc.saveas(modified_path)

        # Serve the modified file for download
        with open(modified_path, "rb") as dxf_file:
            response = HttpResponse(dxf_file.read(), content_type="application/dxf")
            response["Content-Disposition"] = 'attachment; filename="modified_file.dxf"'
            return response

    return HttpResponse("Invalid request", status=400)






# Function to load the MS page
def load_ps_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'PS')
    print('#####')
    print(result)
    print('######')

    # Fetch initial values from DB
    form_name = 'formCalcPS'
    field_configs = FormFieldConfig.objects.filter(form_name=form_name)
    initial_values = {config.field_name: config.initial_value for config in field_configs}

    form1 = formCalcPS(initial=initial_values)  # Pass DB values
    return render(request, 'PS.html', {'form1': form1})

# Function to handle form submission
def handle_ps_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcPS(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')

            # Calculate new values for fields

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "PS"  
            input_data = {
                "PS_walkway_length":     oSec01Field01_value,
                "PS_Friction":       oSec01Field02_value,
                "PS_Velocity":        oSec01Field03_value,
                "PS_FOS":        oSec01Field04_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["PS_out2"]
            oSec02Field02_result = response["PS_out1"]
            oSec02Field03_result = '000'
            oSec02Field04_result = response["PS_out3"]
            oSec02Field05_result = response["PS_out4"]
            
            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field2
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field3
            instance.oSec02Field04 = oSec02Field04_result  # Update S2field1
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "PS"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcPS(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
                'oSec02Field04': oSec02Field04_result,
                'oSec02Field05': oSec02Field05_result,
            })

            return render(request, 'PS.html', {'form1': form1})
    
    
    return redirect('ps_load')  # Redirect to the page if the request is invalid



def generate_ps_report(request):
    
    if request.method == "POST":
        form1 = formCalcPS(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('PST Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
                (form1.fields["oSec02Field04"].label, request.POST.get("oSec02Field04", "N/A")),
                (form1.fields["oSec02Field05"].label, request.POST.get("oSec02Field05", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="PST_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)


def modify_ps_dxf(request):
    if request.method == "POST":
        # Define the path to the DXF file in the static directory
        static_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "PS_General_Drawing.dxf")
        modified_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "PS_General_Drawing_new.dxf")

        # Load the DXF file
        doc = ezdxf.readfile(static_path)

        # Iterate over the modelspace to find all DIMENSION entities
        for entity in doc.modelspace().query("DIMENSION"):
            if entity.dxf.text == "PS_walkwayLength":
                entity.dxf.text = request.POST.get("oSec01Field01", "000")

            # Render the dimension to apply changes
            entity.render()

        # Save the modified DXF file
        doc.saveas(modified_path)

        # Serve the modified file for download
        with open(modified_path, "rb") as dxf_file:
            response = HttpResponse(dxf_file.read(), content_type="application/dxf")
            response["Content-Disposition"] = 'attachment; filename="PS_General_Drawing_new.dxf"'
            return response

    return HttpResponse("Invalid request", status=400)







# Function to load the MS page
def load_th_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'TH')
    print('#####')
    print(result)
    print('######')

    # Fetch initial values from DB
    form_name = 'formCalcTH'
    field_configs = FormFieldConfig.objects.filter(form_name=form_name)
    initial_values = {config.field_name: config.initial_value for config in field_configs}

    form1 = formCalcTH(initial=initial_values)  # Pass DB values
    
    return render(request, 'TH.html', {'form1': form1})

# Function to handle form submission
def handle_th_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcTH(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')
            oSec01Field08_value = form1.cleaned_data.get('oSec01Field08')

            # Calculate new values for fields

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "TH"  
            input_data = {
                "TH_diameter":     oSec01Field01_value,
                "TH_n_arm":       oSec01Field02_value,
                "TH_Velocity":        oSec01Field03_value,
                "TH_FOS":        oSec01Field04_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["TH_w"]
            oSec02Field02_result = response["TH_p"]
            oSec02Field03_result = response["TH_s"]
            
            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field1
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field1
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "TH"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcTH(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,
                'oSec01Field07': oSec01Field07_value,
                'oSec01Field08': oSec01Field08_value,
                
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
            })

            return render(request, 'TH.html', {'form1': form1})
    
    
    return redirect('th_load')  # Redirect to the page if the request is invalid



def generate_th_report(request):
    
    if request.method == "POST":
        form1 = formCalcTH(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Thickener Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
                (form1.fields["oSec02Field04"].label, request.POST.get("oSec02Field04", "N/A")),
                (form1.fields["oSec02Field05"].label, request.POST.get("oSec02Field05", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Thickener_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)


def modify_th_dxf(request):
    if request.method == "POST":
        # Define the path to the DXF file in the static directory
        static_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "TH_General_Drawing.dxf")
        modified_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "TH_General_Drawing_new.dxf")

        # Load the DXF file
        doc = ezdxf.readfile(static_path)

        # Iterate over the modelspace to find all DIMENSION entities
        for entity in doc.modelspace().query("DIMENSION"):
            if entity.dxf.text == "TH_dia":
                entity.dxf.text = request.POST.get("oSec01Field01", "000")

            # Render the dimension to apply changes
            entity.render()

        # Save the modified DXF file
        doc.saveas(modified_path)

        # Serve the modified file for download
        with open(modified_path, "rb") as dxf_file:
            response = HttpResponse(dxf_file.read(), content_type="application/dxf")
            response["Content-Disposition"] = 'attachment; filename="TH_General_Drawing_new.dxf"'
            return response

    return HttpResponse("Invalid request", status=400)





# Function to load the MS page
def load_mx_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'MX')
    print('#####')
    print(result)
    print('######')

    # Fetch initial values from DB
    form_name = 'formCalcMX'
    field_configs = FormFieldConfig.objects.filter(form_name=form_name)
    initial_values = {config.field_name: config.initial_value for config in field_configs}

    form1 = formCalcMX(initial=initial_values)  # Pass DB values
    
    return render(request, 'MX.html', {'form1': form1})

# Function to handle form submission
def handle_mx_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        print('qqqqqqqq')
        form1 = formCalcMX(request.POST)        
        print('qqqqqqqq')
        if form1.is_valid():                   
            print('qqqqqqqq')
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')
            oSec01Field08_value = form1.cleaned_data.get('oSec01Field08')
                   
            print('qqqqqqqq')

            # Calculate new values for fields

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "MX"  
            input_data = {
                "MX_length":        oSec01Field01_value,
                "MX_width":         oSec01Field02_value,
                "MX_water_depth":           oSec01Field03_value,
                "MX_tank_depth":            oSec01Field04_value,
                "MX_impeller_coefficient":  oSec01Field05_value,
                "MX_velocity_gradient":     oSec01Field06_value,
                "MX_impeller_diameter_factor":  oSec01Field07_value,
                "MX_safety_factor":             oSec01Field08_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = '000'
            oSec02Field02_result = response["MX_p"]
            oSec02Field03_result = response["MX_s"]
            oSec02Field04_result = response["MX_d"]
            oSec02Field05_result = response["MX_shaftL"]
            oSec02Field06_result = response["MX_shaftD"]
            oSec02Field07_result = response["MX_Type"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "MX"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcMX(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,
                'oSec01Field07': oSec01Field07_value,
                'oSec01Field08': oSec01Field08_value,
                
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
                'oSec02Field04': oSec02Field04_result,
                'oSec02Field05': oSec02Field05_result,
                'oSec02Field06': oSec02Field06_result,
                'oSec02Field07': oSec02Field07_result,
            })

            return render(request, 'MX.html', {'form1': form1})
    
    
    return redirect('mx_load')  # Redirect to the page if the request is invalid



def generate_mx_report(request):
    
    if request.method == "POST":
        form1 = formCalcMX(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Mixer Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
                (form1.fields["oSec01Field05"].label, request.POST.get("oSec01Field05", "N/A")),
                (form1.fields["oSec01Field06"].label, request.POST.get("oSec01Field06", "N/A")),
                (form1.fields["oSec01Field07"].label, request.POST.get("oSec01Field07", "N/A")),
                (form1.fields["oSec01Field08"].label, request.POST.get("oSec01Field08", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
                (form1.fields["oSec02Field04"].label, request.POST.get("oSec02Field04", "N/A")),
                (form1.fields["oSec02Field05"].label, request.POST.get("oSec02Field05", "N/A")),
                (form1.fields["oSec02Field06"].label, request.POST.get("oSec02Field06", "N/A")),
                (form1.fields["oSec02Field07"].label, request.POST.get("oSec02Field07", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Mixer_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)


def modify_mx_dxf(request):
    if request.method == "POST":
        # Define the path to the DXF file in the static directory
        static_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "MS_DXF_B.dxf")
        modified_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "modified_file.dxf")

        # Load the DXF file
        doc = ezdxf.readfile(static_path)

        # Iterate over the modelspace to find all DIMENSION entities
        for entity in doc.modelspace().query("DIMENSION"):
            if entity.dxf.text == "CHH111":
                entity.dxf.text = request.POST.get("oSec01Field01", "000")
            elif entity.dxf.text == "CHW111":
                entity.dxf.text = request.POST.get("oSec01Field02", "000")
            elif entity.dxf.text == "BeltHeight":
                entity.dxf.text = request.POST.get("oSec01Field03", "000")
            elif entity.dxf.text == "DEG111":
                entity.dxf.text = request.POST.get("oSec01Field08", "000")

            # Render the dimension to apply changes
            entity.render()

        # Save the modified DXF file
        doc.saveas(modified_path)

        # Serve the modified file for download
        with open(modified_path, "rb") as dxf_file:
            response = HttpResponse(dxf_file.read(), content_type="application/dxf")
            response["Content-Disposition"] = 'attachment; filename="modified_file.dxf"'
            return response

    return HttpResponse("Invalid request", status=400)






# Function to load the MS page
def load_rt_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'RT')
    print('#####')
    print(result)
    print('######')
    
    # Fetch initial values from DB
    form_name = 'formCalcRT'
    field_configs = FormFieldConfig.objects.filter(form_name=form_name)
    initial_values = {config.field_name: config.initial_value for config in field_configs}

    form1 = formCalcRT(initial=initial_values)  # Pass DB values
    
    return render(request, 'RT.html', {'form1': form1})

# Function to handle form submission
def handle_rt_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcRT(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')
            oSec01Field08_value = form1.cleaned_data.get('oSec01Field08')

            # Calculate new values for fields

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "RT"  
            input_data = {
                "RT_Length":        oSec01Field01_value,
                "RT_Width":         oSec01Field02_value,
                "RT_Hight":           oSec01Field03_value,
                "RT_ShellTH":            oSec01Field04_value,
                "RT_BaseTH":    oSec01Field05_value,
                "RT_N_Spliter":     oSec01Field06_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["RT_w10"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "RT"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcRT(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,
                'oSec01Field07': oSec01Field07_value,
                'oSec01Field08': oSec01Field08_value,
                
                'oSec02Field01': oSec02Field01_result,
            })

            return render(request, 'RT.html', {'form1': form1})
    
    
    return redirect('rt_load')  # Redirect to the page if the request is invalid


def generate_rt_report(request):
    
    if request.method == "POST":
        form1 = formCalcRT(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Rect Tanks Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
                (form1.fields["oSec01Field05"].label, request.POST.get("oSec01Field05", "N/A")),
                (form1.fields["oSec01Field06"].label, request.POST.get("oSec01Field06", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Rect_Tank_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)


def modify_rt_dxf(request):
    if request.method == "POST":
        # Define the path to the DXF file in the static directory
        static_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "MS_DXF_B.dxf")
        modified_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "modified_file.dxf")

        # Load the DXF file
        doc = ezdxf.readfile(static_path)

        # Iterate over the modelspace to find all DIMENSION entities
        for entity in doc.modelspace().query("DIMENSION"):
            if entity.dxf.text == "CHH111":
                entity.dxf.text = request.POST.get("oSec01Field01", "000")
            elif entity.dxf.text == "CHW111":
                entity.dxf.text = request.POST.get("oSec01Field02", "000")
            elif entity.dxf.text == "BeltHeight":
                entity.dxf.text = request.POST.get("oSec01Field03", "000")
            elif entity.dxf.text == "DEG111":
                entity.dxf.text = request.POST.get("oSec01Field08", "000")

            # Render the dimension to apply changes
            entity.render()

        # Save the modified DXF file
        doc.saveas(modified_path)

        # Serve the modified file for download
        with open(modified_path, "rb") as dxf_file:
            response = HttpResponse(dxf_file.read(), content_type="application/dxf")
            response["Content-Disposition"] = 'attachment; filename="modified_file.dxf"'
            return response

    return HttpResponse("Invalid request", status=400)






# Function to load the MS page
def load_ct_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'CT')
    print('#####')
    print(result)
    print('######')
    
    # Fetch initial values from DB
    form_name = 'formCalcCT'
    field_configs = FormFieldConfig.objects.filter(form_name=form_name)
    initial_values = {config.field_name: config.initial_value for config in field_configs}

    form1 = formCalcCT(initial=initial_values)  # Pass DB values
    
    return render(request, 'PageCT.html', {'form1': form1})


# Function to handle form submission
def handle_ct_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        print("111")
        form1 = formCalcCT(request.POST)
        if form1.is_valid():
            print("222")
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')

            # Calculate new values for fields

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "CT"  
            input_data = {
                "CT_Diameter":        oSec01Field01_value,
                "CT_Height":         oSec01Field02_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["O_Tank_Weight"]
            oSec02Field02_result = response["O_Tank_Volume"]
            oSec02Field03_result = response["O_Tank_Shell_Th"]
            oSec02Field04_result = response["O_Tank_Base_Th"]
            oSec02Field05_result = response["O_Tank_Shell_Weight"]
            oSec02Field06_result = response["O_Tank_Base_Weight"]
            oSec02Field07_result = response["O_Tank_Base_UPN_Weight"]
            oSec02Field08_result = response["O_Tank_Cover_Weight"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field1
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field1
            instance.oSec02Field04 = oSec02Field04_result  # Update S2field1
            instance.oSec02Field05 = oSec02Field05_result  # Update S2field1
            instance.oSec02Field06 = oSec02Field06_result  # Update S2field1
            instance.oSec02Field07 = oSec02Field07_result  # Update S2field1
            instance.oSec02Field08 = oSec02Field08_result  # Update S2field1
            
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "CT"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcCT(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
                'oSec02Field04': oSec02Field04_result,
                'oSec02Field05': oSec02Field05_result,
                'oSec02Field06': oSec02Field06_result,
                'oSec02Field07': oSec02Field07_result,
                'oSec02Field08': oSec02Field08_result,
            })

            return render(request, 'PageCT.html', {'form1': form1})
    
    
    return redirect('ct_load')  # Redirect to the page if the request is invalid


def generate_ct_report(request):
    
    if request.method == "POST":
        form1 = formCalcCT(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Circular Tanks Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
                (form1.fields["oSec02Field04"].label, request.POST.get("oSec02Field04", "N/A")),
                (form1.fields["oSec02Field05"].label, request.POST.get("oSec02Field05", "N/A")),
                (form1.fields["oSec02Field06"].label, request.POST.get("oSec02Field06", "N/A")),
                (form1.fields["oSec02Field07"].label, request.POST.get("oSec02Field07", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Circular_Tank_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)


def modify_ct_dxf(request):
    if request.method == "POST":
        # Define the path to the DXF file in the static directory
        static_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "MS_DXF_B.dxf")
        modified_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "modified_file.dxf")

        # Load the DXF file
        doc = ezdxf.readfile(static_path)

        # Iterate over the modelspace to find all DIMENSION entities
        for entity in doc.modelspace().query("DIMENSION"):
            if entity.dxf.text == "CHH111":
                entity.dxf.text = request.POST.get("oSec01Field01", "000")
            elif entity.dxf.text == "CHW111":
                entity.dxf.text = request.POST.get("oSec01Field02", "000")
            elif entity.dxf.text == "BeltHeight":
                entity.dxf.text = request.POST.get("oSec01Field03", "000")
            elif entity.dxf.text == "DEG111":
                entity.dxf.text = request.POST.get("oSec01Field08", "000")

            # Render the dimension to apply changes
            entity.render()

        # Save the modified DXF file
        doc.saveas(modified_path)

        # Serve the modified file for download
        with open(modified_path, "rb") as dxf_file:
            response = HttpResponse(dxf_file.read(), content_type="application/dxf")
            response["Content-Disposition"] = 'attachment; filename="modified_file.dxf"'
            return response

    return HttpResponse("Invalid request", status=400)

    
    
    
    
    
    
    

# Function to load the MS page
def load_sc_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'SC')
    print('#####')
    print(result)
    print('######')
    
    # Fetch initial values from DB
    form_name = 'formCalcSC'
    field_configs = FormFieldConfig.objects.filter(form_name=form_name)
    initial_values = {config.field_name: config.initial_value for config in field_configs}

    form1 = formCalcSC(initial=initial_values)  # Pass DB values
    
    return render(request, 'PageSC.html', {'form1': form1})


# Function to handle form submission
def handle_sc_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcSC(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "SC"  
            input_data = {
                "SC_Density":        oSec01Field01_value,
                "SC_Length":         oSec01Field02_value,
                "SC_Diameter":       oSec01Field02_value,
                "SC_Safety":         oSec01Field02_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["SC_o1"]
            oSec02Field02_result = response["SC_o2"]
            oSec02Field03_result = response["SC_o3"]
            oSec02Field04_result = response["SC_o4"]
            oSec02Field05_result = response["SC_o5"]
            oSec02Field06_result = response["SC_o6"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field1
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field1
            instance.oSec02Field04 = oSec02Field04_result  # Update S2field1
            instance.oSec02Field05 = oSec02Field05_result  # Update S2field1
            instance.oSec02Field06 = oSec02Field06_result  # Update S2field1
            
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "SC"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcSC(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
                'oSec02Field04': oSec02Field04_result,
                'oSec02Field05': oSec02Field05_result,
                'oSec02Field06': oSec02Field06_result,
            })

            return render(request, 'PageSC.html', {'form1': form1})
    
    
    return redirect('sc_load')  # Redirect to the page if the request is invalid


def generate_sc_report(request):
    
    if request.method == "POST":
        form1 = formCalcCT(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Circular Tanks Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
                (form1.fields["oSec02Field04"].label, request.POST.get("oSec02Field04", "N/A")),
                (form1.fields["oSec02Field05"].label, request.POST.get("oSec02Field05", "N/A")),
                (form1.fields["oSec02Field06"].label, request.POST.get("oSec02Field06", "N/A")),
                (form1.fields["oSec02Field07"].label, request.POST.get("oSec02Field07", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Circular_Tank_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)


def modify_sc_dxf(request):
    if request.method == "POST":
        # Define the path to the DXF file in the static directory
        static_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "MS_DXF_B.dxf")
        modified_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "modified_file.dxf")

        # Load the DXF file
        doc = ezdxf.readfile(static_path)

        # Iterate over the modelspace to find all DIMENSION entities
        for entity in doc.modelspace().query("DIMENSION"):
            if entity.dxf.text == "CHH111":
                entity.dxf.text = request.POST.get("oSec01Field01", "000")
            elif entity.dxf.text == "CHW111":
                entity.dxf.text = request.POST.get("oSec01Field02", "000")
            elif entity.dxf.text == "BeltHeight":
                entity.dxf.text = request.POST.get("oSec01Field03", "000")
            elif entity.dxf.text == "DEG111":
                entity.dxf.text = request.POST.get("oSec01Field08", "000")

            # Render the dimension to apply changes
            entity.render()

        # Save the modified DXF file
        doc.saveas(modified_path)

        # Serve the modified file for download
        with open(modified_path, "rb") as dxf_file:
            response = HttpResponse(dxf_file.read(), content_type="application/dxf")
            response["Content-Disposition"] = 'attachment; filename="modified_file.dxf"'
            return response

    return HttpResponse("Invalid request", status=400)

    
###############
###############
###############
###############

def check_user_autho(username, autho_name):
    try:
        # Fetch the user by username
        user = User.objects.get(username=username)
        
        # Fetch the Autho by name
        autho = Autho.objects.get(name=autho_name)
        
        # Check if the user has a role and if that role has the specified Autho
        user_roles = UserRole.objects.filter(user=user)
        
        for user_role in user_roles:
            # Check if the role associated with the user has the specified Autho
            if RoleAutho.objects.filter(role=user_role.role, autho=autho).exists():
                return "T"  # User has the required Autho
            
        return "N"  # User does not have the required Autho
    
    except User.DoesNotExist:
        return "User not found"
    except Autho.DoesNotExist:
        return "Autho not found"





def interact_with_api(api_url, req_type, input_data):
    """
    Interact with the specified API by sending a POST request.

    Parameters:
        api_url (str): The API endpoint URL.
        req_type (str): The request type (e.g., 'MS').
        input_data (dict): A dictionary of input parameters.

    Returns:
        dict: The API response parsed into a Python dictionary.
    """
    # Prepare the payload
    payload = {
        "reqType": req_type,
        **input_data  # Merge the input data into the payload
    }

    try:
        # Send POST request
        response = requests.post(api_url, json=payload)

        # Raise an error for bad responses
        response.raise_for_status()

        # Parse and return the JSON response
        return response.json()

    except requests.exceptions.RequestException as e:
        print(f"Error while interacting with API: {e}")
        return None






#######################




# Function to load the MS page
def load_bs_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to login page if user is not authenticated

    result = check_user_autho(request.user.username, 'BS')
    print('#####')
    print(result)
    print('######')

    # Fetch initial values from DB
    form_name = 'formCalcBS'
    field_configs = FormFieldConfig.objects.filter(form_name=form_name)
    initial_values = {config.field_name: config.initial_value for config in field_configs}

    form1 = formCalcBS(initial=initial_values)  # Pass DB values

    
    print(form_name)

    return render(request, 'PageBS.html', {'form1': form1})

# Function to handle form submission
def handle_bs_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated


    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcBS(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "BS"  
            input_data = {
                "BS_Bar_Dia":     oSec01Field01_value,
                "BS_Bar_Space":       oSec01Field02_value,
                "BS_Screen_Height":        oSec01Field03_value,
                "BS_Screen_Width":        oSec01Field04_value,
                "BS_Screen_Depth":        oSec01Field05_value,
                "BS_Plate_Th":      oSec01Field06_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["O_Weight_allBars"]
            oSec02Field02_result = response["O_Plate_weight"]
            oSec02Field03_result = response["O_Total_weight"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field2
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field3
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "BS"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcBS(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,

                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
            })

            return render(request, 'PageBS.html', {'form1': form1})

    return redirect('ms_load')  # Redirect to the page if the request is invalid


def generate_bs_report(request):
    
    if request.method == "POST":
        form1 = formCalcBS(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Basket Screen Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
                (form1.fields["oSec01Field05"].label, request.POST.get("oSec01Field05", "N/A")),
                (form1.fields["oSec01Field06"].label, request.POST.get("oSec01Field06", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Basket_Screen_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)


def modify_bs_dxf(request):
    
    if request.method == "POST":
        form1 = formCalcBS(request.POST)
    return render(request, 'PageBS.html', {'form1': form1})

    # if request.method == "POST":
    #     # Define the path to the DXF file in the static directory
    #     static_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "MS_General_Drawing.dxf")
    #     modified_path = os.path.join(settings.BASE_DIR, "static", "aDxfs", "modified_fileNew.dxf")

    #     # Load the DXF file
    #     doc = ezdxf.readfile(static_path)

    #     # Iterate over the modelspace to find all DIMENSION entities
    #     for entity in doc.modelspace().query("DIMENSION"):
    #         if entity.dxf.text == "MS_chHeight":
    #             entity.dxf.text = request.POST.get("oSec01Field01", "000")
    #         elif entity.dxf.text == "MS_chWidth":
    #             entity.dxf.text = request.POST.get("oSec01Field02", "000")
    #         elif entity.dxf.text == "BeltHeight":
    #             entity.dxf.text = request.POST.get("oSec01Field03", "000")
    #         elif entity.dxf.text == "MS_angle":
    #             entity.dxf.text = request.POST.get("oSec01Field08", "000")
    #         elif entity.dxf.text == "MS_barSpace":
    #             entity.dxf.text = request.POST.get("oSec01Field05", "000")
    #         elif entity.dxf.text == "MS_barTh":
    #             entity.dxf.text = request.POST.get("oSec01Field06", "000")

    #         # Render the dimension to apply changes
    #         entity.render()

    #     # Save the modified DXF file
    #     doc.saveas(modified_path)

    #     # Serve the modified file for download
    #     with open(modified_path, "rb") as dxf_file:
    #         response = HttpResponse(dxf_file.read(), content_type="application/dxf")
    #         response["Content-Disposition"] = 'attachment; filename="modified_fileNew.dxf"'
    #         return response

    # return HttpResponse("Invalid request", status=400)

