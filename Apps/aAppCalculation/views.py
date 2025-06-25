

from datetime import datetime
from Apps.aAdmin.models import UserRole, RoleAutho, Autho
from Apps.aAppMechanical.models import aLogEntry
from Apps.aAppSubmittal.models import AddMachine
from Apps.aAppProject.models import APP_Project
from .models import modelcalc
from .models import API_Keys
from Apps.aAppMechanical.models import UserCompany
import requests

from .forms import FormCalculationSheet, FormCalculationSheet_log

from django.http import HttpResponse
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.shortcuts import render, redirect
from django.utils.timezone import now 
from django.contrib.auth.models import User

import os
import ezdxf

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,  WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.shared import Inches
from docx.shared import Pt
# Create your views here.


def check_user_autho(username, autho_name):
    try:
        # Fetch the user by username
        user = User.objects.get(username=username)
        
        # Fetch the Autho by name
        autho = Autho.objects.get(name=autho_name)
        
        # Check if the user has a role and if that role has the specified Autho
        user_roles = UserRole.objects.filter(user=user)
        
        for user_role in user_roles:
            # Check if the role associated with the user has the specified Autho
            if RoleAutho.objects.filter(role=user_role.role, autho=autho).exists():
                return "T"  # User has the required Autho
            
        return "N"  # User does not have the required Autho
    
    except User.DoesNotExist:
        return "User not found"
    except Autho.DoesNotExist:
        return "Autho not found"

def interact_with_api(api_url, req_type, input_data):
    """
    Interact with the specified API by sending a POST request.

    Parameters:
        api_url (str): The API endpoint URL.
        req_type (str): The request type (e.g., 'MS').
        input_data (dict): A dictionary of input parameters.

    Returns:
        dict: The API response parsed into a Python dictionary.
    """
    # Prepare the payload
    payload = {
        **req_type,
        **input_data  # Merge the input data into the payload
    }

    try:
        # Send POST request
        response = requests.post(api_url, json=payload)

        # Raise an error for bad responses
        response.raise_for_status()

        # Parse and return the JSON response
        return response.json()

    except requests.exceptions.RequestException as e:
        print(f"Error while interacting with API: {e}")
        return None


###################################
###################################


def LoadPageCalculationSheet(request):
    machineShow = "Hide"
    # Redirect unauthenticated users
    if not request.user.is_authenticated:
        return redirect("login")
    
    # Get the company of the logged-in user    
    user_company = None
    if request.user.is_authenticated:
        try:
            user_company = UserCompany.objects.get(user=request.user).company
        except UserCompany.DoesNotExist:
            user_company = None

    print(user_company)

    sheet_keys = AddMachine.objects.exclude(nameFormCalcXX__isnull=True).exclude(nameFormCalcXX__exact="None").exclude(nameFormCalcXX__exact="No").filter(company=user_company).order_by('order')

    sheet_key = None

    # If POST, get the selected sheet_key
    if request.method == "POST":
        sheet_key = request.POST.get("sheet_key")
        if sheet_key :
            machineShow = "Yes"

    #pdb.set_trace()
    print(sheet_key)
    

    result = check_user_autho(request.user.username, sheet_key)
    print('#####')
    print(result)
    print('######')
    
    print(request.user)
    print(f"{request.user} accessed Load {sheet_key}")
    ###LOG
    
    
    

    #Define Retrieve values from AddMachine model
    try:
        machine_config = AddMachine.objects.get(keyValue=sheet_key, company=user_company)
        form_type = machine_config.nameFormCalcXX
        aMachineName = machine_config.nameMachine
    except AddMachine.DoesNotExist:
        form_type = "None"
        aMachineName = "None"
        


    # Optional: Handle cases where the sheet_key is invalid
    if form_type is None:
        print(f"Warning: Unknown sheet_key '{sheet_key}'")


    # Assign company filter only if the user has a company
    if user_company:
        machines = modelcalc.objects.filter(oSec00Field03=sheet_key, company=user_company)
        projects = APP_Project.objects.filter(company=user_company)
    else:
        machines = modelcalc.objects.none()  # Return an empty queryset if no company
        projects = APP_Project.objects.none()  # Return an empty queryset if no company

    print(form_type)


    form = FormCalculationSheet(form_type=form_type)
    
    print(f"Initial value for oSec01Field02: {form.fields['oSec01Field02'].initial}")
    
    # Initialize all section variables
    section_vars = {}

    for section in range(1, 3):  # Section 1 and 2
        for field in range(1, 31):
            var_name = f"aSection{section:02d}Field{field:02d}Show"
            section_vars[var_name] = "Yes"

    # Apply conditions to modify the values
    for section in range(1, 3):
        for field in range(1, 31):
            field_name = f"oSec{section:02d}Field{field:02d}"
            var_name = f"aSection{section:02d}Field{field:02d}Show"
            initial_value = form.fields[field_name].initial
            if initial_value in ["oooo", None, ""]:
                section_vars[var_name] = "Hide"

    # Optional: print statements if you still need them
    for section in range(1, 3):
        for field in range(1, 31):
            field_name = f"oSec{section:02d}Field{field:02d}"
            print(form.fields[field_name].initial)

    for var_name in section_vars:
        print(section_vars[var_name])

    # Prepare context dictionary
    context = {
        "form": form,
        "machines": machines,
        "projects": projects,
        "aMachineName": aMachineName,
        "user_company": user_company,
        "sheet_key": sheet_key,
        "sheet_keys": sheet_keys,
        "machineShow": machineShow,
    }

    # Add section variables to context
    context.update(section_vars)

    return render(request, "PageCalculationSheet.html", context)


def HandleCalculationSheetForm(request):
    sheet_key = request.POST.get("sheet_key")
    print(sheet_key)
    if sheet_key :
            machineShow = "Yes"
    if not request.user.is_authenticated:
        return redirect('login')
    
    print(request.user)
    print(f"{request.user} accessed Load {sheet_key}")
    
    # Get the company of the logged-in user    
    user_company = None
    firstletter = None
    if request.user.is_authenticated:
        try:
            user_company = UserCompany.objects.get(user=request.user).company
            firstletter = user_company.nameCompanies[0]
        except UserCompany.DoesNotExist:
            user_company = None

    print(user_company)

    sheet_keys = AddMachine.objects.exclude(nameFormCalcXX__isnull=True).exclude(nameFormCalcXX__exact="None").exclude(nameFormCalcXX__exact="No").filter(company=user_company).order_by('order')
    
    
    #Define Retrieve values from AddMachine model
    try:
        machine_config = AddMachine.objects.get(keyValue=sheet_key, company=user_company)
        form_type = machine_config.nameFormCalcXX
        aMachineName = machine_config.nameMachine
    except AddMachine.DoesNotExist:
        form_type = "None"
        aMachineName = "None"
        

    req_sheet_key = sheet_key[0:-2]
    req_type = {"reqType" : req_sheet_key}
    input_fields = {
        data.apikey : data.fieldname
        for data in API_Keys.objects.filter(sheetkey = req_sheet_key, calctype = "Input")
    }
    output_fields = {
        data.fieldname : data.apikey
        for data in API_Keys.objects.filter(sheetkey = req_sheet_key, calctype = "Output")
    }

    # Assign company filter only if the user has a company
    if user_company:
        machines = modelcalc.objects.filter(oSec00Field03=sheet_key, company=user_company)
        projects = APP_Project.objects.filter(company=user_company)
    else:
        machines = modelcalc.objects.none()  # Return an empty queryset if no company
        projects = APP_Project.objects.none()  # Return an empty queryset if no company
    
    print(user_company)

    ###LOG
    
    

    if request.method == 'POST' and 'form1_submit' in request.POST:
        form = FormCalculationSheet(form_type=form_type, data=request.POST)
        if form.is_valid():
            

            input_data = {
                f"{api_key}": form.cleaned_data.get(field)
                for api_key, field in input_fields.items() if form.cleaned_data.get(field) is not None
            }
            print ("input_data : ", input_data)
            


            response = interact_with_api(
                "https://us-central1-h1000project1.cloudfunctions.net/f01",
                req_type,
                input_data
            )

            print("response : ", response)



            instance = form.save(commit=False)
            for form_field, api_key in output_fields.items():
                if api_key not in ["000", "1111"]:
                    setattr(instance, form_field, response[api_key])
            instance.oSec00Field01 = request.user.username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            instance.oSec00Field03 = sheet_key


            # Handle project assignment
            project_id = request.POST.get("project")
            if project_id:
                try:
                    instance.project = APP_Project.objects.get(id=project_id)
                except APP_Project.DoesNotExist:
                    """ return render(request, "PageCalculationSheet.html", {"form": form, "sheet_keys": sheet_keys, "projects": projects, "machines": machines, "aMachineName": aMachineName, "sheet_key": sheet_key, "machineShow": "Yes",  "error": "Invalid Project ID"})
            else:
                return render(request, "PageCalculationSheet.html", {"form": form, "sheet_keys": sheet_keys, "projects": projects, "machines": machines, "aMachineName": aMachineName, "sheet_key": sheet_key, "machineShow": "Yes",  "error": "Project is required"})
            """ 
            # Get the company associated with the user
            try:
                user_company = UserCompany.objects.get(user=request.user).company
                instance.company = user_company  # Assign company to the instance
            except UserCompany.DoesNotExist:
                return render(request, "PageCalculationSheet.html", 
                              {"form": form, 
                               "error": "User is not associated with a company",
                               "aMachineName": aMachineName,
                               "sheet_key" : sheet_key,
                               "sheet_keys": sheet_keys,
                               "machineShow":machineShow,})

            instance.save()

            # Refill form for display
            initial_data = {form_field: form.cleaned_data.get(form_field) for form_field in input_fields.values()}
            for form_field, api_key in output_fields.items():
                if api_key not in ["000", "1111"]:
                    initial_data[form_field] = response[api_key]

            form = FormCalculationSheet(form_type=form_type, initial=initial_data)

            form1 = FormCalculationSheet_log(form_type=form_type, data=request.POST)
            if form1.is_valid():


                input_data1 = {
                    f"{api_key}": form1.cleaned_data.get(field)
                    for api_key, field in input_fields.items() if form1.cleaned_data.get(field) is not None
                }
                print ("input_data1 : ", input_data1)


                response1 = interact_with_api(
                    "https://us-central1-h1000project1.cloudfunctions.net/f01",
                    req_type,
                    input_data1
                )

                print("response1 : ", response1)



                instance1 = form1.save(commit=False)
                for form_field, api_key in output_fields.items():
                    if api_key not in ["000", "1111"]:
                        setattr(instance1, form_field, response1[api_key])
                instance1.oSec00Field01 = request.user.username
                instance1.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                instance1.oSec00Field03 = sheet_key


                if project_id:
                    try:
                        instance1.project = APP_Project.objects.get(id=project_id)
                    except APP_Project.DoesNotExist:
                        """ return render(request, "PageCalculationSheet.html", {"form": form, "sheet_keys": sheet_keys, "projects": projects, "machines": machines, "aMachineName": aMachineName, "sheet_key": sheet_key, "machineShow": "Yes", "error": "Invalid Project ID"})
                else:
                    return render(request, "PageCalculationSheet.html", {"form": form, "sheet_keys": sheet_keys, "projects": projects, "machines": machines, "aMachineName": aMachineName, "sheet_key": sheet_key, "machineShow": "Yes",  "error": "Project is required"}) """

                # Get the company associated with the user
                try:
                    user_company = UserCompany.objects.get(user=request.user).company
                    instance1.company = user_company  # Assign company to the instance
                except UserCompany.DoesNotExist:
                    return render(request, "PageCalculationSheet.html", 
                                  {"form": form, 
                                   "error": "User is not associated with a company",
                                   "aMachineName": aMachineName,
                                   "sheet_key" : sheet_key,
                                   "sheet_keys": sheet_keys,
                                   "machineShow":machineShow,})

                instance1.save()

                # Refill form for display
                initial_data1 = {form_field: form1.cleaned_data.get(form_field) for form_field in input_fields.values()}
                for form_field, api_key in output_fields.items():
                    if api_key not in ["000", "1111"]:
                        initial_data1[form_field] = response1[api_key]

                form1 = FormCalculationSheet_log(form_type=form_type, initial=initial_data1)
            

            # Initialize visibility dictionaries
            aSection01FieldShow = {f"aSection01Field{str(i).zfill(2)}Show": "Yes" for i in range(1, 31)}
            aSection02FieldShow = {f"aSection02Field{str(i).zfill(2)}Show": "Yes" for i in range(1, 31)}
            
            # Update visibility based on field counts
            for i in range(1, 31):
                if form1.fields[f"oSec01Field{str(i).zfill(2)}"].initial in ["oooo", None , ""]:
                    aSection01FieldShow[f"aSection01Field{str(i).zfill(2)}Show"] = "Hide"
            
            for i in range(1, 31):
                if form1.fields[f"oSec02Field{str(i).zfill(2)}"].initial in ["oooo", None , ""]:
                    aSection02FieldShow[f"aSection02Field{str(i).zfill(2)}Show"] = "Hide"
            
            
            

            return render(request, 'PageCalculationSheet.html', {
                'form': form,
                'sheet_keys': sheet_keys,
                'sheet_key': sheet_key,
                'machines': machines,
                'projects': projects,  
                'machineShow': machineShow,
                'aMachineName': aMachineName, 
                'user_company': user_company, 
                **aSection01FieldShow,
                **aSection02FieldShow,
            })

    return redirect("PageCalculationSheet")

def generate_report(request):
    sheet_key = request.POST.get("sheet_key")
    print(sheet_key)
    try:
        #pdb.set_trace()
        # Log the action
        
        #pdb.set_trace()
        # Get the user’s company and project
        aCompany = UserCompany.objects.get(user=request.user)

        project_id = request.POST.get("project")

        #pdb.set_trace()
        # Determine the company and generate the corresponding report
        if aCompany.company.nameCompanies == "AAAA":
            print("Company 1")
            if project_id:
                return generate_report_AAA(request, project_id, sheet_key)
            else:
                return generate_report_AAA(request, " ", sheet_key)

        elif aCompany.company.nameCompanies == "BBBB":
            print("Company 2")
            if project_id:
                return generate_report_BBB(request, project_id, sheet_key)
            else:
                return generate_report_BBB(request, " ", sheet_key)

        else:
            return HttpResponse("Invalid company ID", status=400)

    except UserCompany.DoesNotExist:
        return HttpResponse("User does not belong to a company", status=403)

    except APP_Project.DoesNotExist:
        return HttpResponse("Project not found", status=404)




def generate_report_AAA(request, project_id, sheet_key):
    
    def add_table(doc, data, title=None):
        """Creates a table and applies a background color to the header."""
        if title:
            doc.add_heading(title, level=3)

        table = doc.add_table(rows=len(data), cols=2)
        table.style = "Table Grid"

        for i, row in enumerate(data):
            for j, text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = text

                # Apply background color only to the header row (first row)
                if i == 0:
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(ns.qn("w:fill"), "FFA500")  # Orange color                    
                    #shading_elm.set(ns.qn("w:fill"), "ADD8E6")  # Blue color
                        
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    
            """Generates a Word report for a given project."""
    
    
    
    def add_header_footer(doc):
        """Adds header and footer with page numbers in the format 'Page X of Y'."""
        section = doc.sections[0]
    
        # Header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        # header_para.add_run("Company Name\n")
        # header_para.add_run("Project Name\n")
        # header_para.add_run("Date: " + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "\n")
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Adding logo
        run_logo = header_para.add_run()  # Corrected reference to header paragraph
        
        try:
            run_logo.add_picture("static/aLogo/LogoAAA.PNG", width=Inches(7.0))  # Adjust width as needed
        except Exception as e:
            print(f"Error adding logo: {e}")

        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Add "Page X of Y" format
        run = footer_para.add_run("Page ")
    
        # PAGE field (Current Page Number)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(ns.qn("w:fldCharType"), "begin")
    
        instrText1 = OxmlElement("w:instrText")
        instrText1.set(ns.qn("xml:space"), "preserve")
        instrText1.text = "PAGE"
    
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar1)
        run._r.append(instrText1)
        run._r.append(fldChar2)
    
        run.add_text(" of ")
    
        # NUMPAGES field (Total Number of Pages)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(ns.qn("w:fldCharType"), "begin")
    
        instrText2 = OxmlElement("w:instrText")
        instrText2.set(ns.qn("xml:space"), "preserve")
        instrText2.text = "NUMPAGES"
    
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)
    
    
    def add_colored_heading(doc, text, level, color):
        """Adds a heading with color."""
        heading = doc.add_paragraph()
        run = heading.add_run(text)
        run.bold = True
        run.font.size = Pt(14) if level == 1 else Pt(12)
        run.font.color.rgb = color
        heading.style = f"Heading {level}"
    
    
    
    try:
        
        ###LOG
        
        print(f"at {now()} {User} accessed Download Report")
        ###LOG

        # Get the company of the logged-in user    
        user_company = None
        if request.user.is_authenticated:
            try:
                user_company = UserCompany.objects.get(user=request.user).company
            except UserCompany.DoesNotExist:
                user_company = None

        print(user_company)

        #Define Retrieve values from AddMachine model
        try:
            machine_config = AddMachine.objects.get(keyValue=sheet_key, company=user_company)
            form_type = machine_config.nameFormCalcXX
            aMachineName = machine_config.nameMachine
        except AddMachine.DoesNotExist:
            form_type = "None"
            aMachineName = "None"

        # Optional: Handle cases where the sheet_key is invalid
        if form_type is None:
            print(f"Warning: Unknown sheet_key '{sheet_key}'")

        aCompany = UserCompany.objects.get(user=request.user)
        if project_id != " ":
            project = APP_Project.objects.get(id=project_id)
        else:
            project = None
        form1 = FormCalculationSheet(form_type=form_type)
        
        
        print(aCompany.id)
        if project_id != " ":
            print(project.id)
        else:
            print("No Project ID")
    
        print("Company 1")
    
    
        # Create a Word document
        doc = Document()

        # Add header and footer with page numbers
        add_header_footer(doc)


        for _ in range(7):  # Adjust this number based on how centered you want it
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Centered content
        if project != None:
            lines = [
                "Project Name: ",
                project.name,
                "Client Name: ",
                project.client_name,
                "Capacity: ",
                project.capacity,
            ]
        else:
            lines = [
                "Project Name: ",
                "None",
                "Client Name: ",
                "None",
                "Capacity: ",
                "None",
            ]

        for line in lines:
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.size = Pt(25) 
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing after if needed
        para = doc.add_paragraph("\n")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()     
        doc.add_paragraph("\n")


        # Extract form data
        Input_data = [
            (request.POST.get("oSec01Field01", "N/A"), request.POST.get("oSec01Field02", "N/A")),
            (request.POST.get("oSec01Field03", "N/A"), request.POST.get("oSec01Field04", "N/A")),
            (request.POST.get("oSec01Field05", "N/A"), request.POST.get("oSec01Field06", "N/A")),
            (request.POST.get("oSec01Field07", "N/A"), request.POST.get("oSec01Field08", "N/A")),
            (request.POST.get("oSec01Field09", "N/A"), request.POST.get("oSec01Field10", "N/A")),
            (request.POST.get("oSec01Field11", "N/A"), request.POST.get("oSec01Field12", "N/A")),
            (request.POST.get("oSec01Field13", "N/A"), request.POST.get("oSec01Field14", "N/A")),
            (request.POST.get("oSec01Field15", "N/A"), request.POST.get("oSec01Field16", "N/A")),
            (request.POST.get("oSec01Field17", "N/A"), request.POST.get("oSec01Field18", "N/A")),
            (request.POST.get("oSec01Field19", "N/A"), request.POST.get("oSec01Field20", "N/A")),
            (request.POST.get("oSec01Field21", "N/A"), request.POST.get("oSec01Field22", "N/A")),
            (request.POST.get("oSec01Field23", "N/A"), request.POST.get("oSec01Field24", "N/A")),
            (request.POST.get("oSec01Field25", "N/A"), request.POST.get("oSec01Field26", "N/A")),
            (request.POST.get("oSec01Field27", "N/A"), request.POST.get("oSec01Field28", "N/A")),
            (request.POST.get("oSec01Field29", "N/A"), request.POST.get("oSec01Field30", "N/A")),
        ]
        Output_data = [
            (request.POST.get("oSec02Field01", "N/A"), request.POST.get("oSec02Field02", "N/A")),
            (request.POST.get("oSec02Field03", "N/A"), request.POST.get("oSec02Field04", "N/A")),
            (request.POST.get("oSec02Field05", "N/A"), request.POST.get("oSec02Field06", "N/A")),
            (request.POST.get("oSec02Field07", "N/A"), request.POST.get("oSec02Field08", "N/A")),
            (request.POST.get("oSec02Field09", "N/A"), request.POST.get("oSec02Field10", "N/A")),
            (request.POST.get("oSec02Field11", "N/A"), request.POST.get("oSec02Field12", "N/A")),
            (request.POST.get("oSec02Field13", "N/A"), request.POST.get("oSec02Field14", "N/A")),
            (request.POST.get("oSec02Field15", "N/A"), request.POST.get("oSec02Field16", "N/A")),
            (request.POST.get("oSec02Field17", "N/A"), request.POST.get("oSec02Field18", "N/A")),
            (request.POST.get("oSec02Field19", "N/A"), request.POST.get("oSec02Field20", "N/A")),
            (request.POST.get("oSec02Field21", "N/A"), request.POST.get("oSec02Field22", "N/A")),
            (request.POST.get("oSec02Field23", "N/A"), request.POST.get("oSec02Field24", "N/A")),
            (request.POST.get("oSec02Field25", "N/A"), request.POST.get("oSec02Field26", "N/A")),
            (request.POST.get("oSec02Field27", "N/A"), request.POST.get("oSec02Field28", "N/A")),
            (request.POST.get("oSec02Field29", "N/A"), request.POST.get("oSec02Field30", "N/A")),
        ]
        

        # Add machine title with font size 14 and numbering
        machine_title = doc.add_paragraph(f" {aMachineName}", style="Heading3")
        machine_title.runs[0].font.size = Pt(14)

        input_section_data = [("Field", "Value")]
        for input_key, input_value in Input_data:
            input_key = input_key.strip() if input_key else ""
            input_value = input_value.strip() if input_value else ""
            if input_key and input_value and input_key.lower() != "oooo" and input_key.lower() != "" and input_key.lower() != "n/a" and input_value.lower() != "oooo" and input_value.lower() != "" and input_value.lower() != "n/a":
                input_section_data.append((input_key, input_value))
        if len(input_section_data) > 1:  # If the section has valid data, create a table
            doc.add_paragraph(f"Input", style="Heading3")  # Only one title now
            add_table(doc, input_section_data)  # Removed redundant title

        output_section_data = [("Field", "Value")]
        for output_key, output_value in Output_data:
            output_key = output_key.strip() if output_key else ""
            output_value = output_value.strip() if output_value else ""
            if output_key and output_value and output_key.lower() != "oooo" and output_key.lower() != "" and output_key.lower() != "n/a" and output_value.lower() != "oooo" and output_value.lower() != "" and output_value.lower() != "n/a":
                output_section_data.append((output_key, output_value))
        if len(output_section_data) > 1:  # If the section has valid data, create a table
            doc.add_paragraph(f"Output", style="Heading3")  # Only one title now
            add_table(doc, output_section_data)  # Removed redundant title

        doc.add_page_break() 
        
        # Save the document to a response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        if project != None:
            response['Content-Disposition'] = f'attachment; filename={project.name}_report.docx'
        else:
            response['Content-Disposition'] = f'attachment; filename=None_report.docx'
        doc.save(response)
        return response

    except APP_Project.DoesNotExist:
        return HttpResponse("Project not found", status=404)





def generate_report_BBB(request, project_id, sheet_key):
    
    def add_table(doc, data, title=None):
        """Creates a borderless table and applies a background color to the header."""
        if title:
            doc.add_heading(title, level=3)

        table = doc.add_table(rows=len(data), cols=2)

        # Remove all table borders manually
        tbl = table._tbl  # Get the table's XML element
        tblPr = tbl.find(ns.qn("w:tblPr"))  # Find existing table properties

        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")  # Create table properties if missing
            tbl.insert(0, tblPr)  # Insert as the first child of <w:tbl>

        tblBorders = OxmlElement("w:tblBorders")  # Create <w:tblBorders>
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(ns.qn("w:val"), "nil")  # Remove the border
            tblBorders.append(border)

        tblPr.append(tblBorders)  # Append border settings to the table properties

        for i, row in enumerate(data):
            for j, text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = text

                # Apply background color only to the header row (first row)
                if i == 0:
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(ns.qn("w:val"), "clear")  # Set shading value
                    shading_elm.set(ns.qn("w:fill"), "ffffff")  # Light blue color
                    cell._tc.get_or_add_tcPr().append(shading_elm)
   
    
    
    def add_header_footer(doc):
        """Adds header and footer with page numbers in the format 'Page X of Y'."""
        section = doc.sections[0]
    
        # Header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        # header_para.add_run("Company Name\n")
        # header_para.add_run("Project Name\n")
        # header_para.add_run("Date: " + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "\n")
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Adding logo
        run_logo = header_para.add_run()  # Corrected reference to header paragraph
        try:
            run_logo.add_picture("static/aLogo/LogoBBB.PNG", width=Inches(7.0))  # Adjust width as needed
        except Exception as e:
            print(f"Error adding logo: {e}")

        
    
        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Add "Page X of Y" format
        run = footer_para.add_run("Page ")
    
        # PAGE field (Current Page Number)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(ns.qn("w:fldCharType"), "begin")
    
        instrText1 = OxmlElement("w:instrText")
        instrText1.set(ns.qn("xml:space"), "preserve")
        instrText1.text = "PAGE"
    
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar1)
        run._r.append(instrText1)
        run._r.append(fldChar2)
    
        run.add_text(" of ")
    
        # NUMPAGES field (Total Number of Pages)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(ns.qn("w:fldCharType"), "begin")
    
        instrText2 = OxmlElement("w:instrText")
        instrText2.set(ns.qn("xml:space"), "preserve")
        instrText2.text = "NUMPAGES"
    
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)
    
    
    def add_colored_heading(doc, text, level, color):
        """Adds a heading with color."""
        heading = doc.add_paragraph()
        run = heading.add_run(text)
        run.bold = True
        run.font.size = Pt(14) if level == 1 else Pt(12)
        run.font.color.rgb = color
        heading.style = f"Heading {level}"
    
    
    
    try:
        
        ###LOG
        
        print(f"at {now()} {User} accessed Download Report")
        ###LOG

        # Get the company of the logged-in user    
        user_company = None
        if request.user.is_authenticated:
            try:
                user_company = UserCompany.objects.get(user=request.user).company
            except UserCompany.DoesNotExist:
                user_company = None

        print(user_company)

        #Define Retrieve values from AddMachine model
        try:
            machine_config = AddMachine.objects.get(keyValue=sheet_key, company=user_company)
            form_type = machine_config.nameFormCalcXX
            aMachineName = machine_config.nameMachine
        except AddMachine.DoesNotExist:
            form_type = "None"
            aMachineName = "None"

        # Optional: Handle cases where the sheet_key is invalid
        if form_type is None:
            print(f"Warning: Unknown sheet_key '{sheet_key}'")

        aCompany = UserCompany.objects.get(user=request.user)
        if project_id != " ":
            project = APP_Project.objects.get(id=project_id)
        else:
            project = None
        form1 = FormCalculationSheet(form_type=form_type)
        
        
        print(aCompany.id)
        if project_id != " ":
            print(project.id)
        else:
            print("No Project ID")
    
        print("Company 2")
    
    
        # Create a Word document
        doc = Document()

        # Add header and footer with page numbers
        add_header_footer(doc)


        for _ in range(7):  # Adjust this number based on how centered you want it
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Centered content
        if project != None:
            lines = [
                "Project Name: ",
                project.name,
                "Client Name: ",
                project.client_name,
                "Capacity: ",
                project.capacity,
            ]
        else:
            lines = [
                "Project Name: ",
                "None",
                "Client Name: ",
                "None",
                "Capacity: ",
                "None",
            ]

        for line in lines:
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.size = Pt(25) 
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing after if needed
        para = doc.add_paragraph("\n")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()     
        doc.add_paragraph("\n")

        # Extract form data
        Input_data = [
            (request.POST.get("oSec01Field01", "N/A"), request.POST.get("oSec01Field02", "N/A")),
            (request.POST.get("oSec01Field03", "N/A"), request.POST.get("oSec01Field04", "N/A")),
            (request.POST.get("oSec01Field05", "N/A"), request.POST.get("oSec01Field06", "N/A")),
            (request.POST.get("oSec01Field07", "N/A"), request.POST.get("oSec01Field08", "N/A")),
            (request.POST.get("oSec01Field09", "N/A"), request.POST.get("oSec01Field10", "N/A")),
            (request.POST.get("oSec01Field11", "N/A"), request.POST.get("oSec01Field12", "N/A")),
            (request.POST.get("oSec01Field13", "N/A"), request.POST.get("oSec01Field14", "N/A")),
            (request.POST.get("oSec01Field15", "N/A"), request.POST.get("oSec01Field16", "N/A")),
            (request.POST.get("oSec01Field17", "N/A"), request.POST.get("oSec01Field18", "N/A")),
            (request.POST.get("oSec01Field19", "N/A"), request.POST.get("oSec01Field20", "N/A")),
            (request.POST.get("oSec01Field21", "N/A"), request.POST.get("oSec01Field22", "N/A")),
            (request.POST.get("oSec01Field23", "N/A"), request.POST.get("oSec01Field24", "N/A")),
            (request.POST.get("oSec01Field25", "N/A"), request.POST.get("oSec01Field26", "N/A")),
            (request.POST.get("oSec01Field27", "N/A"), request.POST.get("oSec01Field28", "N/A")),
            (request.POST.get("oSec01Field29", "N/A"), request.POST.get("oSec01Field30", "N/A")),
        ]
        Output_data = [
            (request.POST.get("oSec02Field01", "N/A"), request.POST.get("oSec02Field02", "N/A")),
            (request.POST.get("oSec02Field03", "N/A"), request.POST.get("oSec02Field04", "N/A")),
            (request.POST.get("oSec02Field05", "N/A"), request.POST.get("oSec02Field06", "N/A")),
            (request.POST.get("oSec02Field07", "N/A"), request.POST.get("oSec02Field08", "N/A")),
            (request.POST.get("oSec02Field09", "N/A"), request.POST.get("oSec02Field10", "N/A")),
            (request.POST.get("oSec02Field11", "N/A"), request.POST.get("oSec02Field12", "N/A")),
            (request.POST.get("oSec02Field13", "N/A"), request.POST.get("oSec02Field14", "N/A")),
            (request.POST.get("oSec02Field15", "N/A"), request.POST.get("oSec02Field16", "N/A")),
            (request.POST.get("oSec02Field17", "N/A"), request.POST.get("oSec02Field18", "N/A")),
            (request.POST.get("oSec02Field19", "N/A"), request.POST.get("oSec02Field20", "N/A")),
            (request.POST.get("oSec02Field21", "N/A"), request.POST.get("oSec02Field22", "N/A")),
            (request.POST.get("oSec02Field23", "N/A"), request.POST.get("oSec02Field24", "N/A")),
            (request.POST.get("oSec02Field25", "N/A"), request.POST.get("oSec02Field26", "N/A")),
            (request.POST.get("oSec02Field27", "N/A"), request.POST.get("oSec02Field28", "N/A")),
            (request.POST.get("oSec02Field29", "N/A"), request.POST.get("oSec02Field30", "N/A")),
        ]


        # Add machine title with font size 14 and numbering
        machine_title = doc.add_paragraph(f" {aMachineName}", style="Heading3")
        machine_title.runs[0].font.size = Pt(14)

        input_section_data = [(" ", " ")]
        for input_key, input_value in Input_data:
            input_key = input_key.strip() if input_key else ""
            input_value = input_value.strip() if input_value else ""
            if input_key and input_value and input_key.lower() != "oooo" and input_key.lower() != "" and input_key.lower() != "n/a" and input_value.lower() != "oooo" and input_value.lower() != "" and input_value.lower() != "n/a":
                input_section_data.append((input_key, input_value))
        if len(input_section_data) > 1:  # If the section has valid data, create a table
            doc.add_paragraph(f" ", style="Heading3")  # Only one title now
            add_table(doc, input_section_data)  # Removed redundant title

        output_section_data = [(" ", " ")]
        for output_key, output_value in Output_data:
            output_key = output_key.strip() if output_key else ""
            output_value = output_value.strip() if output_value else ""
            if output_key and output_value and output_key.lower() != "oooo" and output_key.lower() != "" and output_key.lower() != "n/a" and output_value.lower() != "oooo" and output_value.lower() != "" and output_value.lower() != "n/a":
                output_section_data.append((output_key, output_value))
        if len(output_section_data) > 1:  # If the section has valid data, create a table
            doc.add_paragraph(f" ", style="Heading3")  # Only one title now
            add_table(doc, output_section_data)  # Removed redundant title


        doc.add_page_break()     

        # Save the document to a response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        if project != None:
            response['Content-Disposition'] = f'attachment; filename={project.name}_report.docx'
        else:
            response['Content-Disposition'] = f'attachment; filename=None_report.docx'
        doc.save(response)
        return response

    except APP_Project.DoesNotExist:
        return HttpResponse("Project not found", status=404)


def DeleteCalcMachine(request, machine_id):  
    sheet_key = request.POST.get("sheet_key")
    print(sheet_key)
    if sheet_key :
            machineShow = "Yes"

    

    machine = get_object_or_404(modelcalc, id=machine_id)
    
    machine.delete()

    

     # Get the company of the logged-in user    
    user_company = None
    if request.user.is_authenticated:
        try:
            user_company = UserCompany.objects.get(user=request.user).company
        except UserCompany.DoesNotExist:
            user_company = None

    print(user_company)

    sheet_keys = AddMachine.objects.exclude(nameFormCalcXX__isnull=True).exclude(nameFormCalcXX__exact="None").exclude(nameFormCalcXX__exact="No").filter(company=user_company).order_by('order')

    #Define Retrieve values from AddMachine model
    try:
        machine_config = AddMachine.objects.get(keyValue=sheet_key, company=user_company)
        form_type = machine_config.nameFormCalcXX
        aMachineName = machine_config.nameMachine
    except AddMachine.DoesNotExist:
        form_type = "None"
        aMachineName = "None"
        


    # Optional: Handle cases where the sheet_key is invalid
    if form_type is None:
        print(f"Warning: Unknown sheet_key '{sheet_key}'")


     # Assign company filter only if the user has a company
    if user_company:
        machines = modelcalc.objects.filter(oSec00Field03=sheet_key, company=user_company)
        projects = APP_Project.objects.filter(company=user_company)
    else:
        machines = modelcalc.objects.none()  # Return an empty queryset if no company
        projects = APP_Project.objects.none()  # Return an empty queryset if no company

    print(form_type)


    form = FormCalculationSheet(form_type=form_type)
    
    print(f"Initial value for oSec01Field02: {form.fields['oSec01Field02'].initial}")
    
    
    # Initialize all field show variables for both sections
    field_show = {
        f"aSection{str(s).zfill(2)}Field{str(f).zfill(2)}Show": "Yes"
        for s in range(1, 3)
        for f in range(1, 31)
    }

    # Print initial values for debugging
    for s in range(1, 3):
        for f in range(1, 31):
            print(form.fields[f'oSec{str(s).zfill(2)}Field{str(f).zfill(2)}'].initial)

    # Apply conditions to hide fields if initial value is "oooo", None, or ""
    for s in range(1, 3):
        for f in range(1, 31):
            value = form.fields[f'oSec{str(s).zfill(2)}Field{str(f).zfill(2)}'].initial
            if value in ["oooo", None, ""]:
                field_show[f"aSection{str(s).zfill(2)}Field{str(f).zfill(2)}Show"] = "Hide"

    # Print show/hide states for debugging
    for s in range(1, 3):
        for f in range(1, 31):
            print(field_show[f"aSection{str(s).zfill(2)}Field{str(f).zfill(2)}Show"])

    # Render template
    return render(request, "PageCalculationSheet.html", {
        "form": form,
        "machines": machines,
        "projects": projects,
        "aMachineName": aMachineName,
        "user_company": user_company,
        "sheet_key": sheet_key,
        "sheet_keys": sheet_keys,
        "machineShow": machineShow,
        **field_show
    })


def CalculationSheet_get_data(request, machine_id):
    machine = get_object_or_404(modelcalc, id=machine_id)
    sheet_key = machine.oSec00Field03
    print(sheet_key)

    
    
    
    data = {
        "project": machine.project.name if machine.project else "No Project",
        "oSec01Field01": machine.oSec01Field01,
        "oSec01Field02": machine.oSec01Field02,
        "oSec01Field03": machine.oSec01Field03,
        "oSec01Field04": machine.oSec01Field04,
        "oSec01Field05": machine.oSec01Field05,
        "oSec01Field06": machine.oSec01Field06,
        "oSec01Field07": machine.oSec01Field07,
        "oSec01Field08": machine.oSec01Field08,
        "oSec01Field09": machine.oSec01Field09,
        "oSec01Field10": machine.oSec01Field10,        
        "oSec01Field11": machine.oSec01Field11,
        "oSec01Field12": machine.oSec01Field12,
        "oSec01Field13": machine.oSec01Field13,
        "oSec01Field14": machine.oSec01Field14,
        "oSec01Field15": machine.oSec01Field15,
        "oSec01Field16": machine.oSec01Field16,
        "oSec01Field17": machine.oSec01Field17,
        "oSec01Field18": machine.oSec01Field18,
        "oSec01Field19": machine.oSec01Field19,
        "oSec01Field20": machine.oSec01Field20,
        "oSec01Field21": machine.oSec01Field21,
        "oSec01Field22": machine.oSec01Field22,
        "oSec01Field23": machine.oSec01Field23,
        "oSec01Field24": machine.oSec01Field24,
        "oSec01Field25": machine.oSec01Field25,
        "oSec01Field26": machine.oSec01Field26,
        "oSec01Field27": machine.oSec01Field27,
        "oSec01Field28": machine.oSec01Field28,
        "oSec01Field29": machine.oSec01Field29,
        "oSec01Field30": machine.oSec01Field30,
        
        "oSec02Field01": machine.oSec02Field01,
        "oSec02Field02": machine.oSec02Field02,
        "oSec02Field03": machine.oSec02Field03,
        "oSec02Field04": machine.oSec02Field04,
        "oSec02Field05": machine.oSec02Field05,
        "oSec02Field06": machine.oSec02Field06,
        "oSec02Field07": machine.oSec02Field07,
        "oSec02Field08": machine.oSec02Field08,
        "oSec02Field09": machine.oSec02Field09,
        "oSec02Field10": machine.oSec02Field10,        
        "oSec02Field11": machine.oSec02Field11,
        "oSec02Field12": machine.oSec02Field12,
        "oSec02Field13": machine.oSec02Field13,
        "oSec02Field14": machine.oSec02Field14,
        "oSec02Field15": machine.oSec02Field15,
        "oSec02Field16": machine.oSec02Field16,
        "oSec02Field17": machine.oSec02Field17,
        "oSec02Field18": machine.oSec02Field18,
        "oSec02Field19": machine.oSec02Field19,
        "oSec02Field20": machine.oSec02Field20,
        "oSec02Field21": machine.oSec02Field21,
        "oSec02Field22": machine.oSec02Field22,
        "oSec02Field23": machine.oSec02Field23,
        "oSec02Field24": machine.oSec02Field24,
        "oSec02Field25": machine.oSec02Field25,
        "oSec02Field26": machine.oSec02Field26,
        "oSec02Field27": machine.oSec02Field27,
        "oSec02Field28": machine.oSec02Field28,
        "oSec02Field29": machine.oSec02Field29,
        "oSec02Field30": machine.oSec02Field30,
        
    }

    return JsonResponse(data)




def generate_saved_report(request, machine_id):
    try:
        #pdb.set_trace()
        # Log the action
        
        #pdb.set_trace()
        # Get the user’s company and project
        aCompany = UserCompany.objects.get(user=request.user)



        #pdb.set_trace()
        # Determine the company and generate the corresponding report
        if aCompany.company.nameCompanies == "AAAA":
            print("Company 1")
            return generate_saved_report_AAA(request, machine_id)

        elif aCompany.company.nameCompanies == "BBBB":
            print("Company 2")
            return generate_saved_report_BBB(request, machine_id)

        else:
            return HttpResponse("Invalid company ID", status=400)

    except UserCompany.DoesNotExist:
        return HttpResponse("User does not belong to a company", status=403)

    except APP_Project.DoesNotExist:
        return HttpResponse("Project not found", status=404)




def generate_saved_report_AAA(request, machine_id):
    
    def add_table(doc, data, title=None):
        """Creates a table and applies a background color to the header."""
        if title:
            doc.add_heading(title, level=3)

        table = doc.add_table(rows=len(data), cols=2)
        table.style = "Table Grid"

        for i, row in enumerate(data):
            for j, text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = text

                # Apply background color only to the header row (first row)
                if i == 0:
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(ns.qn("w:fill"), "FFA500")  # Orange color                    
                    #shading_elm.set(ns.qn("w:fill"), "ADD8E6")  # Blue color
                        
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    
            """Generates a Word report for a given project."""
    
    
    
    def add_header_footer(doc):
        """Adds header and footer with page numbers in the format 'Page X of Y'."""
        section = doc.sections[0]
    
        # Header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        # header_para.add_run("Company Name\n")
        # header_para.add_run("Project Name\n")
        # header_para.add_run("Date: " + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "\n")
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Adding logo
        run_logo = header_para.add_run()  # Corrected reference to header paragraph
        
        try:
            run_logo.add_picture("static/aLogo/LogoAAA.PNG", width=Inches(7.0))  # Adjust width as needed
        except Exception as e:
            print(f"Error adding logo: {e}")

        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Add "Page X of Y" format
        run = footer_para.add_run("Page ")
    
        # PAGE field (Current Page Number)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(ns.qn("w:fldCharType"), "begin")
    
        instrText1 = OxmlElement("w:instrText")
        instrText1.set(ns.qn("xml:space"), "preserve")
        instrText1.text = "PAGE"
    
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar1)
        run._r.append(instrText1)
        run._r.append(fldChar2)
    
        run.add_text(" of ")
    
        # NUMPAGES field (Total Number of Pages)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(ns.qn("w:fldCharType"), "begin")
    
        instrText2 = OxmlElement("w:instrText")
        instrText2.set(ns.qn("xml:space"), "preserve")
        instrText2.text = "NUMPAGES"
    
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)
    
    
    def add_colored_heading(doc, text, level, color):
        """Adds a heading with color."""
        heading = doc.add_paragraph()
        run = heading.add_run(text)
        run.bold = True
        run.font.size = Pt(14) if level == 1 else Pt(12)
        run.font.color.rgb = color
        heading.style = f"Heading {level}"
    
    
    
    try:
        
        ###LOG
        # aLogEntry.objects.create(
        #         user=request.user,
        #         message=f"at {now()} {request.user} accessed Load  "
        #     )
        print(f"at {now()} {User} accessed Download Report")
        ###LOG

        # Get the company of the logged-in user    
        user_company = None
        if request.user.is_authenticated:
            try:
                user_company = UserCompany.objects.get(user=request.user).company
            except UserCompany.DoesNotExist:
                user_company = None

        print(user_company)

        aCompany = UserCompany.objects.get(user=request.user)
        machine = get_object_or_404(modelcalc, id=machine_id)
        if machine.project:
            project = APP_Project.objects.get(name=machine.project.name)
        else:
            project = None
        
        print(aCompany.id)
        if project != None:
            print(project.id)
        else:
            print("No Project ID")
    
        print("Company 1")

        sheet_key = machine.oSec00Field03

        #Define Retrieve values from AddMachine model
        try:
            machine_config = AddMachine.objects.get(keyValue=sheet_key, company=user_company)
            form_type = machine_config.nameFormCalcXX
            aMachineName = machine_config.nameMachine
        except AddMachine.DoesNotExist:
            form_type = "None"
            aMachineName = "None"

        # Optional: Handle cases where the sheet_key is invalid
        if form_type is None:
            print(f"Warning: Unknown sheet_key '{sheet_key}'")
    
    
        # Create a Word document
        doc = Document()

        # Add header and footer with page numbers
        add_header_footer(doc)


        for _ in range(7):  # Adjust this number based on how centered you want it
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Centered content
        if project != None:
            lines = [
                "Project Name: ",
                project.name,
                "Client Name: ",
                project.client_name,
                "Capacity: ",
                project.capacity,
            ]
        else:
            lines = [
                "Project Name: ",
                "None",
                "Client Name: ",
                "None",
                "Capacity: ",
                "None",
            ]

        for line in lines:
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.size = Pt(25) 
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing after if needed
        para = doc.add_paragraph("\n")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()     
        doc.add_paragraph("\n")


        # Extract form data
        Input_data = [
            (machine.oSec01Field01, machine.oSec01Field02),
            (machine.oSec01Field03, machine.oSec01Field04),
            (machine.oSec01Field05, machine.oSec01Field06),
            (machine.oSec01Field07, machine.oSec01Field08),
            (machine.oSec01Field09, machine.oSec01Field10),
            (machine.oSec01Field11, machine.oSec01Field12),
            (machine.oSec01Field13, machine.oSec01Field14),
            (machine.oSec01Field15, machine.oSec01Field16),
            (machine.oSec01Field17, machine.oSec01Field18),
            (machine.oSec01Field19, machine.oSec01Field20),
            (machine.oSec01Field21, machine.oSec01Field22),
            (machine.oSec01Field23, machine.oSec01Field24),
            (machine.oSec01Field25, machine.oSec01Field26),
            (machine.oSec01Field27, machine.oSec01Field28),
            (machine.oSec01Field29, machine.oSec01Field30),
        ]
        Output_data = [
            (machine.oSec02Field01, machine.oSec02Field02),
            (machine.oSec02Field03, machine.oSec02Field04),
            (machine.oSec02Field05, machine.oSec02Field06),
            (machine.oSec02Field07, machine.oSec02Field08),
            (machine.oSec02Field09, machine.oSec02Field10),
            (machine.oSec02Field11, machine.oSec02Field12),
            (machine.oSec02Field13, machine.oSec02Field14),
            (machine.oSec02Field15, machine.oSec02Field16),
            (machine.oSec02Field17, machine.oSec02Field18),
            (machine.oSec02Field19, machine.oSec02Field20),
            (machine.oSec02Field21, machine.oSec02Field22),
            (machine.oSec02Field23, machine.oSec02Field24),
            (machine.oSec02Field25, machine.oSec02Field26),
            (machine.oSec02Field27, machine.oSec02Field28),
            (machine.oSec02Field29, machine.oSec02Field30),
        ]
        

        # Add machine title with font size 14 and numbering
        machine_title = doc.add_paragraph(f" {aMachineName}", style="Heading3")
        machine_title.runs[0].font.size = Pt(14)

        input_section_data = [("Field", "Value")]
        for input_key, input_value in Input_data:
            input_key = input_key.strip() if input_key else ""
            input_value = input_value.strip() if input_value else ""
            if input_key and input_value and input_key.lower() != "oooo" and input_key.lower() != "" and input_key.lower() != "n/a" and input_value.lower() != "oooo" and input_value.lower() != "" and input_value.lower() != "n/a":
                input_section_data.append((input_key, input_value))
        if len(input_section_data) > 1:  # If the section has valid data, create a table
            doc.add_paragraph(f"Input", style="Heading3")  # Only one title now
            add_table(doc, input_section_data)  # Removed redundant title

        output_section_data = [("Field", "Value")]
        for output_key, output_value in Output_data:
            output_key = output_key.strip() if output_key else ""
            output_value = output_value.strip() if output_value else ""
            if output_key and output_value and output_key.lower() != "oooo" and output_key.lower() != "" and output_key.lower() != "n/a" and output_value.lower() != "oooo" and output_value.lower() != "" and output_value.lower() != "n/a":
                output_section_data.append((output_key, output_value))
        if len(output_section_data) > 1:  # If the section has valid data, create a table
            doc.add_paragraph(f"Output", style="Heading3")  # Only one title now
            add_table(doc, output_section_data)  # Removed redundant title

        doc.add_page_break() 
        
        # Save the document to a response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        if project != None:
            response['Content-Disposition'] = f'attachment; filename={project.name}_report.docx'
        else:
            response['Content-Disposition'] = f'attachment; filename=None_report.docx'
        doc.save(response)
        return response

    except APP_Project.DoesNotExist:
        return HttpResponse("Project not found", status=404)





def generate_saved_report_BBB(request, machine_id):
    
    def add_table(doc, data, title=None):
        """Creates a borderless table and applies a background color to the header."""
        if title:
            doc.add_heading(title, level=3)

        table = doc.add_table(rows=len(data), cols=2)

        # Remove all table borders manually
        tbl = table._tbl  # Get the table's XML element
        tblPr = tbl.find(ns.qn("w:tblPr"))  # Find existing table properties

        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")  # Create table properties if missing
            tbl.insert(0, tblPr)  # Insert as the first child of <w:tbl>

        tblBorders = OxmlElement("w:tblBorders")  # Create <w:tblBorders>
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(ns.qn("w:val"), "nil")  # Remove the border
            tblBorders.append(border)

        tblPr.append(tblBorders)  # Append border settings to the table properties

        for i, row in enumerate(data):
            for j, text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = text

                # Apply background color only to the header row (first row)
                if i == 0:
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(ns.qn("w:val"), "clear")  # Set shading value
                    shading_elm.set(ns.qn("w:fill"), "ffffff")  # Light blue color
                    cell._tc.get_or_add_tcPr().append(shading_elm)
   
    
    
    def add_header_footer(doc):
        """Adds header and footer with page numbers in the format 'Page X of Y'."""
        section = doc.sections[0]
    
        # Header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        # header_para.add_run("Company Name\n")
        # header_para.add_run("Project Name\n")
        # header_para.add_run("Date: " + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "\n")
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Adding logo
        run_logo = header_para.add_run()  # Corrected reference to header paragraph
        try:
            run_logo.add_picture("static/aLogo/LogoBBB.PNG", width=Inches(7.0))  # Adjust width as needed
        except Exception as e:
            print(f"Error adding logo: {e}")

        
    
        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Add "Page X of Y" format
        run = footer_para.add_run("Page ")
    
        # PAGE field (Current Page Number)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(ns.qn("w:fldCharType"), "begin")
    
        instrText1 = OxmlElement("w:instrText")
        instrText1.set(ns.qn("xml:space"), "preserve")
        instrText1.text = "PAGE"
    
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar1)
        run._r.append(instrText1)
        run._r.append(fldChar2)
    
        run.add_text(" of ")
    
        # NUMPAGES field (Total Number of Pages)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(ns.qn("w:fldCharType"), "begin")
    
        instrText2 = OxmlElement("w:instrText")
        instrText2.set(ns.qn("xml:space"), "preserve")
        instrText2.text = "NUMPAGES"
    
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)
    
    
    def add_colored_heading(doc, text, level, color):
        """Adds a heading with color."""
        heading = doc.add_paragraph()
        run = heading.add_run(text)
        run.bold = True
        run.font.size = Pt(14) if level == 1 else Pt(12)
        run.font.color.rgb = color
        heading.style = f"Heading {level}"
    
    
    
    try:
        
        ###LOG
        
        print(f"at {now()} {User} accessed Download Report")
        ###LOG

        # Get the company of the logged-in user    
        user_company = None
        if request.user.is_authenticated:
            try:
                user_company = UserCompany.objects.get(user=request.user).company
            except UserCompany.DoesNotExist:
                user_company = None

        print(user_company)

        aCompany = UserCompany.objects.get(user=request.user)
        machine = get_object_or_404(modelcalc, id=machine_id)
        if machine.project:
            project = APP_Project.objects.get(name=machine.project.name)
        else:
            project = None
        
        
        print(aCompany.id)
        if project != None:
            print(project.id)
        else:
            print("No Project ID")
    
        print("Company 2")

        
        sheet_key = machine.oSec00Field03
    
        #Define Retrieve values from AddMachine model
        try:
            machine_config = AddMachine.objects.get(keyValue=sheet_key, company=user_company)
            form_type = machine_config.nameFormCalcXX
            aMachineName = machine_config.nameMachine
        except AddMachine.DoesNotExist:
            form_type = "None"
            aMachineName = "None"

        # Optional: Handle cases where the sheet_key is invalid
        if form_type is None:
            print(f"Warning: Unknown sheet_key '{sheet_key}'")

        # Create a Word document
        doc = Document()

        # Add header and footer with page numbers
        add_header_footer(doc)


        for _ in range(7):  # Adjust this number based on how centered you want it
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Centered content
        if project != None:
            lines = [
                "Project Name: ",
                project.name,
                "Client Name: ",
                project.client_name,
                "Capacity: ",
                project.capacity,
            ]
        else:
            lines = [
                "Project Name: ",
                "None",
                "Client Name: ",
                "None",
                "Capacity: ",
                "None",
            ]

        for line in lines:
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.size = Pt(25) 
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing after if needed
        para = doc.add_paragraph("\n")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()     
        doc.add_paragraph("\n")

        # Extract form data
        Input_data = [
            (machine.oSec01Field01, machine.oSec01Field02),
            (machine.oSec01Field03, machine.oSec01Field04),
            (machine.oSec01Field05, machine.oSec01Field06),
            (machine.oSec01Field07, machine.oSec01Field08),
            (machine.oSec01Field09, machine.oSec01Field10),
            (machine.oSec01Field11, machine.oSec01Field12),
            (machine.oSec01Field13, machine.oSec01Field14),
            (machine.oSec01Field15, machine.oSec01Field16),
            (machine.oSec01Field17, machine.oSec01Field18),
            (machine.oSec01Field19, machine.oSec01Field20),
            (machine.oSec01Field21, machine.oSec01Field22),
            (machine.oSec01Field23, machine.oSec01Field24),
            (machine.oSec01Field25, machine.oSec01Field26),
            (machine.oSec01Field27, machine.oSec01Field28),
            (machine.oSec01Field29, machine.oSec01Field30),
        ]
        Output_data = [
            (machine.oSec02Field01, machine.oSec02Field02),
            (machine.oSec02Field03, machine.oSec02Field04),
            (machine.oSec02Field05, machine.oSec02Field06),
            (machine.oSec02Field07, machine.oSec02Field08),
            (machine.oSec02Field09, machine.oSec02Field10),
            (machine.oSec02Field11, machine.oSec02Field12),
            (machine.oSec02Field13, machine.oSec02Field14),
            (machine.oSec02Field15, machine.oSec02Field16),
            (machine.oSec02Field17, machine.oSec02Field18),
            (machine.oSec02Field19, machine.oSec02Field20),
            (machine.oSec02Field21, machine.oSec02Field22),
            (machine.oSec02Field23, machine.oSec02Field24),
            (machine.oSec02Field25, machine.oSec02Field26),
            (machine.oSec02Field27, machine.oSec02Field28),
            (machine.oSec02Field29, machine.oSec02Field30),
        ]


        # Add machine title with font size 14 and numbering
        machine_title = doc.add_paragraph(f" {aMachineName}", style="Heading3")
        machine_title.runs[0].font.size = Pt(14)

        input_section_data = [(" ", " ")]
        for input_key, input_value in Input_data:
            input_key = input_key.strip() if input_key else ""
            input_value = input_value.strip() if input_value else ""
            if input_key and input_value and input_key.lower() != "oooo" and input_key.lower() != "" and input_key.lower() != "n/a" and input_value.lower() != "oooo" and input_value.lower() != "" and input_value.lower() != "n/a":
                input_section_data.append((input_key, input_value))
        if len(input_section_data) > 1:  # If the section has valid data, create a table
            doc.add_paragraph(f" ", style="Heading3")  # Only one title now
            add_table(doc, input_section_data)  # Removed redundant title

        output_section_data = [(" ", " ")]
        for output_key, output_value in Output_data:
            output_key = output_key.strip() if output_key else ""
            output_value = output_value.strip() if output_value else ""
            if output_key and output_value and output_key.lower() != "oooo" and output_key.lower() != "" and output_key.lower() != "n/a" and output_value.lower() != "oooo" and output_value.lower() != "" and output_value.lower() != "n/a":
                output_section_data.append((output_key, output_value))
        if len(output_section_data) > 1:  # If the section has valid data, create a table
            doc.add_paragraph(f" ", style="Heading3")  # Only one title now
            add_table(doc, output_section_data)  # Removed redundant title


        doc.add_page_break()     

        # Save the document to a response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        if project != None:
            response['Content-Disposition'] = f'attachment; filename={project.name}_report.docx'
        else:
            response['Content-Disposition'] = f'attachment; filename=None_report.docx'
        doc.save(response)
        return response

    except APP_Project.DoesNotExist:
        return HttpResponse("Project not found", status=404)