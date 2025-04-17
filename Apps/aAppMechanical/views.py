import pdb

from .models import Project
from .models import Machine
from .models import UserCompany
from .models import aLogEntry
from .models import FormFieldConfig

from .forms import formCalcMS, formCalcBC, formCalcGR, formCalcPS, formCalcTH
from .forms import formCalcMX, formCalcRT, formCalcCT, formCalcSC, formCalcBS
from .forms import formCalcNS, formCalcPNch, formCalcPNwa
from .forms import FormDataSheet  
from .forms import UserCompanyForm
from .forms import ProjectForm
from .forms import FormFieldConfigForm

from datetime import datetime
from Apps.aApp1.models import UserRole, RoleAutho, Autho
import requests

from django.http import HttpResponse
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.shortcuts import render, redirect
from django.urls import reverse
from django.utils.timezone import now 
from django.contrib.auth.models import User
from django.conf import settings


import os
import ezdxf

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Inches
from docx.shared import Pt

###################################
###################################
###################################
###################################
###################################
###################################

def list_configs(request):
    print("LINE52")
    sort_by = request.GET.get('sort', 'id')  # Default sorting by ID
    order = request.GET.get('order', 'asc')  # Default order is ascending

    valid_fields = ['id', 'form_name', 'field_name', 'label', 'initial_value', 'visibility', 'company']
    if sort_by not in valid_fields:
        sort_by = 'id'

    
    print("LINE61")
    
    # Apply sorting order
    if order == 'desc':
        sort_by = f'-{sort_by}'

    configs = FormFieldConfig.objects.all().order_by(sort_by)
    
    
    print("LINE70")
    
    return render(request, 'form_config.html', {
        'configs': configs,
        'sort_by': sort_by.strip(''),  # Remove '-' to keep track of column sorting
        'order': order
    })

def add_config(request):
    if request.method == "POST":
        form = FormFieldConfigForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('list_configs')
    else:
        form = FormFieldConfigForm()
    return render(request, 'form_config_form.html', {'form': form})

def edit_config(request, config_id):
    config = get_object_or_404(FormFieldConfig, id=config_id)
    if request.method == "POST":
        form = FormFieldConfigForm(request.POST, instance=config)
        if form.is_valid():
            form.save()
            return redirect('list_configs')
    else:
        form = FormFieldConfigForm(instance=config)
    return render(request, 'form_config_form.html', {'form': form})

def delete_config(request, config_id):
    config = get_object_or_404(FormFieldConfig, id=config_id)
    config.delete()
    return redirect('list_configs')

###################################
###################################


def check_user_autho(username, autho_name):
    try:
        # Fetch the user by username
        user = User.objects.get(username=username)
        
        # Fetch the Autho by name
        autho = Autho.objects.get(name=autho_name)
        
        # Check if the user has a role and if that role has the specified Autho
        user_roles = UserRole.objects.filter(user=user)
        
        for user_role in user_roles:
            # Check if the role associated with the user has the specified Autho
            if RoleAutho.objects.filter(role=user_role.role, autho=autho).exists():
                return "T"  # User has the required Autho
            
        return "N"  # User does not have the required Autho
    
    except User.DoesNotExist:
        return "User not found"
    except Autho.DoesNotExist:
        return "Autho not found"

def interact_with_api(api_url, req_type, input_data):
    """
    Interact with the specified API by sending a POST request.

    Parameters:
        api_url (str): The API endpoint URL.
        req_type (str): The request type (e.g., 'MS').
        input_data (dict): A dictionary of input parameters.

    Returns:
        dict: The API response parsed into a Python dictionary.
    """
    # Prepare the payload
    payload = {
        "reqType": req_type,
        **input_data  # Merge the input data into the payload
    }

    try:
        # Send POST request
        response = requests.post(api_url, json=payload)

        # Raise an error for bad responses
        response.raise_for_status()

        # Parse and return the JSON response
        return response.json()

    except requests.exceptions.RequestException as e:
        print(f"Error while interacting with API: {e}")
        return None



def assign_user_to_company(request):
    if request.method == "POST":
        form = UserCompanyForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect("assign_user")  # Redirect to a success page
    else:
        form = UserCompanyForm()
    
    return render(request, "assign_user_to_company.html", {"form": form})



###################################
###################################
###################################
###################################
###################################
###################################


def load_ms_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to login page if user is not authenticated

    result = check_user_autho(request.user.username, 'MS')
    print('#####')
    print(result)
    print('######')


    form1 = formCalcMS()  # Pass DB values

    return render(request, 'MS.html', {'form1': form1})

def handle_ms_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated


    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcMS(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')
            oSec01Field08_value = form1.cleaned_data.get('oSec01Field08')
            oSec01Field09_value = form1.cleaned_data.get('oSec01Field09')
            oSec01Field10_value = form1.cleaned_data.get('oSec01Field10')
            oSec01Field11_value = form1.cleaned_data.get('oSec01Field11')

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "MS"  
            input_data = {
                "MS_ChannelHeight":     oSec01Field01_value,
                "MS_ScreenWidth":       oSec01Field02_value,
                "MS_BeltHeight":        oSec01Field03_value,
                "MS_WaterLevel":        oSec01Field04_value,
                "MS_BarSpacing":        oSec01Field05_value,
                "MS_BarThickness":      oSec01Field06_value,
                "MS_BarWidth":          oSec01Field07_value,
                "MS_InclinationDegree": oSec01Field08_value,
                "MS_SprocketDiameter":  oSec01Field09_value,
                "MS_Velocity":          oSec01Field10_value,
                "MS_FOS":               oSec01Field11_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["MS_w"]
            oSec02Field02_result = response["MS_p"]
            oSec02Field03_result = response["MS_s"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field2
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field3
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "MS"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcMS(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,
                'oSec01Field07': oSec01Field07_value,
                'oSec01Field08': oSec01Field08_value,
                'oSec01Field09': oSec01Field09_value,
                'oSec01Field10': oSec01Field10_value,
                'oSec01Field11': oSec01Field11_value,
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
            })

            return render(request, 'MS.html', {'form1': form1})

    return redirect('ms_load')  # Redirect to the page if the request is invalid

def generate_ms_report(request):
    
    if request.method == "POST":
        form1 = formCalcMS(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Mechanical Screen Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
                (form1.fields["oSec01Field05"].label, request.POST.get("oSec01Field05", "N/A")),
                (form1.fields["oSec01Field06"].label, request.POST.get("oSec01Field06", "N/A")),
                (form1.fields["oSec01Field07"].label, request.POST.get("oSec01Field07", "N/A")),
                (form1.fields["oSec01Field08"].label, request.POST.get("oSec01Field08", "N/A")),
                (form1.fields["oSec01Field09"].label, request.POST.get("oSec01Field09", "N/A")),
                (form1.fields["oSec01Field10"].label, request.POST.get("oSec01Field10", "N/A")),
                (form1.fields["oSec01Field11"].label, request.POST.get("oSec01Field11", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Mechanical_Screen_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)



def load_bc_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'BC')
    print('#####')
    print(result)
    print('######')

    form1 = formCalcBC()  # Pass DB values

    return render(request, 'BC.html', {'form1': form1})

def handle_bc_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcBC(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')

            # Calculate new values for fields

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "BC"  
            input_data = {
                "BC_Length":     oSec01Field01_value,
                "BC_Width":       oSec01Field02_value,
                "BC_DrumDia":        oSec01Field03_value,
                "BC_Friction":        oSec01Field04_value,
                "BC_Velocity":        oSec01Field05_value,
                "BC_FOS":      oSec01Field06_value,
                "BC_Belt_weight_per_meter":      oSec01Field07_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["BC_w"]
            oSec02Field02_result = response["BC_p"]
            oSec02Field03_result = response["BC_s"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field2
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field3
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "BC"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcBC(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,
                
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
            })

            return render(request, 'BC.html', {'form1': form1})
    
    
    return redirect('bc_load')  # Redirect to the page if the request is invalid

def generate_bc_report(request):
    
    if request.method == "POST":
        form1 = formCalcBC(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Belt Conveyor Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
                (form1.fields["oSec01Field05"].label, request.POST.get("oSec01Field05", "N/A")),
                (form1.fields["oSec01Field06"].label, request.POST.get("oSec01Field06", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Belt_Conveyor_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)



def load_gr_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'GR')
    print('#####')
    print(result)
    print('######')

    form1 = formCalcGR()  # Pass DB values

    return render(request, 'GR.html', {'form1': form1})

def handle_gr_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcGR(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')
            oSec01Field08_value = form1.cleaned_data.get('oSec01Field08')

            # Calculate new values for fields

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "GR"  
            input_data = {
                "GR_n_channel":     oSec01Field01_value,
                "GR_channel_width":       oSec01Field02_value,
                "GR_civil_width":        oSec01Field03_value,
                "GR_bridge_length":        oSec01Field04_value,
                "GR_wheel_diameter":        oSec01Field05_value,
                "GR_Friction":      oSec01Field06_value,
                "GR_Velocity":      oSec01Field07_value,
                "GR_FOS":      oSec01Field08_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["GR_out3"]
            oSec02Field02_result = response["GR_out1"]
            oSec02Field03_result = response["GR_out2"]
            oSec02Field04_result = response["GR_out4"]
            oSec02Field05_result = response["GR_out5"]
            oSec02Field06_result = response["GR_out6"]
            
            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field2
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field3
            instance.oSec02Field04 = oSec02Field04_result  # Update S2field1
            instance.oSec02Field05 = oSec02Field05_result  # Update S2field2
            instance.oSec02Field06 = oSec02Field06_result  # Update S2field3
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "GR"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcGR(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,
                'oSec01Field07': oSec01Field07_value,
                'oSec01Field08': oSec01Field08_value,
                
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
                'oSec02Field04': oSec02Field04_result,
                'oSec02Field05': oSec02Field05_result,
                'oSec02Field06': oSec02Field06_result,
            })

            return render(request, 'GR.html', {'form1': form1})
    
    
    return redirect('gr_load')  # Redirect to the page if the request is invalid

def generate_gr_report(request):
    
    if request.method == "POST":
        form1 = formCalcGR(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Gritremoval Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
                (form1.fields["oSec01Field05"].label, request.POST.get("oSec01Field05", "N/A")),
                (form1.fields["oSec01Field06"].label, request.POST.get("oSec01Field06", "N/A")),
                (form1.fields["oSec01Field07"].label, request.POST.get("oSec01Field07", "N/A")),
                (form1.fields["oSec01Field08"].label, request.POST.get("oSec01Field08", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
                (form1.fields["oSec02Field04"].label, request.POST.get("oSec02Field04", "N/A")),
                (form1.fields["oSec02Field05"].label, request.POST.get("oSec02Field05", "N/A")),
                (form1.fields["oSec02Field06"].label, request.POST.get("oSec02Field06", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Gritremoval_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)





def load_ps_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'PS')
    print('#####')
    print(result)
    print('######')

    form1 = formCalcPS() 
    return render(request, 'PS.html', {'form1': form1})

def handle_ps_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcPS(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')

            # Calculate new values for fields

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "PS"  
            input_data = {
                "PS_walkway_length":     oSec01Field01_value,
                "PS_Friction":       oSec01Field02_value,
                "PS_Velocity":        oSec01Field03_value,
                "PS_FOS":        oSec01Field04_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["PS_out2"]
            oSec02Field02_result = response["PS_out1"]
            oSec02Field03_result = '000'
            oSec02Field04_result = response["PS_out3"]
            oSec02Field05_result = response["PS_out4"]
            
            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field2
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field3
            instance.oSec02Field04 = oSec02Field04_result  # Update S2field1
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "PS"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcPS(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
                'oSec02Field04': oSec02Field04_result,
                'oSec02Field05': oSec02Field05_result,
            })

            return render(request, 'PS.html', {'form1': form1})
    
    
    return redirect('ps_load')  # Redirect to the page if the request is invalid

def generate_ps_report(request):
    
    if request.method == "POST":
        form1 = formCalcPS(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('PST Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
                (form1.fields["oSec02Field04"].label, request.POST.get("oSec02Field04", "N/A")),
                (form1.fields["oSec02Field05"].label, request.POST.get("oSec02Field05", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="PST_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)



def load_th_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'TH')
    print('#####')
    print(result)
    print('######')

    form1 = formCalcTH() 
    
    return render(request, 'TH.html', {'form1': form1})

def handle_th_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcTH(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')
            oSec01Field08_value = form1.cleaned_data.get('oSec01Field08')

            # Calculate new values for fields

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "TH"  
            input_data = {
                "TH_diameter":     oSec01Field01_value,
                "TH_n_arm":       oSec01Field02_value,
                "TH_Velocity":        oSec01Field03_value,
                "TH_FOS":        oSec01Field04_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["TH_w"]
            oSec02Field02_result = response["TH_p"]
            oSec02Field03_result = response["TH_s"]
            
            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field1
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field1
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "TH"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcTH(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,
                'oSec01Field07': oSec01Field07_value,
                'oSec01Field08': oSec01Field08_value,
                
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
            })

            return render(request, 'TH.html', {'form1': form1})
    
    
    return redirect('th_load')  # Redirect to the page if the request is invalid

def generate_th_report(request):
    
    if request.method == "POST":
        form1 = formCalcTH(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Thickener Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
                (form1.fields["oSec02Field04"].label, request.POST.get("oSec02Field04", "N/A")),
                (form1.fields["oSec02Field05"].label, request.POST.get("oSec02Field05", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Thickener_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)




def load_mx_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'MX')
    print('#####')
    print(result)
    print('######')

    form1 = formCalcMX()
    
    return render(request, 'MX.html', {'form1': form1})

def handle_mx_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        print('qqqqqqqq')
        form1 = formCalcMX(request.POST)        
        print('qqqqqqqq')
        if form1.is_valid():                   
            print('qqqqqqqq')
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')
            oSec01Field08_value = form1.cleaned_data.get('oSec01Field08')
                   
            print('qqqqqqqq')

            # Calculate new values for fields

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "MX"  
            input_data = {
                "MX_length":        oSec01Field01_value,
                "MX_width":         oSec01Field02_value,
                "MX_water_depth":           oSec01Field03_value,
                "MX_tank_depth":            oSec01Field04_value,
                "MX_impeller_coefficient":  oSec01Field05_value,
                "MX_velocity_gradient":     oSec01Field06_value,
                "MX_impeller_diameter_factor":  oSec01Field07_value,
                "MX_safety_factor":             oSec01Field08_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = '000'
            oSec02Field02_result = response["MX_p"]
            oSec02Field03_result = response["MX_s"]
            oSec02Field04_result = response["MX_d"]
            oSec02Field05_result = response["MX_shaftL"]
            oSec02Field06_result = response["MX_shaftD"]
            oSec02Field07_result = response["MX_Type"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "MX"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcMX(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,
                'oSec01Field07': oSec01Field07_value,
                'oSec01Field08': oSec01Field08_value,
                
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
                'oSec02Field04': oSec02Field04_result,
                'oSec02Field05': oSec02Field05_result,
                'oSec02Field06': oSec02Field06_result,
                'oSec02Field07': oSec02Field07_result,
            })

            return render(request, 'MX.html', {'form1': form1})
    
    
    return redirect('mx_load')  # Redirect to the page if the request is invalid

def generate_mx_report(request):
    
    if request.method == "POST":
        form1 = formCalcMX(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Mixer Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
                (form1.fields["oSec01Field05"].label, request.POST.get("oSec01Field05", "N/A")),
                (form1.fields["oSec01Field06"].label, request.POST.get("oSec01Field06", "N/A")),
                (form1.fields["oSec01Field07"].label, request.POST.get("oSec01Field07", "N/A")),
                (form1.fields["oSec01Field08"].label, request.POST.get("oSec01Field08", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
                (form1.fields["oSec02Field04"].label, request.POST.get("oSec02Field04", "N/A")),
                (form1.fields["oSec02Field05"].label, request.POST.get("oSec02Field05", "N/A")),
                (form1.fields["oSec02Field06"].label, request.POST.get("oSec02Field06", "N/A")),
                (form1.fields["oSec02Field07"].label, request.POST.get("oSec02Field07", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Mixer_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)




def load_rt_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'RT')
    print('#####')
    print(result)
    print('######')

    form1 = formCalcRT()  
    
    return render(request, 'RT.html', {'form1': form1})

def handle_rt_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcRT(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')
            oSec01Field08_value = form1.cleaned_data.get('oSec01Field08')

            # Calculate new values for fields

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "RT"  
            input_data = {
                "RT_Length":        oSec01Field01_value,
                "RT_Width":         oSec01Field02_value,
                "RT_Hight":           oSec01Field03_value,
                "RT_ShellTH":            oSec01Field04_value,
                "RT_BaseTH":    oSec01Field05_value,
                "RT_N_Spliter":     oSec01Field06_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["RT_w10"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "RT"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcRT(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,
                'oSec01Field07': oSec01Field07_value,
                'oSec01Field08': oSec01Field08_value,
                
                'oSec02Field01': oSec02Field01_result,
            })

            return render(request, 'RT.html', {'form1': form1})
    
    
    return redirect('rt_load')  # Redirect to the page if the request is invalid

def generate_rt_report(request):
    
    if request.method == "POST":
        form1 = formCalcRT(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Rect Tanks Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
                (form1.fields["oSec01Field05"].label, request.POST.get("oSec01Field05", "N/A")),
                (form1.fields["oSec01Field06"].label, request.POST.get("oSec01Field06", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Rect_Tank_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)




def load_ct_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'CT')
    print('#####')
    print(result)
    print('######')

    form1 = formCalcCT()  
    
    return render(request, 'PageCT.html', {'form1': form1})

def handle_ct_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        print("111")
        form1 = formCalcCT(request.POST)
        if form1.is_valid():
            print("222")
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')

            # Calculate new values for fields

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "CT"  
            input_data = {
                "CT_Diameter":        oSec01Field01_value,
                "CT_Height":         oSec01Field02_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["O_Tank_Weight"]
            oSec02Field02_result = response["O_Tank_Volume"]
            oSec02Field03_result = response["O_Tank_Shell_Th"]
            oSec02Field04_result = response["O_Tank_Base_Th"]
            oSec02Field05_result = response["O_Tank_Shell_Weight"]
            oSec02Field06_result = response["O_Tank_Base_Weight"]
            oSec02Field07_result = response["O_Tank_Base_UPN_Weight"]
            oSec02Field08_result = response["O_Tank_Cover_Weight"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field1
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field1
            instance.oSec02Field04 = oSec02Field04_result  # Update S2field1
            instance.oSec02Field05 = oSec02Field05_result  # Update S2field1
            instance.oSec02Field06 = oSec02Field06_result  # Update S2field1
            instance.oSec02Field07 = oSec02Field07_result  # Update S2field1
            instance.oSec02Field08 = oSec02Field08_result  # Update S2field1
            
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "CT"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcCT(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
                'oSec02Field04': oSec02Field04_result,
                'oSec02Field05': oSec02Field05_result,
                'oSec02Field06': oSec02Field06_result,
                'oSec02Field07': oSec02Field07_result,
                'oSec02Field08': oSec02Field08_result,
            })

            return render(request, 'PageCT.html', {'form1': form1})
    
    
    return redirect('ct_load')  # Redirect to the page if the request is invalid

def generate_ct_report(request):
    
    if request.method == "POST":
        form1 = formCalcCT(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Circular Tanks Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
                (form1.fields["oSec02Field04"].label, request.POST.get("oSec02Field04", "N/A")),
                (form1.fields["oSec02Field05"].label, request.POST.get("oSec02Field05", "N/A")),
                (form1.fields["oSec02Field06"].label, request.POST.get("oSec02Field06", "N/A")),
                (form1.fields["oSec02Field07"].label, request.POST.get("oSec02Field07", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Circular_Tank_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)

  
    
    
    
    

def load_sc_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    result = check_user_autho(request.user.username, 'SC')
    print('#####')
    print(result)
    print('######')

    form1 = formCalcSC()  # Pass DB values
    
    return render(request, 'PageSC.html', {'form1': form1})

def handle_sc_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcSC(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')
            oSec01Field08_value = form1.cleaned_data.get('oSec01Field08')

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "SC"  
            input_data = {
                "aInput01":     oSec01Field01_value,
                "aInput02":     oSec01Field02_value,
                "aInput03":     oSec01Field03_value,
                "aInput04":     oSec01Field04_value,
                "aInput05":     oSec01Field05_value,
                "aInput06":     oSec01Field06_value,
                "aInput07":     oSec01Field07_value,
                "aInput08":     oSec01Field08_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["Pitch"]
            oSec02Field02_result = response["SpeedRPM"]
            oSec02Field03_result = response["MotorPower"]
            oSec02Field04_result = response["ScrewWeight"]
            oSec02Field05_result = response["FrameWeight"]
            oSec02Field06_result = "1111"

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field1
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field1
            instance.oSec02Field04 = oSec02Field04_result  # Update S2field1
            instance.oSec02Field05 = oSec02Field05_result  # Update S2field1
            instance.oSec02Field06 = oSec02Field06_result  # Update S2field1
            
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "SC"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcSC(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                
                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
                'oSec02Field04': oSec02Field04_result,
                'oSec02Field05': oSec02Field05_result,
                'oSec02Field06': oSec02Field06_result,
            })

            return render(request, 'PageSC.html', {'form1': form1})
    
    
    return redirect('sc_load')  # Redirect to the page if the request is invalid

def generate_sc_report(request):
    
    if request.method == "POST":
        form1 = formCalcCT(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Circular Tanks Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
                (form1.fields["oSec02Field04"].label, request.POST.get("oSec02Field04", "N/A")),
                (form1.fields["oSec02Field05"].label, request.POST.get("oSec02Field05", "N/A")),
                (form1.fields["oSec02Field06"].label, request.POST.get("oSec02Field06", "N/A")),
                (form1.fields["oSec02Field07"].label, request.POST.get("oSec02Field07", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Circular_Tank_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)




def load_bs_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to login page if user is not authenticated

    result = check_user_autho(request.user.username, 'BS')
    print('#####')
    print(result)
    print('######')

    form1 = formCalcBS()  # Pass DB values

    return render(request, 'PageBS.html', {'form1': form1})

def handle_bs_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated


    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = formCalcBS(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "BS"  
            input_data = {
                "BS_Bar_Dia":     oSec01Field01_value,
                "BS_Bar_Space":       oSec01Field02_value,
                "BS_Screen_Height":        oSec01Field03_value,
                "BS_Screen_Width":        oSec01Field04_value,
                "BS_Screen_Depth":        oSec01Field05_value,
                "BS_Plate_Th":      oSec01Field06_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["O_Weight_allBars"]
            oSec02Field02_result = response["O_Plate_weight"]
            oSec02Field03_result = response["O_Total_weight"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1
            instance.oSec02Field02 = oSec02Field02_result  # Update S2field2
            instance.oSec02Field03 = oSec02Field03_result  # Update S2field3
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "BS"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcBS(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,

                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
            })

            return render(request, 'PageBS.html', {'form1': form1})

    return redirect('ms_load')  # Redirect to the page if the request is invalid

def generate_bs_report(request):
    
    if request.method == "POST":
        form1 = formCalcBS(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Basket Screen Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
                (form1.fields["oSec01Field05"].label, request.POST.get("oSec01Field05", "N/A")),
                (form1.fields["oSec01Field06"].label, request.POST.get("oSec01Field06", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Basket_Screen_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)





def load_ns_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to login page if user is not authenticated

    result = check_user_autho(request.user.username, 'BS')
    print('#####')
    print(result)
    print('######')

    form1 = formCalcNS()  # Pass DB values

    return render(request, 'PageNS.html', {'form1': form1})

def handle_ns_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    
    print("Line 1866")


    if request.method == 'POST' and 'form1_submit' in request.POST:
        
        print("Line 1871")
        
        form1 = formCalcNS(request.POST)
        
        print("Line 1875")

        if form1.is_valid():
            
            print("Line 1879")
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')
            oSec01Field08_value = form1.cleaned_data.get('oSec01Field08')

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "NS"  
            input_data = {
                "NS_Ch_Height":         oSec01Field01_value,
                "NS_Ch_Width":          oSec01Field02_value,
                "NS_WaterLv":           oSec01Field03_value,
                "NS_WaterLv_Margin":    oSec01Field04_value,
                "NS_Bar_Spacing":       oSec01Field05_value,
                "NS_Bar_Th":            oSec01Field06_value,
                "NS_Bar_Width":         oSec01Field07_value,
                "NS_Angle":             oSec01Field08_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["O_Weight"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result  # Update S2field1            
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "NS"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcNS(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,
                'oSec01Field07': oSec01Field07_value,
                'oSec01Field08': oSec01Field08_value,

                'oSec02Field01': oSec02Field01_result,
            })

            return render(request, 'PageNS.html', {'form1': form1})

    return redirect('ms_load')  # Redirect to the page if the request is invalid

def generate_ns_report(request):
    
    if request.method == "POST":
        form1 = formCalcNS(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Manual Screen Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
                (form1.fields["oSec01Field05"].label, request.POST.get("oSec01Field05", "N/A")),
                (form1.fields["oSec01Field06"].label, request.POST.get("oSec01Field06", "N/A")),
                (form1.fields["oSec01Field07"].label, request.POST.get("oSec01Field07", "N/A")),
                (form1.fields["oSec01Field08"].label, request.POST.get("oSec01Field08", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Manual_Screen_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)





def load_pnch_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to login page if user is not authenticated

    result = check_user_autho(request.user.username, 'PNch')
    print('#####')
    print(result)
    print('######')

    # Fetch initial values from DB

    form1 = formCalcPNch()  # Pass DB values

    return render(request, 'PagePNch.html', {'form1': form1})

def handle_pnch_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    
    print("aaa")


    if request.method == 'POST' and 'form1_submit' in request.POST:
        
        print("aaa")
        
        form1 = formCalcPNch(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')
            oSec01Field08_value = form1.cleaned_data.get('oSec01Field08')
            oSec01Field09_value = form1.cleaned_data.get('oSec01Field09')
            oSec01Field10_value = form1.cleaned_data.get('oSec01Field10')
            
            
            print(oSec01Field08_value)
            print(oSec01Field09_value)
            print(oSec01Field10_value)

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "PNch"  
            input_data = {
                "PNch_Channel_Height":              oSec01Field01_value,
                "PNch_Frame_Height_Over_Channel":   oSec01Field02_value,
                "PNch_Channel_Width":               oSec01Field03_value,
                "PNch_Gate_Margin_Width":           oSec01Field04_value,
                "PNch_Water_Lv":                    oSec01Field05_value,
                "PNch_Gate_Margin_Over_Water_Lv":   oSec01Field06_value,
                "PNch_Gate_Th":                     oSec01Field07_value,
                "PNch_Gate_Other_PLs":              oSec01Field08_value,
                "PNch_HeadStock":                   oSec01Field09_value,
                "PNch_Frame_Weight_Per_M":          oSec01Field10_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["O_Frame_Perimeter"]
            oSec02Field02_result = response["O_Frame_Weight"]
            oSec02Field03_result = response["O_Gate_PL_Weight"]
            oSec02Field04_result = response["O_Gate_Stiffener_N"]
            oSec02Field05_result = response["O_Gate_Stiffener_Weight"]
            oSec02Field06_result = response["O_Gate_Weight"]
            oSec02Field07_result = response["O_Total_Weight"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result      
            instance.oSec02Field02 = oSec02Field02_result        
            instance.oSec02Field03 = oSec02Field03_result      
            instance.oSec02Field04 = oSec02Field04_result        
            instance.oSec02Field05 = oSec02Field05_result       
            instance.oSec02Field06 = oSec02Field06_result         
            instance.oSec02Field07 = oSec02Field07_result            
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "PNch"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcPNch(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,
                'oSec01Field07': oSec01Field07_value,
                'oSec01Field08': oSec01Field08_value,
                'oSec01Field09': oSec01Field09_value,
                'oSec01Field10': oSec01Field10_value,

                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
                'oSec02Field04': oSec02Field04_result,
                'oSec02Field05': oSec02Field05_result,
                'oSec02Field06': oSec02Field06_result,
                'oSec02Field07': oSec02Field07_result,
            })
            
            print(oSec01Field08_value)
            print(oSec01Field09_value)
            print(oSec01Field10_value)

            return render(request, 'PagePNch.html', {'form1': form1})

    return redirect('ms_load')  # Redirect to the page if the request is invalid

def generate_pnch_report(request):
    
    if request.method == "POST":
        form1 = formCalcPNch(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Channel Penstock Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
                (form1.fields["oSec01Field05"].label, request.POST.get("oSec01Field05", "N/A")),
                (form1.fields["oSec01Field06"].label, request.POST.get("oSec01Field06", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
                (form1.fields["oSec02Field04"].label, request.POST.get("oSec02Field04", "N/A")),
                (form1.fields["oSec02Field05"].label, request.POST.get("oSec02Field05", "N/A")),
                (form1.fields["oSec02Field06"].label, request.POST.get("oSec02Field06", "N/A")),
                (form1.fields["oSec02Field07"].label, request.POST.get("oSec02Field07", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Channel_Penstock_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)




def load_pnwa_page(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to login page if user is not authenticated

    result = check_user_autho(request.user.username, 'PNch')
    print('#####')
    print(result)
    print('######')


    form1 = formCalcPNwa()  

    return render(request, 'PagePNwa.html', {'form1': form1})

def handle_pnwa_form(request):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated

    print("aaa")
    if request.method == 'POST' and 'form1_submit' in request.POST:
                
        print("aaa")
        
        form1 = formCalcPNwa(request.POST)
        if form1.is_valid():
            
            print("form is valid")
            # Access the cleaned_data dictionary to get individual field values
            oSec01Field01_value = form1.cleaned_data.get('oSec01Field01')
            oSec01Field02_value = form1.cleaned_data.get('oSec01Field02')
            oSec01Field03_value = form1.cleaned_data.get('oSec01Field03')
            oSec01Field04_value = form1.cleaned_data.get('oSec01Field04')
            oSec01Field05_value = form1.cleaned_data.get('oSec01Field05')
            oSec01Field06_value = form1.cleaned_data.get('oSec01Field06')
            oSec01Field07_value = form1.cleaned_data.get('oSec01Field07')
            oSec01Field08_value = form1.cleaned_data.get('oSec01Field08')
            oSec01Field09_value = form1.cleaned_data.get('oSec01Field09')
            oSec01Field10_value = form1.cleaned_data.get('oSec01Field10')
            

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            req_type = "PNwa"  
            input_data = {
                "aInput01":   oSec01Field01_value,
                "aInput02":   oSec01Field02_value,
                "aInput03":   oSec01Field03_value,
                "aInput04":   oSec01Field04_value,
                "aInput05":   oSec01Field05_value,
                "aInput06":   oSec01Field06_value,
                "aInput07":   oSec01Field07_value,
                "aInput08":   oSec01Field08_value,
                "aInput09":   oSec01Field09_value,
                "aInput10":   oSec01Field10_value,
            }

            # Call the function to interact with the API
            response = interact_with_api(api_url, req_type, input_data)
            
            oSec02Field01_result = response["O_PNwa_Out01"]
            oSec02Field02_result = response["O_PNwa_Out02"]
            oSec02Field03_result = response["O_PNwa_Out03"]
            oSec02Field04_result = response["O_PNwa_Out04"]
            oSec02Field05_result = response["O_PNwa_Out05"]
            oSec02Field06_result = response["O_PNwa_Out06"]
            oSec02Field07_result = response["O_PNwa_Out07"]

            # Save the form with updated values
            instance = form1.save(commit=False)  # Do not save to the database yet
            instance.oSec02Field01 = oSec02Field01_result      
            instance.oSec02Field02 = oSec02Field02_result        
            instance.oSec02Field03 = oSec02Field03_result      
            instance.oSec02Field04 = oSec02Field04_result        
            instance.oSec02Field05 = oSec02Field05_result       
            instance.oSec02Field06 = oSec02Field06_result         
            instance.oSec02Field07 = oSec02Field07_result  
                      
            instance.oSec00Field01 = request.user.username  # Insert username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Insert current time
            instance.oSec00Field03 = "PNwa"  # Insert fixed type
            instance.save()  # Save to the database

            # Refresh the form with initial values to display results
            form1 = formCalcPNwa(initial={
                'oSec01Field01': oSec01Field01_value,
                'oSec01Field02': oSec01Field02_value,
                'oSec01Field03': oSec01Field03_value,
                'oSec01Field04': oSec01Field04_value,
                'oSec01Field05': oSec01Field05_value,
                'oSec01Field06': oSec01Field06_value,
                'oSec01Field07': oSec01Field07_value,
                'oSec01Field08': oSec01Field08_value,
                'oSec01Field09': oSec01Field09_value,
                'oSec01Field10': oSec01Field10_value,

                'oSec02Field01': oSec02Field01_result,
                'oSec02Field02': oSec02Field02_result,
                'oSec02Field03': oSec02Field03_result,
                'oSec02Field04': oSec02Field04_result,
                'oSec02Field05': oSec02Field05_result,
                'oSec02Field06': oSec02Field06_result,
                'oSec02Field07': oSec02Field07_result,
            })
                        
            return render(request, 'PagePNwa.html', {'form1': form1})

    return redirect('ms_load')  # Redirect to the page if the request is invalid

def generate_pnwa_report(request):
    
    if request.method == "POST":
        form1 = formCalcPNwa(request.POST)
        # Create a new Word document
        doc = Document()
        doc.add_heading('Wall Penstock Report', level=1)

        # Extract form data
        form_data = {
            "Input": [
                (form1.fields["oSec01Field01"].label, request.POST.get("oSec01Field01", "N/A")),
                (form1.fields["oSec01Field02"].label, request.POST.get("oSec01Field02", "N/A")),
                (form1.fields["oSec01Field03"].label, request.POST.get("oSec01Field03", "N/A")),
                (form1.fields["oSec01Field04"].label, request.POST.get("oSec01Field04", "N/A")),
                (form1.fields["oSec01Field05"].label, request.POST.get("oSec01Field05", "N/A")),
                (form1.fields["oSec01Field06"].label, request.POST.get("oSec01Field06", "N/A")),
                (form1.fields["oSec01Field07"].label, request.POST.get("oSec01Field07", "N/A")),
                (form1.fields["oSec01Field08"].label, request.POST.get("oSec01Field08", "N/A")),
                (form1.fields["oSec01Field09"].label, request.POST.get("oSec01Field09", "N/A")),
                (form1.fields["oSec01Field10"].label, request.POST.get("oSec01Field10", "N/A")),
            ],
            "Output": [
                (form1.fields["oSec02Field01"].label, request.POST.get("oSec02Field01", "N/A")),
                (form1.fields["oSec02Field02"].label, request.POST.get("oSec02Field02", "N/A")),
                (form1.fields["oSec02Field03"].label, request.POST.get("oSec02Field03", "N/A")),
                (form1.fields["oSec02Field04"].label, request.POST.get("oSec02Field04", "N/A")),
                (form1.fields["oSec02Field05"].label, request.POST.get("oSec02Field05", "N/A")),
                (form1.fields["oSec02Field06"].label, request.POST.get("oSec02Field06", "N/A")),
                (form1.fields["oSec02Field07"].label, request.POST.get("oSec02Field07", "N/A")),
            ]
        }

        # Add form data to the Word document
        for section, fields in form_data.items():
            doc.add_heading(section, level=2)
            for field, value in fields:
                doc.add_paragraph(f"{field}: {value}")

        # Prepare the response to download the document
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="Wall_Penstock_Report.docx"'
        doc.save(response)
        return response

    return HttpResponse("Invalid request", status=400)



############################
############################
############################
############################
############################
############################
############################
############################
############################
############################
############################
############################
############################
############################
############################
############################
############################
############################
############################
############################



def project_list(request):
    if request.method == "POST":
        form = ProjectForm(request.POST)
        if form.is_valid():
            
            instance = form.save(commit=False)  # Don't save to DB yet

            
            # Get the company associated with the user
            try:
                user_company = UserCompany.objects.get(user=request.user).company
                instance.company = user_company  # Assign company to the instance
            except UserCompany.DoesNotExist:
                return render(request, 'project_list', {"form": form, "error": "User is not associated with a company"})


            form.save()
            return redirect('project_list')
    else:
        form = ProjectForm()
        
        
    # Get the company of the logged-in user
    user_company = None
    if request.user.is_authenticated:
        try:
            user_company = UserCompany.objects.get(user=request.user).company
        except UserCompany.DoesNotExist:
            user_company = None

    # Assign company filter only if the user has a company
    if user_company:
        #machines = Machine.objects.filter(oSec00Field03=DB_Name, company=user_company)
        projects = Project.objects.filter(company=user_company)
    else:
        projects = Project.objects.none()  # Return an empty queryset if no company
    
    #projects = Project.objects.all()
    return render(request, 'project_list.html', {'form': form, 'projects': projects})
    
def edit_project(request, project_id):
    project = get_object_or_404(Project, id=project_id)
    
    if request.method == "POST":
        form = ProjectForm(request.POST, instance=project)
        if form.is_valid():
            form.save()
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({'success': True})
            return redirect('project_list')
        else:
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({'success': False, 'form': form.as_p()})
    else:
        form = ProjectForm(instance=project)
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({'success': True, 'form': form.as_p()})

    return render(request, 'edit_project.html', {'form': form})



def delete_project(request, project_id):
    if request.method == "POST":
        project = get_object_or_404(Project, id=project_id)
        project.delete()

        # Otherwise, redirect to the project list page
        return redirect('project_list')

    return JsonResponse({'success': False, 'error': 'Invalid request'})



def get_machines(request, project_id):
    try:
        project = Project.objects.get(id=project_id)
        machines = Machine.objects.filter(project=project)
        
        data = {
            "project_name": project.name,
            "machines": list(machines.values("oSec00Field01", "oSec00Field02", "oSec00Field03",
                                             "oSec01Field04", "oSec01Field05", "oSec01Field06",
                                             "oSec01Field07", "oSec01Field08", "oSec01Field09",
                                             "oSec01Field10", "oSec01Field11", "oSec01Field12",
                                             "oSec01Field13", "oSec01Field14", "oSec01Field15",
                                             "oSec01Field16", "oSec01Field17", "oSec01Field18",
                                             "oSec01Field19", "oSec01Field20"))
        }
        return JsonResponse(data)
    except Project.DoesNotExist:
        return JsonResponse({"error": "Project not found"}, status=404)












def generate_report(request, project_id):
    try:
        #pdb.set_trace()
        # Log the action
        aLogEntry.objects.create(user=request.user, message=f"at {now()} {request.user} accessed Word Report")
        
        #pdb.set_trace()
        # Get the user’s company and project
        aCompany = UserCompany.objects.get(user=request.user)

        #pdb.set_trace()
        # Determine the company and generate the corresponding report
        if aCompany.id == 1:
            print("Company 1")
            return generate_report_AAA(request, project_id)

        elif aCompany.id == 2:
            print("Company 2")
            return generate_report_BBB(request, project_id)

        else:
            return HttpResponse("Invalid company ID", status=400)

    except UserCompany.DoesNotExist:
        return HttpResponse("User does not belong to a company", status=403)

    except Project.DoesNotExist:
        return HttpResponse("Project not found", status=404)




def generate_report_AAA(request, project_id):
    
    def add_table(doc, data, title=None):
        """Creates a table and applies a background color to the header."""
        if title:
            doc.add_heading(title, level=3)

        table = doc.add_table(rows=len(data), cols=2)
        table.style = "Table Grid"

        for i, row in enumerate(data):
            for j, text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = text

                # Apply background color only to the header row (first row)
                if i == 0:
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(ns.qn("w:fill"), "FFA500")  # Orange color                    
                    #shading_elm.set(ns.qn("w:fill"), "ADD8E6")  # Blue color
                        
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    
            """Generates a Word report for a given project."""
    
    
    
    def add_header_footer(doc):
        """Adds header and footer with page numbers in the format 'Page X of Y'."""
        section = doc.sections[0]
    
        # Header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        # header_para.add_run("Company Name\n")
        # header_para.add_run("Project Name\n")
        # header_para.add_run("Date: " + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "\n")
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Adding logo
        run_logo = header_para.add_run()  # Corrected reference to header paragraph
        
        try:
            run_logo.add_picture("LogoAAA.PNG", width=Inches(7.0))  # Adjust width as needed
        except Exception as e:
            print(f"Error adding logo: {e}")

        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Add "Page X of Y" format
        run = footer_para.add_run("Page ")
    
        # PAGE field (Current Page Number)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(ns.qn("w:fldCharType"), "begin")
    
        instrText1 = OxmlElement("w:instrText")
        instrText1.set(ns.qn("xml:space"), "preserve")
        instrText1.text = "PAGE"
    
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar1)
        run._r.append(instrText1)
        run._r.append(fldChar2)
    
        run.add_text(" of ")
    
        # NUMPAGES field (Total Number of Pages)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(ns.qn("w:fldCharType"), "begin")
    
        instrText2 = OxmlElement("w:instrText")
        instrText2.set(ns.qn("xml:space"), "preserve")
        instrText2.text = "NUMPAGES"
    
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)
    
    
    def add_colored_heading(doc, text, level, color):
        """Adds a heading with color."""
        heading = doc.add_paragraph()
        run = heading.add_run(text)
        run.bold = True
        run.font.size = Pt(14) if level == 1 else Pt(12)
        run.font.color.rgb = color
        heading.style = f"Heading {level}"
    
    
    
    try:
        
        ###LOG
        aLogEntry.objects.create(
                user=request.user,
                message=f"at {now()} {request.user} accessed Load  "
            )
        print(f"at {now()} {User} accessed Download Report")
        ###LOG

        aCompany = UserCompany.objects.get(user=request.user)
        project = Project.objects.get(id=project_id)
        machines = Machine.objects.filter(project=project)
        
        
        print(aCompany.id)
        print(project.id)
    
        print("Company 1")
    
    
        # Create a Word document
        doc = Document()

        # Add header and footer with page numbers
        add_header_footer(doc)

        # Add project title
        doc.add_heading(f'Project Report: {project.name}', level=1)

        # Add project details
        doc.add_heading("Project Details", level=2)     

        doc.add_paragraph("\n")
        doc.add_paragraph("Name: " + project.name)
        doc.add_paragraph("Client Name: " + project.client_name)
        doc.add_paragraph("Capacity: " + project.capacity)
        doc.add_paragraph("\n")
        
        doc.add_page_break()     
        doc.add_paragraph("\n")

        # Add machine details
        for index, machine in enumerate(machines, start=1):  # Add numbering
            machine_name = machine.oSec00Field03
            section_titles = []

            if machine_name == "DataSheetNS":
                machine_name = "Manual Screen" 
                section_titles = ["General Data", "Design Data", "Material Data", "Channel Data", " ", " ", " ", " ", " ", " "]

            if machine_name == "DataSheetMSc":
                machine_name = "Mechanical Screen" 
                section_titles = ["General Data", "Design Data", "Gearmotor Data", "Control panel Data", "Material Data", "Other Data", " ", " ", " ", " "]

            if machine_name == "DataSheetBC":
                machine_name = "Belt Conveyor"
                section_titles = ["General Data", "Design Data", "Gearbox Data", "Motor Data", "Material Data", " ", " ", " ", " ", " "]

            if machine_name == "DataSheetCO":
                machine_name = "Container"
                section_titles = ["General Data", "Design Data", "Material Data", " ", " ", " ", " ", " ", " ", " "]

            if machine_name == "DataSheetGR":
                machine_name = "Gritremoval"
                section_titles = ["General Data", "Design Data", "Walkway, Handrail, Wheel Data", "Scrapper Data", "Gearmotor Data", "Scrapper Data", "Drive unit", "Control panel Data", "Material Data ", " "]

            if machine_name == "DataSheetSS":
                machine_name = "Sand Silo"

            if machine_name == "DataSheetPS":
                machine_name = "Primary Sedimentation"

            if machine_name == "DataSheetQV":
                machine_name = "Quick Valve"

            if machine_name == "DataSheetTV":
                machine_name = "Telescopic Valve"
                
            if machine_name == "DataSheetTH":
                machine_name = "Sludge Thickener"

            # Add machine title with font size 14 and numbering
            machine_title = doc.add_paragraph(f"{index}. {machine_name}", style="Heading3")
            machine_title.runs[0].font.size = Pt(14)

            for i in range(1, 11):  # Loop from Sec01 to Sec10
                section_name = f"Sec{i:02d}"
                section_data = [("Field", "Value")]

                for j in range(1, 21, 2):  # Step by 2 to avoid duplication
                    key = getattr(machine, f"o{section_name}Field{j:02d}", "").strip()
                    value = getattr(machine, f"o{section_name}Field{j+1:02d}", "").strip()

                    if key and value and key.lower() != "oooo" and value.lower() != "oooo":
                        section_data.append((key, value))

                if len(section_data) > 1:  # If the section has valid data, create a table
                    section_title = section_titles[i-1] if i-1 < len(section_titles) else f"Section {i}"
                    doc.add_paragraph(f"{section_name}: {section_title}", style="Heading3")  # Only one title now

                    add_table(doc, section_data)  # Removed redundant title

            doc.add_page_break() 
        
        # Save the document to a response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={project.name}_report.docx'
        doc.save(response)
        return response

    except Project.DoesNotExist:
        return HttpResponse("Project not found", status=404)





def generate_report_BBB(request, project_id):
    
    def add_table(doc, data, title=None):
        """Creates a borderless table and applies a background color to the header."""
        if title:
            doc.add_heading(title, level=3)

        table = doc.add_table(rows=len(data), cols=2)

        # Remove all table borders manually
        tbl = table._tbl  # Get the table's XML element
        tblPr = tbl.find(qn("w:tblPr"))  # Find existing table properties

        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")  # Create table properties if missing
            tbl.insert(0, tblPr)  # Insert as the first child of <w:tbl>

        tblBorders = OxmlElement("w:tblBorders")  # Create <w:tblBorders>
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")  # Remove the border
            tblBorders.append(border)

        tblPr.append(tblBorders)  # Append border settings to the table properties

        for i, row in enumerate(data):
            for j, text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = text

                # Apply background color only to the header row (first row)
                if i == 0:
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:val"), "clear")  # Set shading value
                    shading_elm.set(qn("w:fill"), "ADD8E6")  # Light blue color
                    cell._tc.get_or_add_tcPr().append(shading_elm)
   
    
    
    def add_header_footer(doc):
        """Adds header and footer with page numbers in the format 'Page X of Y'."""
        section = doc.sections[0]
    
        # Header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        # header_para.add_run("Company Name\n")
        # header_para.add_run("Project Name\n")
        # header_para.add_run("Date: " + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "\n")
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Adding logo
        run_logo = header_para.add_run()  # Corrected reference to header paragraph
        try:
            run_logo.add_picture("LogoBBB.PNG", width=Inches(7.0))  # Adjust width as needed
        except Exception as e:
            print(f"Error adding logo: {e}")

        
    
        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Add "Page X of Y" format
        run = footer_para.add_run("Page ")
    
        # PAGE field (Current Page Number)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(ns.qn("w:fldCharType"), "begin")
    
        instrText1 = OxmlElement("w:instrText")
        instrText1.set(ns.qn("xml:space"), "preserve")
        instrText1.text = "PAGE"
    
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar1)
        run._r.append(instrText1)
        run._r.append(fldChar2)
    
        run.add_text(" of ")
    
        # NUMPAGES field (Total Number of Pages)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(ns.qn("w:fldCharType"), "begin")
    
        instrText2 = OxmlElement("w:instrText")
        instrText2.set(ns.qn("xml:space"), "preserve")
        instrText2.text = "NUMPAGES"
    
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)
    
    
    def add_colored_heading(doc, text, level, color):
        """Adds a heading with color."""
        heading = doc.add_paragraph()
        run = heading.add_run(text)
        run.bold = True
        run.font.size = Pt(14) if level == 1 else Pt(12)
        run.font.color.rgb = color
        heading.style = f"Heading {level}"
    
    
    
    try:
        
        ###LOG
        aLogEntry.objects.create(
                user=request.user,
                message=f"at {now()} {request.user} accessed Load  "
            )
        print(f"at {now()} {User} accessed Download Report")
        ###LOG

        aCompany = UserCompany.objects.get(user=request.user)
        project = Project.objects.get(id=project_id)
        machines = Machine.objects.filter(project=project)
        
        
        print(aCompany.id)
        print(project.id)
    
        print("Company 2")
    
    
        # Create a Word document
        doc = Document()

        # Add header and footer with page numbers
        add_header_footer(doc)

        # Add project title
        doc.add_heading(f'Project Report: {project.name}', level=1)

        # Add project details
        doc.add_heading("Project Details", level=2)     

        doc.add_paragraph("\n")
        doc.add_paragraph("Name: " + project.name)
        doc.add_paragraph("Client Name: " + project.client_name)
        doc.add_paragraph("Capacity: " + project.capacity)
        doc.add_paragraph("\n")
        
        doc.add_page_break()     
        doc.add_paragraph("\n")

        # Add machine details
        for index, machine in enumerate(machines, start=1):  # Add numbering
            machine_name = machine.oSec00Field03
            section_titles = []

            if machine_name == "DataSheetNS":
                machine_name = "Manual Screen" 
                section_titles = ["General Data", "Design Data", "Material Data", "Channel Data", " ", " ", " ", " ", " ", " "]

            if machine_name == "DataSheetMSc":
                machine_name = "Mechanical Screen" 
                section_titles = ["General Data", "Design Data", "Gearmotor Data", "Control panel Data", "Material Data", "Other Data", " ", " ", " ", " "]

            if machine_name == "DataSheetBC":
                machine_name = "Belt Conveyor"
                section_titles = ["General Data", "Design Data", "Gearbox Data", "Motor Data", "Material Data", " ", " ", " ", " ", " "]

            if machine_name == "DataSheetCO":
                machine_name = "Container"
                section_titles = ["General Data", "Design Data", "Material Data", " ", " ", " ", " ", " ", " ", " "]

            if machine_name == "DataSheetGR":
                machine_name = "Gritremoval"
                section_titles = ["General Data", "Design Data", "Walkway, Handrail, Wheel Data", "Scrapper Data", "Gearmotor Data", "Scrapper Data", "Drive unit", "Control panel Data", "Material Data ", " "]

            if machine_name == "DataSheetSS":
                machine_name = "Sand Silo"

            if machine_name == "DataSheetPS":
                machine_name = "Primary Sedimentation"

            if machine_name == "DataSheetQV":
                machine_name = "Quick Valve"

            if machine_name == "DataSheetTV":
                machine_name = "Telescopic Valve"
                
            if machine_name == "DataSheetTH":
                machine_name = "Sludge Thickener"

            # Add machine title with font size 14 and numbering
            machine_title = doc.add_paragraph(f"{index}. {machine_name}", style="Heading3")
            machine_title.runs[0].font.size = Pt(14)

            for i in range(1, 11):  # Loop from Sec01 to Sec10
                section_name = f"Sec{i:02d}"
                section_data = [("Field", "Value")]

                for j in range(1, 21, 2):  # Step by 2 to avoid duplication
                    key = getattr(machine, f"o{section_name}Field{j:02d}", "").strip()
                    value = getattr(machine, f"o{section_name}Field{j+1:02d}", "").strip()

                    if key and value and key.lower() != "oooo" and value.lower() != "oooo":
                        section_data.append((key, value))

                if len(section_data) > 1:  # If the section has valid data, create a table
                    section_title = section_titles[i-1] if i-1 < len(section_titles) else f"Section {i}"
                    doc.add_paragraph(f"{section_name}: {section_title}", style="Heading3")  # Only one title now

                    add_table(doc, section_data)  # Removed redundant title

            doc.add_page_break()     

        # Save the document to a response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={project.name}_report.docx'
        doc.save(response)
        return response

    except Project.DoesNotExist:
        return HttpResponse("Project not found", status=404)



###
###




###
###






#######
#######
#######
#######
#######
#######
#######
#######


# Helper function to get the user's company
def get_user_company(request):
    if request.user.is_authenticated:
        try:
            return UserCompany.objects.get(user=request.user).company
        except UserCompany.DoesNotExist:
            return None
    return None

# Helper function to define DXF paths
def get_dxf_paths(user_company, category):
    company_paths = {
        1: os.path.join(settings.BASE_DIR, "static", "aDxfs", "AAA", f"AAA_{category}.dxf"),
        2: os.path.join(settings.BASE_DIR, "static", "aDxfs", "BBB", f"{category}.dxf"),
    }

    static_path = company_paths.get(user_company.id, "")
    modified_path = static_path.replace(".dxf", "_new.dxf")

    return static_path, modified_path

# Helper function to modify DXF files
def modify_dxf_file(static_path, modified_path, modifications):
    doc = ezdxf.readfile(static_path)

    for entity in doc.modelspace().query("DIMENSION"):
        if entity.dxf.text in modifications:
            entity.dxf.text = modifications[entity.dxf.text]

        # Update text height and arrow size
        dimstyle = doc.dimstyles.get(entity.dxf.dimstyle)
        if dimstyle:
            dimstyle.dxf.dimtxt = 0.1  # Set text height
            dimstyle.dxf.dimasz = 0.1  # Set arrow size

        entity.render()

    doc.saveas(modified_path)

# Main DXF Processing Function
def process_dxf(request, aMachine_ID, category, modifications, output_filename):
    # Log the request
    aLogEntry.objects.create(
        user=request.user,
        message=f"at {now()} {request.user} Download DXF {category} {aMachine_ID}"
    )

    user_company = get_user_company(request)
    if not user_company:
        return HttpResponse("Unauthorized", status=403)

    static_path, modified_path = get_dxf_paths(user_company, category)
    if not os.path.exists(static_path):
        return HttpResponse("File not found", status=404)

    machine = Machine.objects.get(id=aMachine_ID)

    if request.method == "POST":
        modify_dxf_file(static_path, modified_path, modifications(machine))

        # Serve the modified file for download
        with open(modified_path, "rb") as dxf_file:
            response = HttpResponse(dxf_file.read(), content_type="application/dxf")
            response["Content-Disposition"] = f'attachment; filename="{output_filename}"'
            return response

    return HttpResponse("Invalid request", status=400)

# DXF Download Views
def General_DXF_ALL(request, aMachine_ID, aType):
    
    print(aType)
    
    # Redirect unauthenticated users
    if not request.user.is_authenticated:
        return redirect("login")  
    
    
    ###LOG
    aLogEntry.objects.create(
            user=request.user,
            message=f"at {now()} {request.user} DXF download {aType} "
        )
    
    
    # Get the company of the logged-in user    
    user_company = None
    if request.user.is_authenticated:
        try:
            user_company = UserCompany.objects.get(user=request.user).company
        except UserCompany.DoesNotExist:
            user_company = None

    
    if aType == "NS":
        return process_dxf(
            request,
            aMachine_ID,
            "NS",
            lambda machine: {
                "ScreenLength": machine.oSec02Field06,
                "BarLength": "500",
                "ScreenWidth": machine.oSec02Field04,
                "BarTh": "10",
                "BarSpacing": machine.oSec02Field10,
            },
            f"new_NS_{user_company}.dxf"
        )
        
    
    if aType == "MS":
        return process_dxf(
            request,
            aMachine_ID,
            "MS",
            lambda machine: {
                "ChannelHeight": "0000",
                "WaterLevel": "000",
                "Width": machine.oSec02Field08,
                "Length": machine.oSec02Field10,
                "Angle": machine.oSec02Field20,
            },
            f"new_MS_{user_company}.dxf"
        )
        
    if aType == "BC":
        return process_dxf(
            request,
            aMachine_ID,
            "BC",
            lambda machine: {
                "Length": machine.oSec02Field04,
                "Width": machine.oSec02Field02,
                "WidB": "000",
                "BarTh": "10",
                "BarSpacing": "25",
            },
            f"new_BC_{user_company}.dxf"
        )
        
    if aType == "CO":
        return process_dxf(
            request,
            aMachine_ID,
            "CO",
            lambda machine: {
                "ScreenLength": "1000",
                "BarLength": "500",
                "ScreenWidth": "600",
                "BarTh": "10",
                "BarSpacing": "25",
            },
            f"new_CO_{user_company}.dxf"
        )
        
    if aType == "GR":
        return process_dxf(
            request,
            aMachine_ID,
            "GR",
            lambda machine: {
                "ScreenLength": "1000",
                "BarLength": "500",
                "ScreenWidth": "600",
                "BarTh": "10",
                "BarSpacing": "25",
            },
            f"new_GR_{user_company}.dxf"
        )
        
    if aType == "SS":
        return process_dxf(
            request,
            aMachine_ID,
            "SS",
            lambda machine: {
                "ScreenLength": "1000",
                "BarLength": "500",
                "ScreenWidth": "600",
                "BarTh": "10",
                "BarSpacing": "25",
            },
            f"new_SS_{user_company}.dxf"
        )
        
    if aType == "PS":
        return process_dxf(
            request,
            aMachine_ID,
            "PS",
            lambda machine: {
                "ScreenLength": "1000",
                "BarLength": "500",
                "ScreenWidth": "600",
                "BarTh": "10",
                "BarSpacing": "25",
            },
            f"new_PS_{user_company}.dxf"
        )
        
    if aType == "QV":
        return process_dxf(
            request,
            aMachine_ID,
            "QV",
            lambda machine: {
                "ScreenLength": "1000",
                "BarLength": "500",
                "ScreenWidth": "600",
                "BarTh": "10",
                "BarSpacing": "25",
            },
            f"new_QV_{user_company}.dxf"
        )
        
    if aType == "TV":
        return process_dxf(
            request,
            aMachine_ID,
            "TV",
            lambda machine: {
                "ScreenLength": "1000",
                "BarLength": "500",
                "ScreenWidth": "600",
                "BarTh": "10",
                "BarSpacing": "25",
            },
            f"new_TV_{user_company}.dxf"
        )
        
    if aType == "TH":
        return process_dxf(
            request,
            aMachine_ID,
            "TH",
            lambda machine: {
                "ScreenLength": "1000",
                "BarLength": "500",
                "ScreenWidth": "600",
                "BarTh": "10",
                "BarSpacing": "25",
            },
            f"new_TH_{user_company}.dxf"
        )
        
    
    
    
    
  

# DXF Download Views
def FullDrawing(request, aMachine_ID, aType):
    
    # Helper function to modify DXF files
    def FullDrawing_modify_dxf_file(static_path, modified_path, modifications):
        doc = ezdxf.readfile(static_path)

        # for entity in doc.modelspace().query("DIMENSION"):
        #     if entity.dxf.text in modifications:
        #         entity.dxf.text = modifications[entity.dxf.text]

        #     # Update text height and arrow size
        #     dimstyle = doc.dimstyles.get(entity.dxf.dimstyle)
        #     if dimstyle:
        #         dimstyle.dxf.dimtxt = 0.1  # Set text height
        #         dimstyle.dxf.dimasz = 0.1  # Set arrow size

        #     entity.render()

        doc.saveas(modified_path)    
        
    
    # Main DXF Processing Function
    def FullDrawing_process_dxf(request, aMachine_ID, category, modifications, output_filename):

        user_company = get_user_company(request)
        if not user_company:
            return HttpResponse("Unauthorized", status=403)

        static_path  = os.path.join(settings.BASE_DIR, "static", "aDxfs", "AAA", "FullDrawing", "Full Drawing NS.dxf")
        modified_path = static_path.replace(".dxf", "_newFullDrawing.dxf")
        
        if not os.path.exists(static_path):
            return HttpResponse("File not found", status=404)

        machine = Machine.objects.get(id=aMachine_ID)

        if request.method == "POST":
            FullDrawing_modify_dxf_file(static_path, modified_path, modifications(machine))

            # Serve the modified file for download
            with open(modified_path, "rb") as dxf_file:
                response = HttpResponse(dxf_file.read(), content_type="application/dxf")
                response["Content-Disposition"] = f'attachment; filename="{output_filename}"'
                return response

        return HttpResponse("Invalid request", status=400)
        

    
    # Redirect unauthenticated users
    if not request.user.is_authenticated:
        return redirect("login") 
        
    # Get the company of the logged-in user    
    user_company = None
    if request.user.is_authenticated:
        try:
            user_company = UserCompany.objects.get(user=request.user).company
        except UserCompany.DoesNotExist:
            user_company = None 
        
    ###LOG
    aLogEntry.objects.create(
            user=request.user,
            message=f"at {now()} {request.user} DXF download {aType} "
        )

    
    #if aType == "NS":
    return FullDrawing_process_dxf(
        request,
        aMachine_ID,
        "NS",
        lambda machine: {
            "ScreenLength": machine.oSec02Field06,
            "BarLength": "500",
            "ScreenWidth": machine.oSec02Field04,
            "BarTh": "10",
            "BarSpacing": machine.oSec02Field10,
        },
        f"newFullDrawing_ManualScreen_{user_company}.dxf"
    )
        
    
    


###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###
###






####


def LoadPageDataSheet(request, sheet_key):
    #pdb.set_trace()
    print(sheet_key)
    
    # Redirect unauthenticated users
    if not request.user.is_authenticated:
        return redirect("login")  
    
    print(request.user)
    print(f"{request.user} accessed Load {sheet_key}")
    ###LOG
    
    aLogEntry.objects.create(
        user=request.user,
        message=f"{request.user} >>> {sheet_key}"
    )
    
    
    
    # Get the company of the logged-in user    
    user_company = None
    if request.user.is_authenticated:
        try:
            user_company = UserCompany.objects.get(user=request.user).company
        except UserCompany.DoesNotExist:
            user_company = None

    print(user_company)




    # Define a dictionary mapping sheet keys to their corresponding values
    sheet_mapping = {
        "NS": ("formDataSheetNS",               "DataSheetNS",      "Manual Screen"),
        "MS": ("FDS_MechanicalScreen",          "DataSheetMS",      "Mechanical Screen"),
        "BC": ("FDS_BeltConveyor",              "DataSheetBC",      "Belt Conveyor"),
        "CO": ("FDS_Container",                 "DataSheetCO",      "Container"),
        "GR": ("FDS_GritGreaseRemoval",         "DataSheetGR",      "Gritremoval"),
        "SS": ("FDS_SandSilo",                  "DataSheetSS",      "Sand Silo"),
        "PS": ("FDS_PrimarySedimentationTank",  "DataSheetPS",      "Primary Sedimentation Tank"),
        "QV": ("FDS_QuickValve",                "DataSheetQV",      "Quick Valve"),
        "TV": ("FDS_TelescopicValve",           "DataSheetTV",      "Telescopic Valve"),
        "TH": ("FDS_SludgeThickener",           "DataSheetTH",      "Thickener"),
    }

    # Retrieve values using the dictionary
    form_type, DB_Name, aMachineName = sheet_mapping.get(sheet_key, ("None", "None", "None"))

    # Optional: Handle cases where the sheet_key is invalid
    if form_type is None:
        print(f"Warning: Unknown sheet_key '{sheet_key}'")




    # Assign company filter only if the user has a company
    if user_company:
        machines = Machine.objects.filter(oSec00Field03=DB_Name, company=user_company)
        projects = Project.objects.filter(company=user_company)
    else:
        machines = Machine.objects.none()  # Return an empty queryset if no company
        projects = Project.objects.none()  # Return an empty queryset if no company



    form = FormDataSheet(user=request.user, form_type=form_type)
    
    print(f"Initial value for oSec01Field02: {form.fields['oSec01Field02'].initial}")
    
    
    # Initialize all section variables
    aSection01Show = "Yes"
    aSection02Show = "Yes"
    aSection03Show = "Yes"
    aSection04Show = "Yes"
    aSection05Show = "Yes"
    aSection06Show = "Yes"
    aSection07Show = "Yes"
    aSection08Show = "Yes"
    aSection09Show = "Yes"
    aSection10Show = "Yes"
    
    print(form.fields['oSec01Field01'].initial)
    print(form.fields['oSec02Field01'].initial)
    print(form.fields['oSec03Field01'].initial)
    print(form.fields['oSec04Field01'].initial)
    print(form.fields['oSec05Field01'].initial)
    print(form.fields['oSec06Field01'].initial)
    print(form.fields['oSec07Field01'].initial)
    print(form.fields['oSec08Field01'].initial)
    print(form.fields['oSec09Field01'].initial)
    print(form.fields['oSec10Field01'].initial)

    # Apply conditions to modify the values
    if form.fields['oSec01Field01'].initial in ["oooo", None]:
        aSection01Show = "Hide"

    if form.fields['oSec02Field01'].initial in ["oooo", None]:
        aSection02Show = "Hide"

    if form.fields['oSec03Field01'].initial in ["oooo", None]:
        aSection03Show = "Hide"

    if form.fields['oSec04Field01'].initial in ["oooo", None]:
        aSection04Show = "Hide"

    if form.fields['oSec05Field01'].initial in ["oooo", None]:
        aSection05Show = "Hide"

    if form.fields['oSec06Field01'].initial in ["oooo", None]:
        aSection06Show = "Hide"

    if form.fields['oSec07Field01'].initial in ["oooo", None]:
        aSection07Show = "Hide"

    if form.fields['oSec08Field01'].initial in ["oooo", None]:
        aSection08Show = "Hide"

    if form.fields['oSec09Field01'].initial in ["oooo", None]:
        aSection09Show = "Hide"

    if form.fields['oSec10Field01'].initial in ["oooo", None]:
        aSection10Show = "Hide"
    
    print(aSection01Show)
    print(aSection02Show)
    print(aSection03Show)
    print(aSection04Show)
    print(aSection05Show)
    print(aSection06Show)
    print(aSection07Show)
    print(aSection08Show)
    print(aSection09Show)
    print(aSection10Show)
    
    # print(projects)

    return render(request, "PageDataSheet.html", {
    "form": form,
    "machines": machines,
    "projects": projects,  
    "aMachineName": aMachineName,  
    "user_company": user_company,
    "sheet_key": sheet_key,
    "aSection01Show": aSection01Show,
    "aSection02Show": aSection02Show,
    "aSection03Show": aSection03Show,
    "aSection04Show": aSection04Show,
    "aSection05Show": aSection05Show,
    "aSection06Show": aSection06Show,
    "aSection07Show": aSection07Show,
    "aSection08Show": aSection08Show,
    "aSection09Show": aSection09Show,
    "aSection10Show": aSection10Show,
})




def SavePageDataSheet(request, sheet_key):    
    print(sheet_key)
    
    # Redirect unauthenticated users
    if not request.user.is_authenticated:
        return redirect("login")  
    
    
    ###LOG
    aLogEntry.objects.create(
            user=request.user,
            message=f"at {now()} {request.user} accessed Load {sheet_key} "
        )
    
    
    # Get the company of the logged-in user    
    user_company = None
    if request.user.is_authenticated:
        try:
            user_company = UserCompany.objects.get(user=request.user).company
        except UserCompany.DoesNotExist:
            user_company = None


    
    # Define a dictionary mapping sheet keys to their corresponding values
    sheet_mapping = {
        "NS": ("formDataSheetNS",               "DataSheetNS",      "Manual Screen"),
        "MS": ("FDS_MechanicalScreen",          "DataSheetMS",      "Mechanical Screen"),
        "BC": ("FDS_BeltConveyor",              "DataSheetBC",      "Belt Conveyor"),
        "CO": ("FDS_Container",                 "DataSheetCO",      "Container"),
        "GR": ("FDS_GritGreaseRemoval",         "DataSheetGR",      "Gritremoval"),
        "SS": ("FDS_SandSilo",                  "DataSheetSS",      "Sand Silo"),
        "PS": ("FDS_PrimarySedimentationTank",  "DataSheetPS",      "Primary Sedimentation Tank"),
        "QV": ("FDS_QuickValve",                "DataSheetQV",      "Quick Valve"),
        "TV": ("FDS_TelescopicValve",           "DataSheetTV",      "Telescopic Valve"),
        "TH": ("FDS_SludgeThickener",           "DataSheetTH",      "Thickener"),
    }

    # Retrieve values using the dictionary
    form_type, DB_Name, aMachineName = sheet_mapping.get(sheet_key, ("None", "None", "None"))
    


    # Assign company filter only if the user has a company
    if user_company:
        machines = Machine.objects.filter(oSec00Field03=DB_Name, company=user_company)
        projects = Project.objects.filter(company=user_company)
    else:
        machines = Machine.objects.none()  # Return an empty queryset if no company
        projects = Project.objects.none()  # Return an empty queryset if no company

    print(form_type)
    

    if request.method == "POST":
        form = FormDataSheet(form_type=form_type, data=request.POST)

        if form.is_valid():
            instance = form.save(commit=False)  # Don't save to DB yet

            # Assign common fields
            instance.oSec00Field01 = request.user.username  # Username
            instance.oSec00Field02 = datetime.now().strftime("%Y-%m-%d %H:%M:%S")  # Timestamp
            instance.oSec00Field03 = DB_Name  # Fixed type

            # Handle project assignment
            project_id = request.POST.get("project")
            if project_id:
                try:
                    instance.project = Project.objects.get(id=project_id)
                except Project.DoesNotExist:
                    return render(request, "PageDataSheet.html", {"form": form, "error": "Invalid Project ID"})
            else:
                return render(request, "PageDataSheet.html", {"form": form, "error": "Project is required"})

            # Get the company associated with the user
            try:
                user_company = UserCompany.objects.get(user=request.user).company
                instance.company = user_company  # Assign company to the instance
            except UserCompany.DoesNotExist:
                return render(request, "PageDataSheet.html", 
                              {"form": form, 
                               "error": "User is not associated with a company",
                               "aMachineName": aMachineName,
                               "sheet_key" : sheet_key})

            # Save the instance to the database
            instance.save()

            # Refresh form with initial values
            form = FormDataSheet(initial=form.cleaned_data, form_type=form_type)

            # Filter machines by the user’s company
            machines = Machine.objects.filter(oSec00Field03=DB_Name, company=user_company)

            #######################################
            #######################################
            #######################################
            print("#######################")
        
            # Initialize all section variables
            aSection01Show = "Yes"
            aSection02Show = "Yes"
            aSection03Show = "Yes"
            aSection04Show = "Yes"
            aSection05Show = "Yes"
            aSection06Show = "Yes"
            aSection07Show = "Yes"
            aSection08Show = "Yes"
            aSection09Show = "Yes"
            aSection10Show = "Yes"
            
            print(form.fields['oSec01Field01'].initial)
            print(form.fields['oSec02Field01'].initial)
            print(form.fields['oSec03Field01'].initial)
            print(form.fields['oSec04Field01'].initial)
            print(form.fields['oSec05Field01'].initial)
            print(form.fields['oSec06Field01'].initial)
            print(form.fields['oSec07Field01'].initial)
            print(form.fields['oSec08Field01'].initial)
            print(form.fields['oSec09Field01'].initial)
            print(form.fields['oSec10Field01'].initial)
        
            # Apply conditions to modify the values
            if form.fields['oSec01Field01'].initial in ["oooo", None]:
                aSection01Show = "Hide"
        
            if form.fields['oSec02Field01'].initial in ["oooo", None]:
                aSection02Show = "Hide"
        
            if form.fields['oSec03Field01'].initial in ["oooo", None]:
                aSection03Show = "Hide"
        
            if form.fields['oSec04Field01'].initial in ["oooo", None]:
                aSection04Show = "Hide"
        
            if form.fields['oSec05Field01'].initial in ["oooo", None]:
                aSection05Show = "Hide"
        
            if form.fields['oSec06Field01'].initial in ["oooo", None]:
                aSection06Show = "Hide"
        
            if form.fields['oSec07Field01'].initial in ["oooo", None]:
                aSection07Show = "Hide"
        
            if form.fields['oSec08Field01'].initial in ["oooo", None]:
                aSection08Show = "Hide"
        
            if form.fields['oSec09Field01'].initial in ["oooo", None]:
                aSection09Show = "Hide"
        
            if form.fields['oSec10Field01'].initial in ["oooo", None]:
                aSection10Show = "Hide"
            
            print(aSection01Show)
            print(aSection02Show)
            print(aSection03Show)
            print(aSection04Show)
            print(aSection05Show)
            print(aSection06Show)
            print(aSection07Show)
            print(aSection08Show)
            print(aSection09Show)
            print(aSection10Show)
            
            

            return render(request, "PageDataSheet.html", {
                "form": form,
                "machines": machines,
                "projects": projects,  
                "aMachineName": aMachineName,  
                "user_company": user_company,
                "sheet_key": sheet_key,
                "aSection01Show": aSection01Show,
                "aSection02Show": aSection02Show,
                "aSection03Show": aSection03Show,
                "aSection04Show": aSection04Show,
                "aSection05Show": aSection05Show,
                "aSection06Show": aSection06Show,
                "aSection07Show": aSection07Show,
                "aSection08Show": aSection08Show,
                "aSection09Show": aSection09Show,
                "aSection10Show": aSection10Show,
            })

        else:
            # If the form has errors, return all machines for this DB_Name (no company filtering)
            machines = Machine.objects.filter(oSec00Field03=DB_Name)
            return render(request, "PageDataSheet.html", {"form": form, "error": "Form contains errors", "machines": machines})

    return redirect("ms_load")  # Redirect for invalid requests



def DeleteMachine(request, machine_id, aType):  
    machine = get_object_or_404(Machine, id=machine_id)
    machine.delete()

    # Redirect after deletion (if needed)
    return redirect(reverse('PageDataSheet', kwargs={'sheet_key': aType}))



def edit_machine(request, machine_id):
    machine = get_object_or_404(Machine, id=machine_id)  # Fetch the existing machine

    if request.method == "POST":
        #form = FDS_CO(request.POST)  # Bind the form with posted data
        form = FormDataSheet(data=request.POST)
        
        if form.is_valid():
            # Manually update machine fields
            machine.project = form.cleaned_data.get('project', None)
            machine.oSec01Field01 = form.cleaned_data.get('oSec01Field01', '')
            machine.oSec01Field02 = form.cleaned_data.get('oSec01Field02', '')
            machine.oSec01Field03 = form.cleaned_data.get('oSec01Field03', '')
            machine.oSec01Field04 = form.cleaned_data.get('oSec01Field04', '')
            machine.oSec01Field05 = form.cleaned_data.get('oSec01Field05', '')
            machine.oSec01Field06 = form.cleaned_data.get('oSec01Field06', '')
            machine.oSec01Field07 = form.cleaned_data.get('oSec01Field07', '')
            machine.oSec01Field08 = form.cleaned_data.get('oSec01Field08', '')
            machine.oSec01Field09 = form.cleaned_data.get('oSec01Field09', '')
            machine.oSec01Field10 = form.cleaned_data.get('oSec01Field10', '')
            machine.oSec01Field11 = form.cleaned_data.get('oSec01Field11', '')
            machine.oSec01Field12 = form.cleaned_data.get('oSec01Field12', '')
            machine.oSec01Field13 = form.cleaned_data.get('oSec01Field13', '')
            machine.oSec01Field14 = form.cleaned_data.get('oSec01Field14', '')
            machine.oSec01Field15 = form.cleaned_data.get('oSec01Field15', '')
            machine.oSec01Field16 = form.cleaned_data.get('oSec01Field16', '')
            machine.oSec01Field17 = form.cleaned_data.get('oSec01Field17', '')
            machine.oSec01Field18 = form.cleaned_data.get('oSec01Field18', '')
            machine.oSec01Field19 = form.cleaned_data.get('oSec01Field19', '')
            machine.oSec01Field20 = form.cleaned_data.get('oSec01Field20', '')

            machine.save()  # Save updates to the database
            return JsonResponse({"success": True})
        else:
            return JsonResponse({"success": False, "errors": form.errors})  # Send form validation errors

    return JsonResponse({"success": False, "error": "Invalid request"})




def DataSheetNS_get_datasheet_data(request, machine_id):
    machine = get_object_or_404(Machine, id=machine_id)
    
    data = {
        "project": machine.project.name if machine.project else "No Project",
        "oSec01Field01": machine.oSec01Field01,
        "oSec01Field02": machine.oSec01Field02,
        "oSec01Field03": machine.oSec01Field03,
        "oSec01Field04": machine.oSec01Field04,
        "oSec01Field05": machine.oSec01Field05,
        "oSec01Field06": machine.oSec01Field06,
        "oSec01Field07": machine.oSec01Field07,
        "oSec01Field08": machine.oSec01Field08,
        "oSec01Field09": machine.oSec01Field09,
        "oSec01Field10": machine.oSec01Field10,        
        "oSec01Field11": machine.oSec01Field11,
        "oSec01Field12": machine.oSec01Field12,
        "oSec01Field13": machine.oSec01Field13,
        "oSec01Field14": machine.oSec01Field14,
        "oSec01Field15": machine.oSec01Field15,
        "oSec01Field16": machine.oSec01Field16,
        "oSec01Field17": machine.oSec01Field17,
        "oSec01Field18": machine.oSec01Field18,
        "oSec01Field19": machine.oSec01Field19,
        "oSec01Field20": machine.oSec01Field20,
        
        "oSec02Field01": machine.oSec02Field01,
        "oSec02Field02": machine.oSec02Field02,
        "oSec02Field03": machine.oSec02Field03,
        "oSec02Field04": machine.oSec02Field04,
        "oSec02Field05": machine.oSec02Field05,
        "oSec02Field06": machine.oSec02Field06,
        "oSec02Field07": machine.oSec02Field07,
        "oSec02Field08": machine.oSec02Field08,
        "oSec02Field09": machine.oSec02Field09,
        "oSec02Field10": machine.oSec02Field10,        
        "oSec02Field11": machine.oSec02Field11,
        "oSec02Field12": machine.oSec02Field12,
        "oSec02Field13": machine.oSec02Field13,
        "oSec02Field14": machine.oSec02Field14,
        "oSec02Field15": machine.oSec02Field15,
        "oSec02Field16": machine.oSec02Field16,
        "oSec02Field17": machine.oSec02Field17,
        "oSec02Field18": machine.oSec02Field18,
        "oSec02Field19": machine.oSec02Field19,
        "oSec02Field20": machine.oSec02Field20,
        
        "oSec03Field01": machine.oSec03Field01,
        "oSec03Field02": machine.oSec03Field02,
        "oSec03Field03": machine.oSec03Field03,
        "oSec03Field04": machine.oSec03Field04,
        "oSec03Field05": machine.oSec03Field05,
        "oSec03Field06": machine.oSec03Field06,
        "oSec03Field07": machine.oSec03Field07,
        "oSec03Field08": machine.oSec03Field08,
        "oSec03Field09": machine.oSec03Field09,
        "oSec03Field10": machine.oSec03Field10,        
        "oSec03Field11": machine.oSec03Field11,
        "oSec03Field12": machine.oSec03Field12,
        "oSec03Field13": machine.oSec03Field13,
        "oSec03Field14": machine.oSec03Field14,
        "oSec03Field15": machine.oSec03Field15,
        "oSec03Field16": machine.oSec03Field16,
        "oSec03Field17": machine.oSec03Field17,
        "oSec03Field18": machine.oSec03Field18,
        "oSec03Field19": machine.oSec03Field19,
        "oSec03Field20": machine.oSec03Field20,
        
        "oSec04Field01": machine.oSec04Field01,
        "oSec04Field02": machine.oSec04Field02,
        "oSec04Field03": machine.oSec04Field03,
        "oSec04Field04": machine.oSec04Field04,
        "oSec04Field05": machine.oSec04Field05,
        "oSec04Field06": machine.oSec04Field06,
        "oSec04Field07": machine.oSec04Field07,
        "oSec04Field08": machine.oSec04Field08,
        "oSec04Field09": machine.oSec04Field09,
        "oSec04Field10": machine.oSec04Field10,        
        "oSec04Field11": machine.oSec04Field11,
        "oSec04Field12": machine.oSec04Field12,
        "oSec04Field13": machine.oSec04Field13,
        "oSec04Field14": machine.oSec04Field14,
        "oSec04Field15": machine.oSec04Field15,
        "oSec04Field16": machine.oSec04Field16,
        "oSec04Field17": machine.oSec04Field17,
        "oSec04Field18": machine.oSec04Field18,
        "oSec04Field19": machine.oSec04Field19,
        "oSec04Field20": machine.oSec04Field20,
        
        "oSec05Field01": machine.oSec05Field01,
        "oSec05Field02": machine.oSec05Field02,
        "oSec05Field03": machine.oSec05Field03,
        "oSec05Field04": machine.oSec05Field04,
        "oSec05Field05": machine.oSec05Field05,
        "oSec05Field06": machine.oSec05Field06,
        "oSec05Field07": machine.oSec05Field07,
        "oSec05Field08": machine.oSec05Field08,
        "oSec05Field09": machine.oSec05Field09,
        "oSec05Field10": machine.oSec05Field10,        
        "oSec05Field11": machine.oSec05Field11,
        "oSec05Field12": machine.oSec05Field12,
        "oSec05Field13": machine.oSec05Field13,
        "oSec05Field14": machine.oSec05Field14,
        "oSec05Field15": machine.oSec05Field15,
        "oSec05Field16": machine.oSec05Field16,
        "oSec05Field17": machine.oSec05Field17,
        "oSec05Field18": machine.oSec05Field18,
        "oSec05Field19": machine.oSec05Field19,
        "oSec05Field20": machine.oSec05Field20,
        
        "oSec06Field01": machine.oSec06Field01,
        "oSec06Field02": machine.oSec06Field02,
        "oSec06Field03": machine.oSec06Field03,
        "oSec06Field04": machine.oSec06Field04,
        "oSec06Field05": machine.oSec06Field05,
        "oSec06Field06": machine.oSec06Field06,
        "oSec06Field07": machine.oSec06Field07,
        "oSec06Field08": machine.oSec06Field08,
        "oSec06Field09": machine.oSec06Field09,
        "oSec06Field10": machine.oSec06Field10,        
        "oSec06Field11": machine.oSec06Field11,
        "oSec06Field12": machine.oSec06Field12,
        "oSec06Field13": machine.oSec06Field13,
        "oSec06Field14": machine.oSec06Field14,
        "oSec06Field15": machine.oSec06Field15,
        "oSec06Field16": machine.oSec06Field16,
        "oSec06Field17": machine.oSec06Field17,
        "oSec06Field18": machine.oSec06Field18,
        "oSec06Field19": machine.oSec06Field19,
        "oSec06Field20": machine.oSec06Field20,
        
        "oSec07Field01": machine.oSec07Field01,
        "oSec07Field02": machine.oSec07Field02,
        "oSec07Field03": machine.oSec07Field03,
        "oSec07Field04": machine.oSec07Field04,
        "oSec07Field05": machine.oSec07Field05,
        "oSec07Field06": machine.oSec07Field06,
        "oSec07Field07": machine.oSec07Field07,
        "oSec07Field08": machine.oSec07Field08,
        "oSec07Field09": machine.oSec07Field09,
        "oSec07Field10": machine.oSec07Field10,        
        "oSec07Field11": machine.oSec07Field11,
        "oSec07Field12": machine.oSec07Field12,
        "oSec07Field13": machine.oSec07Field13,
        "oSec07Field14": machine.oSec07Field14,
        "oSec07Field15": machine.oSec07Field15,
        "oSec07Field16": machine.oSec07Field16,
        "oSec07Field17": machine.oSec07Field17,
        "oSec07Field18": machine.oSec07Field18,
        "oSec07Field19": machine.oSec07Field19,
        "oSec07Field20": machine.oSec07Field20,
        
        "oSec08Field01": machine.oSec08Field01,
        "oSec08Field02": machine.oSec08Field02,
        "oSec08Field03": machine.oSec08Field03,
        "oSec08Field04": machine.oSec08Field04,
        "oSec08Field05": machine.oSec08Field05,
        "oSec08Field06": machine.oSec08Field06,
        "oSec08Field07": machine.oSec08Field07,
        "oSec08Field08": machine.oSec08Field08,
        "oSec08Field09": machine.oSec08Field09,
        "oSec08Field10": machine.oSec08Field10,        
        "oSec08Field11": machine.oSec08Field11,
        "oSec08Field12": machine.oSec08Field12,
        "oSec08Field13": machine.oSec08Field13,
        "oSec08Field14": machine.oSec08Field14,
        "oSec08Field15": machine.oSec08Field15,
        "oSec08Field16": machine.oSec08Field16,
        "oSec08Field17": machine.oSec08Field17,
        "oSec08Field18": machine.oSec08Field18,
        "oSec08Field19": machine.oSec08Field19,
        "oSec08Field20": machine.oSec08Field20,
        
        "oSec09Field01": machine.oSec09Field01,
        "oSec09Field02": machine.oSec09Field02,
        "oSec09Field03": machine.oSec09Field03,
        "oSec09Field04": machine.oSec09Field04,
        "oSec09Field05": machine.oSec09Field05,
        "oSec09Field06": machine.oSec09Field06,
        "oSec09Field07": machine.oSec09Field07,
        "oSec09Field08": machine.oSec09Field08,
        "oSec09Field09": machine.oSec09Field09,
        "oSec09Field10": machine.oSec09Field10,        
        "oSec09Field11": machine.oSec09Field11,
        "oSec09Field12": machine.oSec09Field12,
        "oSec09Field13": machine.oSec09Field13,
        "oSec09Field14": machine.oSec09Field14,
        "oSec09Field15": machine.oSec09Field15,
        "oSec09Field16": machine.oSec09Field16,
        "oSec09Field17": machine.oSec09Field17,
        "oSec09Field18": machine.oSec09Field18,
        "oSec09Field19": machine.oSec09Field19,
        "oSec09Field20": machine.oSec09Field20,
        
        "oSec10Field01": machine.oSec10Field01,
        "oSec10Field02": machine.oSec10Field02,
        "oSec10Field03": machine.oSec10Field03,
        "oSec10Field04": machine.oSec10Field04,
        "oSec10Field05": machine.oSec10Field05,
        "oSec10Field06": machine.oSec10Field06,
        "oSec10Field07": machine.oSec10Field07,
        "oSec10Field08": machine.oSec10Field08,
        "oSec10Field09": machine.oSec10Field09,
        "oSec10Field10": machine.oSec10Field10,        
        "oSec10Field11": machine.oSec10Field11,
        "oSec10Field12": machine.oSec10Field12,
        "oSec10Field13": machine.oSec10Field13,
        "oSec10Field14": machine.oSec10Field14,
        "oSec10Field15": machine.oSec10Field15,
        "oSec10Field16": machine.oSec10Field16,
        "oSec10Field17": machine.oSec10Field17,
        "oSec10Field18": machine.oSec10Field18,
        "oSec10Field19": machine.oSec10Field19,
        "oSec10Field20": machine.oSec10Field20,
        
        
        
        # Add other fields if necessary
    }

    return JsonResponse(data)


