import pdb

from . import forms


from datetime import datetime
from docx import Document
import requests

from django.http import HttpResponse
from django.shortcuts import render, redirect

def interact_with_api(api_url, req_type, input_data):
    """
    Interact with the specified API by sending a POST request.

    Parameters:
        api_url (str): The API endpoint URL.
        req_type (str): The request type (e.g., 'MS').
        input_data (dict): A dictionary of input parameters.

    Returns:
        dict: The API response parsed into a Python dictionary.
    """
    # Prepare the payload
    payload = {
        "reqType": req_type,
        **input_data  # Merge the input data into the payload
    }

    try:
        # Send POST request
        response = requests.post(api_url, json=payload)

        # Raise an error for bad responses
        response.raise_for_status()

        # Parse and return the JSON response
        return response.json()

    except requests.exceptions.RequestException as e:
        print(f"Error while interacting with API: {e}")
        return None


SHEET_CONFIG = {
    'MS': {
        'form_class': forms.formCalcMS,
        'aMachineName': 'Mechanical Screen',
        'template': 'PageCalculationsSheet.html',
        'api_type': 'MS',
        'input_fields': [
            'oSec01Field01', 'oSec01Field02', 'oSec01Field03',
            'oSec01Field04', 'oSec01Field05', 'oSec01Field06', 'oSec01Field07',
            'oSec01Field08', 'oSec01Field09', 'oSec01Field10', 'oSec01Field11',
        ],
        'output_fields': ['oSec02Field01', 'oSec02Field02', 'oSec02Field03'],
        'api_fields': {
            "MS_ChannelHeight":     'oSec01Field01',
                "MS_ScreenWidth":       'oSec01Field02',
                "MS_BeltHeight":        'oSec01Field03',
                "MS_WaterLevel":        'oSec01Field04',
                "MS_BarSpacing":        'oSec01Field05',
                "MS_BarThickness":      'oSec01Field06',
                "MS_BarWidth":          'oSec01Field07',
                "MS_InclinationDegree": 'oSec01Field08',
                "MS_SprocketDiameter":  'oSec01Field09',
                "MS_Velocity":          'oSec01Field10',
                "MS_FOS":               'oSec01Field11',
        },
    },
    'BC': {
        'form_class': forms.formCalcBC,
        'aMachineName': 'Belt Conveyor',
        'template': 'PageCalculationsSheet.html',
        'api_type': 'BC',
        'input_fields': [
            'oSec01Field01', 'oSec01Field02', 'oSec01Field03',
            'oSec01Field04', 'oSec01Field05', 'oSec01Field06', 'oSec01Field07',
        ],
        'output_fields': ['oSec02Field01', 'oSec02Field02', 'oSec02Field03'],
        'api_fields': {
            'BC_Length': 'oSec01Field01',
            'BC_Width': 'oSec01Field02',
            'BC_DrumDia': 'oSec01Field03',
            'BC_Friction': 'oSec01Field04',
            'BC_Velocity': 'oSec01Field05',
            'BC_FOS': 'oSec01Field06',
            'BC_Belt_weight_per_meter': 'oSec01Field07',
        },
    },
    'GR': {
        'form_class': forms.formCalcGR,
        'aMachineName': 'Grit Removal',
        'template': 'PageCalculationsSheet.html',
        'api_type': 'GR',
        'input_fields': [
            'oSec01Field01', 'oSec01Field02', 'oSec01Field03', 'oSec01Field04',
            'oSec01Field05', 'oSec01Field06', 'oSec01Field07', 'oSec01Field08',
        ],
        'output_fields': [
            'oSec02Field01', 'oSec02Field02', 'oSec02Field03',
            'oSec02Field04', 'oSec02Field05', 'oSec02Field06',
        ],
        'api_fields': {
            'GR_n_channel': 'oSec01Field01',
            'GR_channel_width': 'oSec01Field02',
            'GR_civil_width': 'oSec01Field03',
            'GR_bridge_length': 'oSec01Field04',
            'GR_wheel_diameter': 'oSec01Field05',
            'GR_Friction': 'oSec01Field06',
            'GR_Velocity': 'oSec01Field07',
            'GR_FOS': 'oSec01Field08',
        },
    },
    'PS': {
        'form_class': forms.formCalcPS,
        'aMachineName': 'PST',
        'template': 'PageCalculationsSheet.html',
        'api_type': 'PS',
        'input_fields': [
            'oSec01Field01', 'oSec01Field02', 'oSec01Field03', 'oSec01Field04',
        ],
        'output_fields': [
            'oSec02Field01', 'oSec02Field02', 'oSec02Field03',
            'oSec02Field04', 'oSec02Field05'
        ],
        'api_fields': {
                "PS_walkway_length":     'oSec01Field01',
                "PS_Friction":       'oSec01Field02',
                "PS_Velocity":        'oSec01Field03',
                "PS_FOS":        'oSec01Field04',
        },
    },
    'TH': {
        'form_class': forms.formCalcTH,
        'aMachineName': 'Thickener',
        'template': 'PageCalculationsSheet.html',
        'api_type': 'TH',
        'input_fields': [
            'oSec01Field01', 'oSec01Field02', 'oSec01Field03', 'oSec01Field04',
        ],
        'output_fields': [
            'oSec02Field01', 'oSec02Field02', 'oSec02Field03',
            'oSec02Field04', 'oSec02Field05',
        ],
        'api_fields': {
            "TH_diameter":     'oSec01Field01',
                "TH_n_arm":       'oSec01Field02',
                "TH_Velocity":        'oSec01Field03',
                "TH_FOS":        'oSec01Field04',
        },
    },
    'MX': {
        'form_class': forms.formCalcMX,
        'aMachineName': 'Rectangular Mixers',
        'template': 'PageCalculationsSheet.html',
        'api_type': 'MX',
        'input_fields': [
            'oSec01Field01', 'oSec01Field02', 'oSec01Field03', 'oSec01Field04',
            'oSec01Field05', 'oSec01Field06', 'oSec01Field07', 'oSec01Field08',
        ],
        'output_fields': [
            'oSec02Field01', 'oSec02Field02', 'oSec02Field03', 'oSec02Field04', 
            'oSec02Field05', 'oSec02Field06', 'oSec02Field07',
        ],
        'api_fields': {
            "MX_length":        'oSec01Field01',
             "MX_width":         'oSec01Field02',
             "MX_water_depth":           'oSec01Field03',
             "MX_tank_depth":            'oSec01Field04',
             "MX_impeller_coefficient":  'oSec01Field05',
             "MX_velocity_gradient":     'oSec01Field06',
             "MX_impeller_diameter_factor":  'oSec01Field07',
             "MX_safety_factor":             'oSec01Field08',
        },
    },
    'RT': {
        'form_class': forms.formCalcRT,
        'aMachineName': 'Rectangular Tanks',
        'template': 'PageCalculationsSheet.html',
        'api_type': 'RT',
        'input_fields': [
            'oSec01Field01', 'oSec01Field02', 'oSec01Field03', 'oSec01Field04',
            'oSec01Field05', 'oSec01Field06',
        ],
        'output_fields': [
            'oSec02Field01', 
        ],
        'api_fields': {
            "RT_Length":        'oSec01Field01',
            "RT_Width":         'oSec01Field02',
            "RT_Hight":           'oSec01Field03',
            "RT_ShellTH":            'oSec01Field04',
            "RT_BaseTH":    'oSec01Field05',
            "RT_N_Spliter":     'oSec01Field06',
        },
    },
    'CT': {
        'form_class': forms.formCalcCT,
        'aMachineName': 'Circular Tanks',
        'template': 'PageCalculationsSheet.html',
        'api_type': 'GR',
        'input_fields': [
            'oSec01Field01', 'oSec01Field02',
        ],
        'output_fields': [
            'oSec02Field01', 'oSec02Field02', 'oSec02Field03', 'oSec02Field04', 
            'oSec02Field05', 'oSec02Field06', 'oSec02Field07', 'oSec02Field08',
        ],
        'api_fields': {
            "CT_Diameter":        'oSec01Field01',
            "CT_Height":         'oSec01Field02',
        },
    },
    'SC': {
        'form_class': forms.formCalcSC,
        'aMachineName': 'Screw Conveyor',
        'template': 'PageCalculationsSheet.html',
        'api_type': 'SC',
        'input_fields': [
            'oSec01Field01', 'oSec01Field02', 'oSec01Field03', 'oSec01Field04',
            'oSec01Field05', 'oSec01Field06', 'oSec01Field07', 'oSec01Field08',
        ],
        'output_fields': [
            'oSec02Field01', 'oSec02Field02', 'oSec02Field03',
            'oSec02Field04', 'oSec02Field05', 'oSec02Field06',
        ],
        'api_fields': {
            'aInput01': 'oSec01Field01',
            'aInput02': 'oSec01Field02',
            'aInput03': 'oSec01Field03',
            'aInput04': 'oSec01Field04',
            'aInput05': 'oSec01Field05',
            'aInput06': 'oSec01Field06',
            'aInput07': 'oSec01Field07',
            'aInput08': 'oSec01Field08',
        },
    },
    'BS': {
        'form_class': forms.formCalcBS,
        'aMachineName': 'Basket Screen',
        'template': 'PageCalculationsSheet.html',
        'api_type': 'BS',
        'input_fields': [
            'oSec01Field01', 'oSec01Field02', 'oSec01Field03', 'oSec01Field04',
            'oSec01Field05', 'oSec01Field06',
        ],
        'output_fields': [
            'oSec02Field01', 'oSec02Field02', 'oSec02Field03',
        ],
        'api_fields': {
            'BS_Bar_Dia': 'oSec01Field01',
            'BS_Bar_Space': 'oSec01Field02',
            'BS_Screen_Height': 'oSec01Field03',
            'BS_Screen_Width': 'oSec01Field04',
            'BS_Screen_Depth': 'oSec01Field05',
            'BS_Plate_Th': 'oSec01Field06',
        },
    },
    'NS': {
        'form_class': forms.formCalcNS,
        'aMachineName': 'Manual Screen',
        'template': 'PageCalculationsSheet.html',
        'api_type': 'NS',
        'input_fields': [
            'oSec01Field01', 'oSec01Field02', 'oSec01Field03', 'oSec01Field04',
            'oSec01Field05', 'oSec01Field06', 'oSec01Field07', 'oSec01Field08',
        ],
        'output_fields': [
            'oSec02Field01',
        ],
        'api_fields': {
            'NS_Ch_Height': 'oSec01Field01',
            'NS_Ch_Width': 'oSec01Field02',
            'NS_WaterLv': 'oSec01Field03',
            'NS_WaterLv_Margin': 'oSec01Field04',
            'NS_Bar_Spacing': 'oSec01Field05',
            'NS_Bar_Th': 'oSec01Field06',
            'NS_Bar_Width': 'oSec01Field07',
            'NS_Angle': 'oSec01Field08',
        },
    },
    'PNch': {
        'form_class': forms.formCalcPNch,
        'aMachineName': 'Channel Penstock',
        'template': 'PageCalculationsSheet.html',
        'api_type': 'PNch',
        'input_fields': [
            'oSec01Field01', 'oSec01Field02', 'oSec01Field03', 'oSec01Field04',
            'oSec01Field05', 'oSec01Field06', 'oSec01Field07', 'oSec01Field08',
            'oSec01Field09', 'oSec01Field10',
        ],
        'output_fields': [
            'oSec02Field01', 'oSec02Field02', 'oSec02Field03', 'oSec02Field04', 
            'oSec02Field05', 'oSec02Field06', 'oSec02Field07', 
        ],
        'api_fields': {
            'PNch_Channel_Height': 'oSec01Field01',
            'PNch_Frame_Height_Over_Channel': 'oSec01Field02',
            'PNch_Channel_Width': 'oSec01Field03',
            'PNch_Gate_Margin_Width': 'oSec01Field04',
            'PNch_Water_Lv': 'oSec01Field05',
            'PNch_Gate_Margin_Over_Water_Lv': 'oSec01Field06',
            'PNch_Gate_Th': 'oSec01Field07',
            'PNch_Gate_Other_PLs': 'oSec01Field08',
            'PNch_HeadStock': 'oSec01Field09',
            'PNch_Frame_Weight_Per_M': 'oSec01Field10',
        },
    },
    'PNwa': {
        'form_class': forms.formCalcPNwa,
        'aMachineName': 'Wall Penstock',
        'template': 'PageCalculationsSheet.html',
        'api_type': 'PNwa',
        'input_fields': [
            'oSec01Field01', 'oSec01Field02', 'oSec01Field03', 'oSec01Field04',
            'oSec01Field05', 'oSec01Field06', 'oSec01Field07', 'oSec01Field08',
            'oSec01Field09', 'oSec01Field10',
        ],
        'output_fields': [
            'oSec02Field01', 'oSec02Field02', 'oSec02Field03', 'oSec02Field04', 
            'oSec02Field05', 'oSec02Field06', 'oSec02Field07', 
        ],
        'api_fields': {
            'aInput01': 'oSec01Field01',
            'aInput02': 'oSec01Field02',
            'aInput03': 'oSec01Field03',
            'aInput04': 'oSec01Field04',
            'aInput05': 'oSec01Field05',
            'aInput06': 'oSec01Field06',
            'aInput07': 'oSec01Field07',
            'aInput08': 'oSec01Field08',
            'aInput09': 'oSec01Field09',
            'aInput010': 'oSec01Field10',
        },
    },
    
}

def LoadPageCalculationsSheet(request, sheet_key):
    if not request.user.is_authenticated:
        return redirect('login')

    config = SHEET_CONFIG.get(sheet_key)
    if not config:
        return HttpResponse("Invalid sheet key", status=404)

    aMachineName = config['aMachineName']
    form_class = config['form_class']
    template = config['template']
    api_type = config['api_type']
    api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"

    if request.method == "GET":
        form = form_class()
        # Initialize all section variables
        aSection01Show = "Yes"
        aSection02Show = "Yes"

        print(form.fields['oSec01Field01'].initial)
        print(form.fields['oSec02Field01'].initial)

        # Apply conditions to modify the values
        if form.fields['oSec01Field01'].initial in ["oooo", None]:
            aSection01Show = "Hide"

        if form.fields['oSec02Field01'].initial in ["oooo", None]:
            aSection02Show = "Hide"

    

        print(aSection01Show)
        print(aSection02Show)
        return render(request, template, {
            'form1': form,
            'aMachineName':aMachineName,
            "aSection01Show": aSection01Show,
            "aSection02Show": aSection02Show,
            })

    if request.method == "POST":
        form = form_class(request.POST)
        if 'generate_report' in request.POST:
            doc = Document()
            doc.add_heading(f'{sheet_key} Report', level=1)

            sections = {
                "Input": config['input_fields'],
                "Output": config['output_fields'],
            }

            for section, fields in sections.items():
                doc.add_heading(section, level=2)
                for field in fields:
                    label = form.fields[field].label
                    value = request.POST.get(field, "N/A")
                    doc.add_paragraph(f"{label}: {value}")

            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{sheet_key}_Report.docx"'
            doc.save(response)
            return response
        
        if 'form1_submit' in request.POST:
            form = form_class(request.POST)
            if form.is_valid():
                cleaned = form.cleaned_data
                input_data = {
                    api_key: cleaned.get(form_field)
                    for api_key, form_field in config['api_fields'].items()
                }

                # Call external API
                response = interact_with_api(api_url, api_type, input_data)
                

                # Update output fields
                instance = form.save(commit=False)
                for field in config['output_fields']:
                    if field in response:
                        setattr(instance, field, response[field])

                instance.oSec00Field01 = request.user.username
                instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                instance.oSec00Field03 = sheet_key
                instance.save()

                # Re-initialize with new values for display
                initial_data = {field: cleaned.get(field) for field in config['input_fields']}
                initial_data.update({field: response.get(field) for field in config['output_fields']})
                form = form_class(initial=initial_data)
                # Initialize all section variables
                aSection01Show = "Yes"
                aSection02Show = "Yes"

                print(form.fields['oSec01Field01'].initial)
                print(form.fields['oSec02Field01'].initial)

                # Apply conditions to modify the values
                if form.fields['oSec01Field01'].initial in ["oooo", None]:
                    aSection01Show = "Hide"

                if form.fields['oSec02Field01'].initial in ["oooo", None]:
                    aSection02Show = "Hide"



                print(aSection01Show)
                print(aSection02Show)

            return render(request, template, {
                'form1': form,
                'aMachineName':aMachineName,
                "aSection01Show": aSection01Show,
                "aSection02Show": aSection02Show,
                })

    return HttpResponse("Invalid request method", status=405)


""" def handle_form(request , sheet_key):
    if not request.user.is_authenticated:
        return redirect('login')  # Redirect to the login page if the user is not authenticated
    
    config = SHEET_CONFIG.get(sheet_key)
    if not config:
        return HttpResponse("Invalid sheet key", status=404)

    
    form_class = config['form_class']
    template = config['template']
    api_type = config['api_type']


    if request.method == 'POST' and 'form1_submit' in request.POST:
        form1 = form_class(request.POST)
        if form1.is_valid():
            # Access the cleaned_data dictionary to get individual field values
            cleaned = form1.cleaned_data
            input_data = {
                api_key: cleaned.get(form_field)
                for api_key, form_field in config['api_fields'].items()
            }

            # Calculate new values for fields            
            api_url = "https://us-central1-h1000project1.cloudfunctions.net/f01"
            # Call external API
            response = interact_with_api(api_url, api_type, input_data)
            # Update output fields
            instance = form1.save(commit=False)
            for field in config['output_fields']:
                if field in response:
                    setattr(instance, field, response[field])
            instance.oSec00Field01 = request.user.username
            instance.oSec00Field02 = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            instance.oSec00Field03 = sheet_key
            instance.save()
            # Re-initialize with new values for display
            initial_data = {field: cleaned.get(field) for field in config['input_fields']}
            initial_data.update({field: response.get(field) for field in config['output_fields']})
            form = form_class(initial=initial_data)

        return render(request, template, {'form1': form1})

    return redirect('ms_load')  # Redirect to the page if the request is invalid """