import pdb

from Apps.aAppMechanical.models import Project
from Apps.aAppMechanical.models import Machine
from Apps.aAppMechanical.models import UserCompany
from Apps.aAppMechanical.models import aLogEntry


from .forms import ProjectForm


import requests

from django.http import HttpResponse
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.shortcuts import render, redirect
from django.urls import reverse
from django.utils.timezone import now 
from django.contrib.auth.models import User
from django.conf import settings


from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Inches
from docx.shared import Pt
# Create your views here.


def project_list(request):
    if request.method == "POST":
        form = ProjectForm(request.POST)
        if form.is_valid():
            
            instance = form.save(commit=False)  # Don't save to DB yet

            
            # Get the company associated with the user
            try:
                user_company = UserCompany.objects.get(user=request.user).company
                instance.company = user_company  # Assign company to the instance
            except UserCompany.DoesNotExist:
                return render(request, 'project_list', {"form": form, "error": "User is not associated with a company"})


            form.save()
            return redirect('project_list')
    else:
        form = ProjectForm()
        
        
    # Get the company of the logged-in user
    user_company = None
    if request.user.is_authenticated:
        try:
            user_company = UserCompany.objects.get(user=request.user).company
        except UserCompany.DoesNotExist:
            user_company = None

    # Assign company filter only if the user has a company
    if user_company:
        #machines = Machine.objects.filter(oSec00Field03=DB_Name, company=user_company)
        projects = Project.objects.filter(company=user_company)
    else:
        projects = Project.objects.none()  # Return an empty queryset if no company
    
    #projects = Project.objects.all()
    return render(request, 'project_list.html', {'form': form, 'projects': projects})
    
def edit_project(request, project_id):
    project = get_object_or_404(Project, id=project_id)
    
    if request.method == "POST":
        form = ProjectForm(request.POST, instance=project)
        if form.is_valid():
            form.save()
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({'success': True})
            return redirect('project_list')
        else:
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({'success': False, 'form': form.as_p()})
    else:
        form = ProjectForm(instance=project)
        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            return JsonResponse({'success': True, 'form': form.as_p()})

    return render(request, 'edit_project.html', {'form': form})



def delete_project(request, project_id):
    if request.method == "POST":
        project = get_object_or_404(Project, id=project_id)
        project.delete()

        # Otherwise, redirect to the project list page
        return redirect('project_list')

    return JsonResponse({'success': False, 'error': 'Invalid request'})



def get_machines(request, project_id):
    try:
        project = Project.objects.get(id=project_id)
        machines = Machine.objects.filter(project=project)
        
        data = {
            "project_name": project.name,
            "machines": list(machines.values("oSec00Field01", "oSec00Field02", "oSec00Field03",
                                             "oSec01Field01", "oSec01Field02", "oSec01Field03",
                                             "oSec01Field04", "oSec01Field05", "oSec01Field06",
                                             "oSec01Field07", "oSec01Field08", "oSec01Field09",
                                             "oSec01Field10", "oSec01Field11", "oSec01Field12",
                                             "oSec01Field13", "oSec01Field14", "oSec01Field15",
                                             "oSec01Field16", "oSec01Field17", "oSec01Field18",
                                             "oSec01Field19", "oSec01Field20",
                                             "oSec02Field01", "oSec02Field02", "oSec02Field03",
                                             "oSec02Field04", "oSec02Field05", "oSec02Field06",
                                             "oSec02Field07", "oSec02Field08", "oSec02Field09",
                                             "oSec02Field10", "oSec02Field11", "oSec02Field12",
                                             "oSec02Field13", "oSec02Field14", "oSec02Field15",
                                             "oSec02Field16", "oSec02Field17", "oSec02Field18",
                                             "oSec02Field19", "oSec02Field20",
                                             "oSec03Field01", "oSec03Field02", "oSec03Field03",
                                             "oSec03Field04", "oSec03Field05", "oSec03Field06",
                                             "oSec03Field07", "oSec03Field08", "oSec03Field09",
                                             "oSec03Field10", "oSec03Field11", "oSec03Field12",
                                             "oSec03Field13", "oSec03Field14", "oSec03Field15",
                                             "oSec03Field16", "oSec03Field17", "oSec03Field18",
                                             "oSec03Field19", "oSec03Field20",
                                             "oSec04Field01", "oSec04Field02", "oSec04Field03",
                                             "oSec04Field04", "oSec04Field05", "oSec04Field06",
                                             "oSec04Field07", "oSec04Field08", "oSec04Field09",
                                             "oSec04Field10", "oSec04Field11", "oSec04Field12",
                                             "oSec04Field13", "oSec04Field14", "oSec04Field15",
                                             "oSec04Field16", "oSec04Field17", "oSec04Field18",
                                             "oSec04Field19", "oSec04Field20",
                                             "oSec05Field01", "oSec05Field02", "oSec05Field03",
                                             "oSec05Field04", "oSec05Field05", "oSec05Field06",
                                             "oSec05Field07", "oSec05Field08", "oSec05Field09",
                                             "oSec05Field10", "oSec05Field11", "oSec05Field12",
                                             "oSec05Field13", "oSec05Field14", "oSec05Field15",
                                             "oSec05Field16", "oSec05Field17", "oSec05Field18",
                                             "oSec05Field19", "oSec05Field20",
                                             "oSec06Field01", "oSec06Field02", "oSec06Field03",
                                             "oSec06Field04", "oSec06Field05", "oSec06Field06",
                                             "oSec06Field07", "oSec06Field08", "oSec06Field09",
                                             "oSec06Field10", "oSec06Field11", "oSec06Field12",
                                             "oSec06Field13", "oSec06Field14", "oSec06Field15",
                                             "oSec06Field16", "oSec06Field17", "oSec06Field18",
                                             "oSec06Field19", "oSec06Field20",
                                             "oSec07Field01", "oSec07Field02", "oSec07Field03",
                                             "oSec07Field04", "oSec07Field05", "oSec07Field06",
                                             "oSec07Field07", "oSec07Field08", "oSec07Field09",
                                             "oSec07Field10", "oSec07Field11", "oSec07Field12",
                                             "oSec07Field13", "oSec07Field14", "oSec07Field15",
                                             "oSec07Field16", "oSec07Field17", "oSec07Field18",
                                             "oSec07Field19", "oSec07Field20",
                                             "oSec08Field01", "oSec08Field02", "oSec08Field03",
                                             "oSec08Field04", "oSec08Field05", "oSec08Field06",
                                             "oSec08Field07", "oSec08Field08", "oSec08Field09",
                                             "oSec08Field10", "oSec08Field11", "oSec08Field12",
                                             "oSec08Field13", "oSec08Field14", "oSec08Field15",
                                             "oSec08Field16", "oSec08Field17", "oSec08Field18",
                                             "oSec08Field19", "oSec08Field20",
                                             "oSec09Field01", "oSec09Field02", "oSec09Field03",
                                             "oSec09Field04", "oSec09Field05", "oSec09Field06",
                                             "oSec09Field07", "oSec09Field08", "oSec09Field09",
                                             "oSec09Field10", "oSec09Field11", "oSec09Field12",
                                             "oSec09Field13", "oSec09Field14", "oSec09Field15",
                                             "oSec09Field16", "oSec09Field17", "oSec09Field18",
                                             "oSec09Field19", "oSec09Field20",
                                             "oSec10Field01", "oSec10Field02", "oSec10Field03",
                                             "oSec10Field04", "oSec10Field05", "oSec10Field06",
                                             "oSec10Field07", "oSec10Field08", "oSec10Field09",
                                             "oSec10Field10", "oSec10Field11", "oSec10Field12",
                                             "oSec10Field13", "oSec10Field14", "oSec10Field15",
                                             "oSec10Field16", "oSec10Field17", "oSec10Field18",
                                             "oSec10Field19", "oSec10Field20"))
        }
        return JsonResponse(data)
    except Project.DoesNotExist:
        return JsonResponse({"error": "Project not found"}, status=404)












def generate_report(request, project_id):
    try:
        #pdb.set_trace()
        # Log the action
        aLogEntry.objects.create(user=request.user, message=f"at {now()} {request.user} accessed Word Report")
        
        #pdb.set_trace()
        # Get the user’s company and project
        aCompany = UserCompany.objects.get(user=request.user)

        #pdb.set_trace()
        # Determine the company and generate the corresponding report
        if aCompany.id == 1:
            print("Company 1")
            return generate_report_AAA(request, project_id)

        elif aCompany.id == 2:
            print("Company 2")
            return generate_report_BBB(request, project_id)

        else:
            return HttpResponse("Invalid company ID", status=400)

    except UserCompany.DoesNotExist:
        return HttpResponse("User does not belong to a company", status=403)

    except Project.DoesNotExist:
        return HttpResponse("Project not found", status=404)




def generate_report_AAA(request, project_id):
    
    def add_table(doc, data, title=None):
        """Creates a table and applies a background color to the header."""
        if title:
            doc.add_heading(title, level=3)

        table = doc.add_table(rows=len(data), cols=2)
        table.style = "Table Grid"

        for i, row in enumerate(data):
            for j, text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = text

                # Apply background color only to the header row (first row)
                if i == 0:
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(ns.qn("w:fill"), "FFA500")  # Orange color                    
                    #shading_elm.set(ns.qn("w:fill"), "ADD8E6")  # Blue color
                        
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    
            """Generates a Word report for a given project."""
    
    
    
    def add_header_footer(doc):
        """Adds header and footer with page numbers in the format 'Page X of Y'."""
        section = doc.sections[0]
    
        # Header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        # header_para.add_run("Company Name\n")
        # header_para.add_run("Project Name\n")
        # header_para.add_run("Date: " + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "\n")
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Adding logo
        run_logo = header_para.add_run()  # Corrected reference to header paragraph
        
        try:
            run_logo.add_picture("LogoAAA.PNG", width=Inches(7.0))  # Adjust width as needed
        except Exception as e:
            print(f"Error adding logo: {e}")

        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Add "Page X of Y" format
        run = footer_para.add_run("Page ")
    
        # PAGE field (Current Page Number)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(ns.qn("w:fldCharType"), "begin")
    
        instrText1 = OxmlElement("w:instrText")
        instrText1.set(ns.qn("xml:space"), "preserve")
        instrText1.text = "PAGE"
    
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar1)
        run._r.append(instrText1)
        run._r.append(fldChar2)
    
        run.add_text(" of ")
    
        # NUMPAGES field (Total Number of Pages)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(ns.qn("w:fldCharType"), "begin")
    
        instrText2 = OxmlElement("w:instrText")
        instrText2.set(ns.qn("xml:space"), "preserve")
        instrText2.text = "NUMPAGES"
    
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)
    
    
    def add_colored_heading(doc, text, level, color):
        """Adds a heading with color."""
        heading = doc.add_paragraph()
        run = heading.add_run(text)
        run.bold = True
        run.font.size = Pt(14) if level == 1 else Pt(12)
        run.font.color.rgb = color
        heading.style = f"Heading {level}"
    
    
    
    try:
        
        ###LOG
        aLogEntry.objects.create(
                user=request.user,
                message=f"at {now()} {request.user} accessed Load  "
            )
        print(f"at {now()} {User} accessed Download Report")
        ###LOG

        aCompany = UserCompany.objects.get(user=request.user)
        project = Project.objects.get(id=project_id)
        machines = Machine.objects.filter(project=project)
        
        
        print(aCompany.id)
        print(project.id)
    
        print("Company 1")
    
    
        # Create a Word document
        doc = Document()

        # Add header and footer with page numbers
        add_header_footer(doc)

        # Add project title
        doc.add_heading(f'Project Report: {project.name}', level=1)

        # Add project details
        doc.add_heading("Project Details", level=2)     

        doc.add_paragraph("\n")
        doc.add_paragraph("Name: " + project.name)
        doc.add_paragraph("Client Name: " + project.client_name)
        doc.add_paragraph("Capacity: " + project.capacity)
        doc.add_paragraph("\n")
        
        doc.add_page_break()     
        doc.add_paragraph("\n")

        # Add machine details
        for index, machine in enumerate(machines, start=1):  # Add numbering
            machine_name = machine.oSec00Field03
            section_titles = []

            if machine_name == "DataSheetNS":
                machine_name = "Manual Screen" 
                section_titles = ["General Data", "Design Data", "Material Data", "Channel Data", " ", " ", " ", " ", " ", " "]

            if machine_name == "DataSheetMSc":
                machine_name = "Mechanical Screen" 
                section_titles = ["General Data", "Design Data", "Gearmotor Data", "Control panel Data", "Material Data", "Other Data", " ", " ", " ", " "]

            if machine_name == "DataSheetBC":
                machine_name = "Belt Conveyor"
                section_titles = ["General Data", "Design Data", "Gearbox Data", "Motor Data", "Material Data", " ", " ", " ", " ", " "]

            if machine_name == "DataSheetCO":
                machine_name = "Container"
                section_titles = ["General Data", "Design Data", "Material Data", " ", " ", " ", " ", " ", " ", " "]

            if machine_name == "DataSheetGR":
                machine_name = "Gritremoval"
                section_titles = ["General Data", "Design Data", "Walkway, Handrail, Wheel Data", "Scrapper Data", "Gearmotor Data", "Scrapper Data", "Drive unit", "Control panel Data", "Material Data ", " "]

            if machine_name == "DataSheetSS":
                machine_name = "Sand Silo"

            if machine_name == "DataSheetPS":
                machine_name = "Primary Sedimentation"

            if machine_name == "DataSheetQV":
                machine_name = "Quick Valve"

            if machine_name == "DataSheetTV":
                machine_name = "Telescopic Valve"
                
            if machine_name == "DataSheetTH":
                machine_name = "Sludge Thickener"

            # Add machine title with font size 14 and numbering
            machine_title = doc.add_paragraph(f"{index}. {machine_name}", style="Heading3")
            machine_title.runs[0].font.size = Pt(14)

            for i in range(1, 11):  # Loop from Sec01 to Sec10
                section_name = f"Sec{i:02d}"
                section_data = [("Field", "Value")]

                for j in range(1, 21, 2):  # Step by 2 to avoid duplication
                    key = getattr(machine, f"o{section_name}Field{j:02d}", "").strip()
                    value = getattr(machine, f"o{section_name}Field{j+1:02d}", "").strip()

                    if key and value and key.lower() != "oooo" and value.lower() != "oooo":
                        section_data.append((key, value))

                if len(section_data) > 1:  # If the section has valid data, create a table
                    section_title = section_titles[i-1] if i-1 < len(section_titles) else f"Section {i}"
                    doc.add_paragraph(f"{section_name}: {section_title}", style="Heading3")  # Only one title now

                    add_table(doc, section_data)  # Removed redundant title

            doc.add_page_break() 
        
        # Save the document to a response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={project.name}_report.docx'
        doc.save(response)
        return response

    except Project.DoesNotExist:
        return HttpResponse("Project not found", status=404)





def generate_report_BBB(request, project_id):
    
    def add_table(doc, data, title=None):
        """Creates a borderless table and applies a background color to the header."""
        if title:
            doc.add_heading(title, level=3)

        table = doc.add_table(rows=len(data), cols=2)

        # Remove all table borders manually
        tbl = table._tbl  # Get the table's XML element
        tblPr = tbl.find(qn("w:tblPr"))  # Find existing table properties

        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")  # Create table properties if missing
            tbl.insert(0, tblPr)  # Insert as the first child of <w:tbl>

        tblBorders = OxmlElement("w:tblBorders")  # Create <w:tblBorders>
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")  # Remove the border
            tblBorders.append(border)

        tblPr.append(tblBorders)  # Append border settings to the table properties

        for i, row in enumerate(data):
            for j, text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = text

                # Apply background color only to the header row (first row)
                if i == 0:
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:val"), "clear")  # Set shading value
                    shading_elm.set(qn("w:fill"), "ADD8E6")  # Light blue color
                    cell._tc.get_or_add_tcPr().append(shading_elm)
   
    
    
    def add_header_footer(doc):
        """Adds header and footer with page numbers in the format 'Page X of Y'."""
        section = doc.sections[0]
    
        # Header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        # header_para.add_run("Company Name\n")
        # header_para.add_run("Project Name\n")
        # header_para.add_run("Date: " + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "\n")
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Adding logo
        run_logo = header_para.add_run()  # Corrected reference to header paragraph
        try:
            run_logo.add_picture("LogoBBB.PNG", width=Inches(7.0))  # Adjust width as needed
        except Exception as e:
            print(f"Error adding logo: {e}")

        
    
        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Add "Page X of Y" format
        run = footer_para.add_run("Page ")
    
        # PAGE field (Current Page Number)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(ns.qn("w:fldCharType"), "begin")
    
        instrText1 = OxmlElement("w:instrText")
        instrText1.set(ns.qn("xml:space"), "preserve")
        instrText1.text = "PAGE"
    
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar1)
        run._r.append(instrText1)
        run._r.append(fldChar2)
    
        run.add_text(" of ")
    
        # NUMPAGES field (Total Number of Pages)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(ns.qn("w:fldCharType"), "begin")
    
        instrText2 = OxmlElement("w:instrText")
        instrText2.set(ns.qn("xml:space"), "preserve")
        instrText2.text = "NUMPAGES"
    
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)
    
    
    def add_colored_heading(doc, text, level, color):
        """Adds a heading with color."""
        heading = doc.add_paragraph()
        run = heading.add_run(text)
        run.bold = True
        run.font.size = Pt(14) if level == 1 else Pt(12)
        run.font.color.rgb = color
        heading.style = f"Heading {level}"
    
    
    
    try:
        
        ###LOG
        aLogEntry.objects.create(
                user=request.user,
                message=f"at {now()} {request.user} accessed Load  "
            )
        print(f"at {now()} {User} accessed Download Report")
        ###LOG

        aCompany = UserCompany.objects.get(user=request.user)
        project = Project.objects.get(id=project_id)
        machines = Machine.objects.filter(project=project)
        
        
        print(aCompany.id)
        print(project.id)
    
        print("Company 2")
    
    
        # Create a Word document
        doc = Document()

        # Add header and footer with page numbers
        add_header_footer(doc)

        # Add project title
        doc.add_heading(f'Project Report: {project.name}', level=1)

        # Add project details
        doc.add_heading("Project Details", level=2)     

        doc.add_paragraph("\n")
        doc.add_paragraph("Name: " + project.name)
        doc.add_paragraph("Client Name: " + project.client_name)
        doc.add_paragraph("Capacity: " + project.capacity)
        doc.add_paragraph("\n")
        
        doc.add_page_break()     
        doc.add_paragraph("\n")

        # Add machine details
        for index, machine in enumerate(machines, start=1):  # Add numbering
            machine_name = machine.oSec00Field03
            section_titles = []

            if machine_name == "DataSheetNS":
                machine_name = "Manual Screen" 
                section_titles = ["General Data", "Design Data", "Material Data", "Channel Data", " ", " ", " ", " ", " ", " "]

            if machine_name == "DataSheetMSc":
                machine_name = "Mechanical Screen" 
                section_titles = ["General Data", "Design Data", "Gearmotor Data", "Control panel Data", "Material Data", "Other Data", " ", " ", " ", " "]

            if machine_name == "DataSheetBC":
                machine_name = "Belt Conveyor"
                section_titles = ["General Data", "Design Data", "Gearbox Data", "Motor Data", "Material Data", " ", " ", " ", " ", " "]

            if machine_name == "DataSheetCO":
                machine_name = "Container"
                section_titles = ["General Data", "Design Data", "Material Data", " ", " ", " ", " ", " ", " ", " "]

            if machine_name == "DataSheetGR":
                machine_name = "Gritremoval"
                section_titles = ["General Data", "Design Data", "Walkway, Handrail, Wheel Data", "Scrapper Data", "Gearmotor Data", "Scrapper Data", "Drive unit", "Control panel Data", "Material Data ", " "]

            if machine_name == "DataSheetSS":
                machine_name = "Sand Silo"

            if machine_name == "DataSheetPS":
                machine_name = "Primary Sedimentation"

            if machine_name == "DataSheetQV":
                machine_name = "Quick Valve"

            if machine_name == "DataSheetTV":
                machine_name = "Telescopic Valve"
                
            if machine_name == "DataSheetTH":
                machine_name = "Sludge Thickener"

            # Add machine title with font size 14 and numbering
            machine_title = doc.add_paragraph(f"{index}. {machine_name}", style="Heading3")
            machine_title.runs[0].font.size = Pt(14)

            for i in range(1, 11):  # Loop from Sec01 to Sec10
                section_name = f"Sec{i:02d}"
                section_data = [("Field", "Value")]

                for j in range(1, 21, 2):  # Step by 2 to avoid duplication
                    key = getattr(machine, f"o{section_name}Field{j:02d}", "").strip()
                    value = getattr(machine, f"o{section_name}Field{j+1:02d}", "").strip()

                    if key and value and key.lower() != "oooo" and value.lower() != "oooo":
                        section_data.append((key, value))

                if len(section_data) > 1:  # If the section has valid data, create a table
                    section_title = section_titles[i-1] if i-1 < len(section_titles) else f"Section {i}"
                    doc.add_paragraph(f"{section_name}: {section_title}", style="Heading3")  # Only one title now

                    add_table(doc, section_data)  # Removed redundant title

            doc.add_page_break()     

        # Save the document to a response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={project.name}_report.docx'
        doc.save(response)
        return response

    except Project.DoesNotExist:
        return HttpResponse("Project not found", status=404)