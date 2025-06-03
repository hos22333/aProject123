import requests

import time
import os
import ezdxf
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt


from django.shortcuts import render
from config import settings
from .models import APP_Project
from .drive import create_folder, service, check_folder_exists, get_folder_id_by_name, upload_files, upload_files_directly
from Apps.aAppMechanical.models import UserCompany
from Apps.aAppSubmittal.models import Machine
from Apps.aAppSubmittal.models import AddMachine
from Apps.aAppSubmittal.models import DXF_data
from Apps.aAppMechanical.models import aLogEntry
from Apps.aAppCalculation.models import modelcalc
from Apps.aAppSubmittal.views import get_user_company
from django.http import FileResponse, HttpResponse
from django.shortcuts import render, redirect
from django.utils.timezone import now 
from django.contrib.auth.models import User
from django.utils.text import slugify
from io import BytesIO, StringIO
from ezdxf import recover
from ezdxf.addons.drawing import Frontend, RenderContext
from ezdxf.addons.drawing import matplotlib as ezdxf_matplotlib


from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,  WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.shared import Inches
from docx.shared import Pt

from fpdf import FPDF
from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from openpyxl import Workbook



def save_word_pdf_submittal_report(user, project_id, logo, color):  
    print("start, save_word_pdf_submittal_report, project_id : ", project_id)
    
    def add_table(doc, data, title=None):
        """Creates a borderless table and applies a background color to the header."""
        if title:
            doc.add_heading(title, level=3)

        table = doc.add_table(rows=len(data), cols=2)

        if logo == "LogoBBB": 
            # Remove all table borders manually
            tbl = table._tbl  # Get the table's XML element
            tblPr = tbl.find(ns.qn("w:tblPr"))  # Find existing table properties

            if tblPr is None:
                tblPr = OxmlElement("w:tblPr")  # Create table properties if missing
                tbl.insert(0, tblPr)  # Insert as the first child of <w:tbl>

            tblBorders = OxmlElement("w:tblBorders")  # Create <w:tblBorders>
            for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(ns.qn("w:val"), "nil")  # Remove the border
                tblBorders.append(border)

            tblPr.append(tblBorders)  # Append border settings to the table properties
        elif logo == "LogoAAA":
            table.style = "Table Grid"

        for i, row in enumerate(data):
            for j, text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = text

                # Apply background color only to the header row (first row)
                if i == 0:
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(ns.qn("w:val"), "clear")  # Set shading value
                    shading_elm.set(ns.qn("w:fill"), color)  # Light blue color
                    cell._tc.get_or_add_tcPr().append(shading_elm)
   
    
    
    def add_header_footer(doc):
        """Adds header and footer with page numbers in the format 'Page X of Y'."""
        section = doc.sections[0]
    
        # Header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        # header_para.add_run("Company Name\n")
        # header_para.add_run("Project Name\n")
        # header_para.add_run("Date: " + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "\n")
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Adding logo
        run_logo = header_para.add_run()  # Corrected reference to header paragraph
        try:
            print(f"static/aLogo/{logo}.png")
            run_logo.add_picture(f"static/aLogo/{logo}.png", width=Inches(7.0))  # Adjust width as needed
        except Exception as e:
            print(f"Error adding logo: {e}")

        
    
        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Add "Page X of Y" format
        run = footer_para.add_run("Page ")
    
        # PAGE field (Current Page Number)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(ns.qn("w:fldCharType"), "begin")
    
        instrText1 = OxmlElement("w:instrText")
        instrText1.set(ns.qn("xml:space"), "preserve")
        instrText1.text = "PAGE"
    
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar1)
        run._r.append(instrText1)
        run._r.append(fldChar2)
    
        run.add_text(" of ")
    
        # NUMPAGES field (Total Number of Pages)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(ns.qn("w:fldCharType"), "begin")
    
        instrText2 = OxmlElement("w:instrText")
        instrText2.set(ns.qn("xml:space"), "preserve")
        instrText2.text = "NUMPAGES"
    
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)
    
    
    def add_colored_heading(doc, text, level, color):
        """Adds a heading with color."""
        heading = doc.add_paragraph()
        run = heading.add_run(text)
        run.bold = True
        run.font.size = Pt(14) if level == 1 else Pt(12)
        run.font.color.rgb = color
        heading.style = f"Heading {level}"
    
    
    # Define a custom class for the PDF layout
    class PDF(FPDF):
        def __init__(self):
            super().__init__()
            self.set_auto_page_break(auto=True, margin=15)

            # Add the TrueType Unicode font (DejaVuSans)
            self.add_font('DejaVu', '', 'static/aApp1/fonts/DejaVu/DejaVuSans.ttf', uni=True)  # Regular
            self.add_font('DejaVu', 'B', 'static/aApp1/fonts/DejaVu/DejaVuSans-Bold.ttf', uni=True)  # Bold
            self.add_font('DejaVu', 'I', 'static/aApp1/fonts/DejaVu/DejaVuSerif-Italic.ttf', uni=True)  # Italic
            self.set_font('DejaVu', '', 12)

        def header_footer(self, projectname, clientname, capacity):
            self.set_y(50)
            
            
            page_height = self.h
            top_margin = self.t_margin
            bottom_margin = self.b_margin

            
            lines = [
                ("Project Name: ", "", 14, (0, 0, 0)),
                (projectname, "", 14, (0, 0, 0)),
                ("Client Name: ", "", 14, (0, 0, 0)),
                (clientname, "", 14, (0, 0, 0)),
                ("Capacity: ", "", 14, (0, 0, 0)),
                (capacity, "", 14, (0, 0, 0)),
            ]

            
            line_height = 16
            total_content_height = line_height * len(lines)

            
            y_position = (page_height - total_content_height) / 2
            self.set_y(y_position)

            
            for text, style, size, color in lines:
                self.set_font("DejaVu", style, size)
                self.set_text_color(*color)
                self.cell(0, line_height, text, ln=True, align="C")

        def header(self):
            # Logo
            page_width = self.w - 2 * self.l_margin
            
            try:
                self.image(f"static/aLogo/{logo}.png", x=self.l_margin, y=10, w=page_width)  

                # Move cursor below image to avoid overlapping next content
                self.set_y(10 + 50)
            except Exception as e:
                print(f"Error loading logo: {e}")



        def footer(self):
            self.set_y(-15)  # Position 15 mm from bottom
            self.set_font("DejaVu", "I", 8)
            self.set_text_color(128)

            # Add "Page X of Y"
            self.cell(0, 10, f"Page {self.page_no()} of {{nb}}", align='C')

        def colored_header(self, number, name):
            self.set_font("DejaVu", "B", 16)
            self.set_text_color(33, 66, 133)
            self.cell(0, 10, f"{number}. {name}", ln=True)

        def section_title(self, title):
            self.set_font("DejaVu", "B", 12)
            self.set_text_color(33, 66, 133)
            self.ln(5)
            self.cell(0, 10, title, ln=True)

        def add_table(self, data):
            self.set_font("DejaVu", "B", 10)
            if logo == "LogoAAA":
                self.set_fill_color(255, 153, 0)
                tableborder = 1
                self.set_text_color(0)
                self.cell(80, 8, "Field", border=tableborder, fill=True)
                self.cell(110, 8, "Value", border=tableborder, ln=True, fill=True)
            elif logo == "LogoBBB": 
                self.set_fill_color(255, 255, 255)
                tableborder = 0
                self.set_text_color(0)
                self.cell(80, 8, " ", border=tableborder, fill=True)
                self.cell(110, 8, " ", border=tableborder, ln=True, fill=True)

            self.set_font("DejaVu", "", 10)
            for field, value in data:
                self.cell(80, 8, field, border=tableborder)
                self.cell(110, 8, value, border=tableborder, ln=True)
    
    
    try:
        
        ###LOG
        aLogEntry.objects.create(
                user=user,
                message=f"at {now()} {user} accessed Load  "
            )
        print(f"at {now()} {User} accessed Download Report")
        ###LOG

        aCompany = UserCompany.objects.get(user=user)
        company_id = aCompany.company
        project = APP_Project.objects.get(id=project_id)
        machines = Machine.objects.filter(project=project)
        
        # Create a Word document
        doc = Document()

        pdf = PDF()

        # Add header and footer with page numbers
        add_header_footer(doc)

        pdf.alias_nb_pages()  # Important for "of {nb}" to work
        pdf.add_page()

        pdf.header_footer(project.name,  project.client_name, project.capacity) 

        for _ in range(4):  # Adjust this number based on how centered you want it
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Centered content
        lines = [
            "Project Name: ",
            project.name,
            "Client Name: ",
            project.client_name,
            "Capacity: ",
            project.capacity,
        ]

        for line in lines:
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.size = Pt(25) 
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing after if needed
        para = doc.add_paragraph("\n")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()     
        doc.add_paragraph("\n")

        # Add machine details
        for index, machine in enumerate(machines, start=1):  # Add numbering
            machine_name = machine.oSec00Field03
            machine_DB = machine.oSec00Field03
            machine_ID = machine.id
            
            try:
                themachines = AddMachine.objects.get(nameDB=machine_DB, company=company_id)
            except AddMachine.DoesNotExist:
                print(f"Skipping machine '{machine_DB}' for company ID {company_id}: not found in AddMachine.")
                continue  # Skip this machine and continue with the rest
            
            sheet_key = themachines.keyValue
            themachinename = themachines.nameMachine
            General_saved_DXF_ALL(user, machine_ID, sheet_key, project_id)
            SavedFullDrawing(user, machine_ID, sheet_key, project_id)
            
            #SavedFullDrawing(request, machine_ID, sheet_key)


            
            section_titles = [" ", " ", " ", " ", " ", " ", " ", " ", " ", " "]

            if machine_name == "DataSheetNS":
                machine_name = "Manual Screen" 
                section_titles = ["General Data", "Design Data", "Material Data", "Channel Data", " ", " ", " ", " ", " ", " "]

            elif machine_name == "DataSheetMS":
                machine_name = "Mechanical Screen" 
                section_titles = ["General Data", "Design Data", "Gearmotor Data", "Control panel Data", "Material Data", "Other Data", " ", " ", " ", " "]

            elif machine_name == "DataSheetBC":
                machine_name = "Belt Conveyor"
                section_titles = ["General Data", "Design Data", "Gearbox Data", "Motor Data", "Material Data", " ", " ", " ", " ", " "]

            elif machine_name == "DataSheetCO":
                machine_name = "Container"
                section_titles = ["General Data", "Design Data", "Material Data", " ", " ", " ", " ", " ", " ", " "]

            elif machine_name == "DataSheetGR":
                machine_name = "Gritremoval"
                section_titles = ["General Data", "Design Data", "Walkway, Handrail, Wheel Data", "Scrapper Data", "Gearmotor Data", "Scrapper Data", "Drive unit", "Control panel Data", "Material Data ", " "]

            # Add machine title with font size 14 and numbering
            machine_title = doc.add_paragraph(f"{index}. {themachinename}", style="Heading3")
            machine_title.runs[0].font.size = Pt(14)
            
            pdf.alias_nb_pages()  # Important for "of {nb}" to work
            pdf.add_page()
            pdf.colored_header(index, themachinename)

            for i in range(1, 11):  # Loop from Sec01 to Sec10
                section_name = f"Sec{i:02d}"
                if aCompany.company.nameCompanies == "AAAA":
                    section_data = [("Field", "Value")]
                elif aCompany.company.nameCompanies == "BBBB":
                    section_data = [(" ", " ")]
                pdf_section_data = []

                for j in range(1, 21, 2):  # Step by 2 to avoid duplication
                    key = getattr(machine, f"o{section_name}Field{j:02d}", "").strip()
                    value = getattr(machine, f"o{section_name}Field{j+1:02d}", "").strip()

                    if key and value and key.lower() != "oooo" and value.lower() != "oooo":
                        section_data.append((key, value))
                        pdf_section_data.append((key, value))

                if len(section_data) > 1:  # If the section has valid data, create a table
                    if aCompany.company.nameCompanies == "AAAA":
                        section_title = section_titles[i-1] if i-1 < len(section_titles) else f"Section {i}"
                        doc.add_paragraph(f"{section_name}: {section_title}", style="Heading3")  # Only one title now

                        pdf.section_title(f"{section_name}: {section_title}")

                    add_table(doc, section_data)  # Removed redundant title

                    pdf.add_table(pdf_section_data)

            doc.add_page_break()     

        # Define the folder path
        company_slug = slugify(project.company.nameCompanies)
        company_name = company_slug.upper()
        project_slug = slugify(project.name)
        folder_name = slugify(f"{project_id}_{company_slug}_{project_slug}")

        project_folder = os.path.join(settings.BASE_DIR, 'static', 'aReports', company_slug, folder_name)
        os.makedirs(project_folder, exist_ok=True)  # Create folder if it doesn't exist

        # Save the file to that path
        file_path = os.path.join(project_folder, f"{project_slug}_report.docx")
        doc.save(file_path)

        """ doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0) """
        
        file_name = f"{project_slug}_report.docx"
        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        folder_exist, folder_data = check_folder_exists(service, "aReports")
        if folder_exist == True :
            folder_id = folder_data['id']
            company_folder_exist, company_folder_data = check_folder_exists(service, company_name, folder_id)
            if company_folder_exist == True:
                company_folder_id = get_folder_id_by_name(service, company_name)
                project_folder_exist, project_folder_data = check_folder_exists(service, folder_name, company_folder_id)
                if project_folder_exist == True:
                    project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                    upload_files(service,file_path, file_name, mime_type, project_folder_id)
                    """ upload_files_directly(service, doc_buffer, file_name, mime_type, project_folder_id) """
                else:
                    create_folder(service, folder_name, company_folder_id)
                    project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                    upload_files(service,file_path, file_name, mime_type, project_folder_id)
                    """ upload_files_directly(service, doc_buffer, file_name, mime_type, project_folder_id) """
            else:
                create_folder(service, company_name, folder_id)
                company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
                create_folder(service, folder_name, company_folder_id)
                project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                upload_files(service,file_path, file_name, mime_type, project_folder_id)
                """ upload_files_directly(service, doc_buffer, file_name, mime_type, project_folder_id) """
        else:
            create_folder(service, "aReports")
            folder_id = get_folder_id_by_name(service, "aReports")
            create_folder(service, company_name, folder_id)
            company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
            create_folder(service, folder_name, company_folder_id)
            project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
            upload_files(service,file_path, file_name, mime_type, project_folder_id)

        pdf_file_path = os.path.join(project_folder, f"{project_slug}_report.pdf")
        pdf.output(pdf_file_path)

        """ pdf_bytes = pdf.output(dest='S').encode('latin1')  # Get PDF as bytes
        pdf_buffer = BytesIO(pdf_bytes)
        pdf_buffer.seek(0) """


        file_name = f"{project_slug}_report.pdf"
        mime_type = 'application/pdf'

        folder_exist, folder_data = check_folder_exists(service, "aReports")
        if folder_exist == True :
            folder_id = folder_data['id']
            company_folder_exist, company_folder_data = check_folder_exists(service, company_name, folder_id)
            if company_folder_exist == True:
                company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
                project_folder_exist, project_folder_data = check_folder_exists(service, folder_name, company_folder_id)
                if project_folder_exist == True:
                    project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                    upload_files(service,pdf_file_path, file_name, mime_type, project_folder_id)
                    """ upload_files_directly(service, pdf_buffer, file_name, mime_type, project_folder_id) """
                else:
                    create_folder(service, folder_name, company_folder_id)
                    project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                    upload_files(service,pdf_file_path, file_name, mime_type, project_folder_id)
                    """ upload_files_directly(service, pdf_buffer, file_name, mime_type, project_folder_id) """
            else:
                create_folder(service, company_name, folder_id)
                company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
                create_folder(service, folder_name, company_folder_id)
                project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                upload_files(service,pdf_file_path, file_name, mime_type, project_folder_id)
                """ upload_files_directly(service, pdf_buffer, file_name, mime_type, project_folder_id) """
        else:
            create_folder(service, "aReports")
            folder_id = get_folder_id_by_name(service, "aReports")
            create_folder(service, company_name, folder_id)
            company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
            create_folder(service, folder_name, company_folder_id)
            project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
            upload_files(service,pdf_file_path, file_name, mime_type, project_folder_id)

        return HttpResponse(status=204)

        

    except APP_Project.DoesNotExist:
        return HttpResponse("Project not found", status=404)



def convert_dxf_to_pdf_ezdxf(input_path, output_path):
    print("start, convert_dxf_to_pdf_ezdxf")
    
    try:
        doc = ezdxf.recover.readfile(input_path)[0]
        ezdxf_matplotlib.qsave(
            doc.modelspace(),
            output_path,
            dpi=300,
            size_inches=(36, 24),  # Custom paper size
            bg='#FFFFFF'
        )

        print("PDF successfully saved")

    except Exception as e:
        print(f"Error converting DXF to PDF: {e}")
        raise




def is_number(value):
    try:
        float(value)
        return True
    except ValueError:
        return False
    
def resolve_fieldvalue(machine, fieldvalue):
    if is_number(fieldvalue):
        return fieldvalue

    # Loop through sections 1 to 10, and odd fields 01–19
    for sec in range(1, 11):
        for i in range(1, 21, 2):
            odd_field = f"oSec{str(sec).zfill(2)}Field{str(i).zfill(2)}"
            even_field = f"oSec{str(sec).zfill(2)}Field{str(i+1).zfill(2)}"
            if fieldvalue == getattr(machine, odd_field, None):
                return getattr(machine, even_field, fieldvalue)

    # If not found, return as-is
    return fieldvalue



# DXF Download Views
def General_saved_DXF_ALL(user, aMachine_ID, aType, project_id): 
    print("start, General_saved_DXF_ALL", project_id, aMachine_ID, aType)
    
    
    
    # Helper function to modify DXF files
    def modify_saved_dxf_file(static_path, modified_path, modifications):
        doc = ezdxf.readfile(static_path)

        for entity in doc.modelspace().query("DIMENSION"):
            if entity.dxf.text in modifications:
                entity.dxf.text = modifications[entity.dxf.text]

            # Update text height and arrow size
            dimstyle = doc.dimstyles.get(entity.dxf.dimstyle)
            if dimstyle:
                dimstyle.dxf.dimtxt = 0.1  # Set text height
                dimstyle.dxf.dimasz = 0.1  # Set arrow size

            entity.render()

        doc.saveas(modified_path)


    
    # Helper function to define DXF paths
    def get_saved_dxf_paths(user_company, category, project_id, output_filename, aType):
        print("start, get_saved_dxf_paths", user_company, category, project_id, output_filename)
        
        
        # Get project info
        project = APP_Project.objects.get(id=project_id)
        company_slug = slugify(user_company.nameCompanies)
        project_slug = slugify(project.name)
        folder_name = f"{project_id}_{company_slug}_{project_slug}"

        
        dxf_name = None
        if company_slug == "aaaa":
            dxf_name = f"AAA_{category}"
        elif company_slug == "bbbb":
            dxf_name = category

        machine = AddMachine.objects.get(keyValue = aType)
        dxffile_model_name = machine.primarynameDXF
        
        if dxffile_model_name not in ["", None] :
            input_filename = f"{dxffile_model_name}.dxf"
        else :
            input_filename = f"{dxf_name}.dxf"
        
        # Load original path (base DXF)
        company_dxf_path = {
            1: os.path.join(settings.BASE_DIR, "static", "aDxfs", input_filename),
            2: os.path.join(settings.BASE_DIR, "static", "aDxfs", input_filename),
        }
        static_path = company_dxf_path.get(user_company.id)
        if not static_path or not os.path.exists(static_path):
            #raise FileNotFoundError(f"DXF not found: {static_path}")
            print(f"DXF not found: {static_path}")

        # Target output path for modified DXF
        target_dir = os.path.join(
            settings.BASE_DIR,
            "static",
            "aReports",
            company_slug.upper(),
            folder_name
        )
        os.makedirs(target_dir, exist_ok=True)

        
        modified_path = os.path.join(target_dir, output_filename)

        return static_path, modified_path
        

        
    # Main DXF Processing Function
    def process_saved_dxf(user, aMachine_ID, category, project_id, modifications, output_filename, aType, company_name, folder_name):
        
        # Log the request
        aLogEntry.objects.create(
            user=user,
            message=f"at {now()} {user} Download DXF {category} {aMachine_ID}"
        )

        try:
            user_company = UserCompany.objects.get(user=user).company
        except UserCompany.DoesNotExist:
            return HttpResponse("Unauthorized", status=403)

        static_path, modified_path = get_saved_dxf_paths(user_company, category, project_id, output_filename, aType)
        if not os.path.exists(static_path):
            return HttpResponse("File not found", status=404)

        machine = Machine.objects.get(id=aMachine_ID)

        
        modify_saved_dxf_file(static_path, modified_path, modifications(machine))

        file_name = output_filename
        mime_type = 'application/dxf'

        folder_exist, folder_data = check_folder_exists(service, "aReports")
        if folder_exist == True :
            folder_id = folder_data['id']
            company_folder_exist, company_folder_data = check_folder_exists(service, company_name, folder_id)
            if company_folder_exist == True:
                company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
                project_folder_exist, project_folder_data = check_folder_exists(service, folder_name, company_folder_id)
                if project_folder_exist == True:
                    project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                    upload_files(service,modified_path, file_name, mime_type, project_folder_id)
                else:
                    create_folder(service, folder_name, company_folder_id)
                    project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                    upload_files(service,modified_path, file_name, mime_type, project_folder_id)
            else:
                create_folder(service, company_name, folder_id)
                company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
                create_folder(service, folder_name, company_folder_id)
                project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                upload_files(service,modified_path, file_name, mime_type, project_folder_id)
        else:
            create_folder(service, "aReports")
            folder_id = get_folder_id_by_name(service, "aReports")
            create_folder(service, company_name, folder_id)
            company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
            create_folder(service, folder_name, company_folder_id)
            project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
            upload_files(service,modified_path, file_name, mime_type, project_folder_id)


        # Define PDF output path
        pdf_output_path = modified_path.replace(".dxf", ".pdf")

        try:
            convert_dxf_to_pdf_ezdxf(modified_path, pdf_output_path)
            
            base, ext = os.path.splitext(output_filename)
            new_filename = base + ".pdf"
            file_name = new_filename
            mime_type = 'application/pdf'

            folder_exist, folder_data = check_folder_exists(service, "aReports")
            if folder_exist == True :
                folder_id = folder_data['id']
                company_folder_exist, company_folder_data = check_folder_exists(service, company_name, folder_id)
                if company_folder_exist == True:
                    company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
                    project_folder_exist, project_folder_data = check_folder_exists(service, folder_name, company_folder_id)
                    if project_folder_exist == True:
                        project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                        upload_files(service,pdf_output_path, file_name, mime_type, project_folder_id)
                    else:
                        create_folder(service, folder_name, company_folder_id)
                        project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                        upload_files(service,pdf_output_path, file_name, mime_type, project_folder_id)
                else:
                    create_folder(service, company_name, folder_id)
                    company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
                    create_folder(service, folder_name, company_folder_id)
                    project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                    upload_files(service,pdf_output_path, file_name, mime_type, project_folder_id)
            else:
                create_folder(service, "aReports")
                folder_id = get_folder_id_by_name(service, "aReports")
                create_folder(service, company_name, folder_id)
                company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
                create_folder(service, folder_name, company_folder_id)
                project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                upload_files(service,pdf_output_path, file_name, mime_type, project_folder_id)

                
            print(f"PDF saved to {pdf_output_path}")

        except Exception as e:
            print("DXF to PDF conversion failed:", e)
            return HttpResponse("DXF to PDF conversion failed", status=500)


    
    # Redirect unauthenticated users
    if not user.is_authenticated:
        return redirect("login")  
    
    
    ###LOG
    aLogEntry.objects.create(
            user=user,
            message=f"at {now()} {user} DXF download {aType} "
        )
    
    themachine = AddMachine.objects.get(keyValue = aType)
    file_model_name = themachine.nameDXF
    sheetkey = aType[0:-2]

    if file_model_name not in ["", None] :
        file_name = file_model_name
    else :
        file_name = f"{sheetkey}_new"

    # Get the company of the logged-in user    
    user_company = None
    company_name = None
    if user.is_authenticated:
        try:
            user_company = UserCompany.objects.get(user=user).company
            company_name = user_company.nameCompanies
        except UserCompany.DoesNotExist:
            user_company = None
    
    project = APP_Project.objects.get(id=project_id)

    # Define the folder path
    company_slug = slugify(company_name)
    project_slug = slugify(project.name)
    folder_name = slugify(f"{project_id}_{company_slug}_{project_slug}")

    datas = DXF_data.objects.filter(sheetkey = sheetkey)


    return process_saved_dxf(
        user,
        aMachine_ID,
        sheetkey,
        project_id,
        lambda machine: {
            data.fieldname : resolve_fieldvalue(machine, data.fieldvalue)
            for data in datas
        },
        f"{file_name}.dxf",
        aType,
        company_name,
        folder_name
    )
   
   
   
   
   
   
   
   
   
   
   
   
   
   
   
    
# DXF Download Views
def SavedFullDrawing(user, aMachine_ID, aType, project_id): 
    print("start, SavedFullDrawing", aMachine_ID, aType)
    
    # Helper function to modify DXF files
    def SavedFullDrawing_modify_dxf_file(static_path, modified_path, modifications):
        doc = ezdxf.readfile(static_path)
        """ msp = doc.modelspace() """

        # Apply modifications here (using your `modifications` logic)
        """ modifications(msp) """

        # for entity in doc.modelspace().query("DIMENSION"):
        #     if entity.dxf.text in modifications:
        #         entity.dxf.text = modifications[entity.dxf.text]

        #     # Update text height and arrow size
        #     dimstyle = doc.dimstyles.get(entity.dxf.dimstyle)
        #     if dimstyle:
        #         dimstyle.dxf.dimtxt = 0.1  # Set text height
        #         dimstyle.dxf.dimasz = 0.1  # Set arrow size

        #     entity.render()

        """ doc.write(buffer)   """
        doc.saveas(modified_path)
        
        
    # Helper function to define DXF paths
    def get_saved_dxf_paths(user_company, category, project_id, output_filename, aType):
        print("start, get_saved_dxf_paths", user_company, category, project_id, output_filename)
        
        
        # Get project info
        project = APP_Project.objects.get(id=project_id)
        company_slug = slugify(user_company.nameCompanies)
        project_slug = slugify(project.name)
        folder_name = f"{project_id}_{company_slug}_{project_slug}"

        fulldrawing_dxf_name = None
        if company_slug == "aaaa":
            fulldrawing_dxf_name = f"Full Drawing {category}"
        elif company_slug == "bbbb":
            fulldrawing_dxf_name = f"BBB_Full Drawing {category}"

        machine = AddMachine.objects.get(keyValue = aType)
        fulldrawingdxffile_model_name = machine.primarynameFullDrawing
        
        if fulldrawingdxffile_model_name not in ["", None] :
            input_filename = f"{fulldrawingdxffile_model_name}.dxf"
        else :
            input_filename = f"{fulldrawing_dxf_name}.dxf"

        # Load original path (base DXF)
        company_dxf_path = {
            1: os.path.join(settings.BASE_DIR, "static", "aDxfs", input_filename),
            2: os.path.join(settings.BASE_DIR, "static", "aDxfs", input_filename),
        }
        static_path = company_dxf_path.get(user_company.id)
        if not static_path or not os.path.exists(static_path):
            #raise FileNotFoundError(f"DXF not found: {static_path}")
            print(f"Full Drawing DXF not found: {static_path}")

        # Target output path for modified DXF
        target_dir = os.path.join(
            settings.BASE_DIR,
            "static",
            "aReports",
            company_slug.upper(),
            folder_name
        )
        os.makedirs(target_dir, exist_ok=True)

        
        modified_path = os.path.join(target_dir, output_filename)

        return static_path, modified_path
        
        
        
        
    # Main DXF Processing Function
    def SavedFullDrawing_process_dxf(user, aMachine_ID, category, project_id, modifications, output_filename, aType):

        try:
            user_company = UserCompany.objects.get(user=user).company
        except UserCompany.DoesNotExist:
            return HttpResponse("Unauthorized", status=403)

        # Get machine and project
        machine = Machine.objects.get(id=aMachine_ID)
        project = machine.project  # Assumes FK from Machine to Project
        company_slug = slugify(user_company.nameCompanies)
        company_name = company_slug.upper()
        project_slug = slugify(project.name)
        folder_name = f"{project.id}_{company_slug}_{project_slug}"

        

        static_path, modified_path = get_saved_dxf_paths(user_company, category, project_id, output_filename, aType)
        
        
        #static_path  = os.path.join(settings.BASE_DIR, "static", "aDxfs", "AAA", "FullDrawing", f"Full Drawing {category}.dxf")
        
        
        if not os.path.exists(static_path):
            return HttpResponse("File not found", status=404)

        # Use a BytesIO buffer to store the modified file in memory
        """ text_buffer  = StringIO() """
        # Destination path
        target_dir = os.path.join(
            settings.BASE_DIR, "static", "aReports", company_slug.upper(), folder_name
        )
        os.makedirs(target_dir, exist_ok=True)
        if output_filename != "None.dxf" :
            modified_path = os.path.join(target_dir, output_filename)
        else :
            modified_path = os.path.join(target_dir, f"{category}_newFullDrawing.dxf")

        machine = Machine.objects.get(id=aMachine_ID)

        """ mod_func = modifications(machine)

        # Modify the DXF and write to the buffer
        SavedFullDrawing_modify_dxf_file(static_path, text_buffer , mod_func)
        text_data = text_buffer.getvalue().encode("utf-8")  # Convert to bytes
        dxf_buffer = BytesIO(text_data)  # For uploading
        dxf_buffer.seek(0)

        # Prepare upload
        file_name = output_filename if output_filename != "None.dxf" else f"{category}_newFullDrawing.dxf"
        mime_type = 'application/dxf' """
        
        SavedFullDrawing_modify_dxf_file(static_path, modified_path, modifications(machine))
        
        file_name = output_filename
        mime_type = 'application/dxf'

        folder_exist, folder_data = check_folder_exists(service, "aReports")
        if folder_exist == True :
            folder_id = folder_data['id']
            company_folder_exist, company_folder_data = check_folder_exists(service, company_name, folder_id)
            if company_folder_exist == True:
                company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
                project_folder_exist, project_folder_data = check_folder_exists(service, folder_name, company_folder_id)
                if project_folder_exist == True:
                    project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                    upload_files(service,modified_path, file_name, mime_type, project_folder_id)
                    """ upload_files_directly(service, dxf_buffer, file_name, mime_type, project_folder_id) """
                else:
                    create_folder(service, folder_name, company_folder_id)
                    project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                    upload_files(service,modified_path, file_name, mime_type, project_folder_id)
                    """ upload_files_directly(service, dxf_buffer, file_name, mime_type, project_folder_id) """
            else:
                create_folder(service, company_name, folder_id)
                company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
                create_folder(service, folder_name, company_folder_id)
                project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                upload_files(service,modified_path, file_name, mime_type, project_folder_id)
                """ upload_files_directly(service, dxf_buffer, file_name, mime_type, project_folder_id) """
        else:
            create_folder(service, "aReports")
            folder_id = get_folder_id_by_name(service, "aReports")
            create_folder(service, company_name, folder_id)
            company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
            create_folder(service, folder_name, company_folder_id)
            project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
            upload_files(service,modified_path, file_name, mime_type, project_folder_id)
        

        

        return HttpResponse("Invalid request", status=400)
        

    
    # Redirect unauthenticated users
    if not user.is_authenticated:
        return redirect("login") 
    
    
    themachine = AddMachine.objects.get(keyValue = aType)
    file_model_name = themachine.nameFullDrawing
    sheetkey = aType[0:-2]

    if file_model_name not in ["", None] :
        file_name = file_model_name
    else :
        file_name = f"{sheetkey}_newFullDrawing"

    # Get the company of the logged-in user    
    user_company = None
    if user.is_authenticated:
        try:
            user_company = UserCompany.objects.get(user=user).company
        except UserCompany.DoesNotExist:
            user_company = None 
        
    ###LOG
    aLogEntry.objects.create(
            user=user,
            message=f"at {now()} {user} DXF download {aType} "
        )

    datas = DXF_data.objects.filter(sheetkey = sheetkey)


    return SavedFullDrawing_process_dxf(
        user,
        aMachine_ID,
        sheetkey,
        project_id,
        lambda machine: {
            data.fieldname : resolve_fieldvalue(machine, data.fieldvalue)
            for data in datas
        },
        f"{file_name}.dxf",
        aType
    )











def save_word_pdf_calculation_report(user, project_id, logo, color):
    
    def add_table(doc, data, title=None):
        """Creates a borderless table and applies a background color to the header."""
        if title:
            doc.add_heading(title, level=3)

        table = doc.add_table(rows=len(data), cols=2)

        if logo == "LogoBBB": 
            # Remove all table borders manually
            tbl = table._tbl  # Get the table's XML element
            tblPr = tbl.find(ns.qn("w:tblPr"))  # Find existing table properties

            if tblPr is None:
                tblPr = OxmlElement("w:tblPr")  # Create table properties if missing
                tbl.insert(0, tblPr)  # Insert as the first child of <w:tbl>

            tblBorders = OxmlElement("w:tblBorders")  # Create <w:tblBorders>
            for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(ns.qn("w:val"), "nil")  # Remove the border
                tblBorders.append(border)

            tblPr.append(tblBorders)  # Append border settings to the table properties
        elif logo == "LogoAAA":
            table.style = "Table Grid"

        for i, row in enumerate(data):
            for j, text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = text

                # Apply background color only to the header row (first row)
                if i == 0:
                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(ns.qn("w:val"), "clear")  # Set shading value
                    shading_elm.set(ns.qn("w:fill"), color)  # Light blue color
                    cell._tc.get_or_add_tcPr().append(shading_elm)
   
    
    
    def add_header_footer(doc):
        """Adds header and footer with page numbers in the format 'Page X of Y'."""
        section = doc.sections[0]
    
        # Header
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        # header_para.add_run("Company Name\n")
        # header_para.add_run("Project Name\n")
        # header_para.add_run("Date: " + datetime.now().strftime('%Y-%m-%d %H:%M:%S') + "\n")
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Adding logo
        run_logo = header_para.add_run()  # Corrected reference to header paragraph
        try:
            run_logo.add_picture(f"static/aLogo/{logo}.png", width=Inches(7.0))  # Adjust width as needed
        except Exception as e:
            print(f"Error adding logo: {e}")

        
    
        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
        # Add "Page X of Y" format
        run = footer_para.add_run("Page ")
    
        # PAGE field (Current Page Number)
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(ns.qn("w:fldCharType"), "begin")
    
        instrText1 = OxmlElement("w:instrText")
        instrText1.set(ns.qn("xml:space"), "preserve")
        instrText1.text = "PAGE"
    
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar1)
        run._r.append(instrText1)
        run._r.append(fldChar2)
    
        run.add_text(" of ")
    
        # NUMPAGES field (Total Number of Pages)
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(ns.qn("w:fldCharType"), "begin")
    
        instrText2 = OxmlElement("w:instrText")
        instrText2.set(ns.qn("xml:space"), "preserve")
        instrText2.text = "NUMPAGES"
    
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(ns.qn("w:fldCharType"), "end")
    
        run._r.append(fldChar3)
        run._r.append(instrText2)
        run._r.append(fldChar4)
    
    
    def add_colored_heading(doc, text, level, color):
        """Adds a heading with color."""
        heading = doc.add_paragraph()
        run = heading.add_run(text)
        run.bold = True
        run.font.size = Pt(14) if level == 1 else Pt(12)
        run.font.color.rgb = color
        heading.style = f"Heading {level}"
    
    
    # Define a custom class for the PDF layout
    class PDF(FPDF):
        def __init__(self):
            super().__init__()
            self.set_auto_page_break(auto=True, margin=15)

            # Add the TrueType Unicode font (DejaVuSans)
            self.add_font('DejaVu', '', 'static/aApp1/fonts/DejaVu/DejaVuSans.ttf', uni=True)  # Regular
            self.add_font('DejaVu', 'B', 'static/aApp1/fonts/DejaVu/DejaVuSans-Bold.ttf', uni=True)  # Bold
            self.add_font('DejaVu', 'I', 'static/aApp1/fonts/DejaVu/DejaVuSerif-Italic.ttf', uni=True)  # Italic
            self.set_font('DejaVu', '', 12)

        def header_footer(self, projectname, clientname, capacity):
            self.set_y(50)
            
            page_height = self.h
            top_margin = self.t_margin
            bottom_margin = self.b_margin

            
            lines = [
                ("Project Name: ", "", 14, (0, 0, 0)),
                (projectname, "", 14, (0, 0, 0)),
                ("Client Name: ", "", 14, (0, 0, 0)),
                (clientname, "", 14, (0, 0, 0)),
                ("Capacity: ", "", 14, (0, 0, 0)),
                (capacity, "", 14, (0, 0, 0)),
            ]

            
            line_height = 16
            total_content_height = line_height * len(lines)

            
            y_position = (page_height - total_content_height) / 2
            self.set_y(y_position)

            
            for text, style, size, color in lines:
                self.set_font("DejaVu", style, size)
                self.set_text_color(*color)
                self.cell(0, line_height, text, ln=True, align="C")

        def header(self):
            # Logo
            page_width = self.w - 2 * self.l_margin
            
            try:
                self.image(f"static/aLogo/{logo}.png", x=self.l_margin, y=10, w=page_width)  

                # Move cursor below image to avoid overlapping next content
                self.set_y(10 + 50)
            except Exception as e:
                print(f"Error loading logo: {e}")


        def footer(self):
            self.set_y(-15)  # Position 15 mm from bottom
            self.set_font("DejaVu", "I", 8)
            self.set_text_color(128)

            # Add "Page X of Y"
            self.cell(0, 10, f"Page {self.page_no()} of {{nb}}", align='C')

        def colored_header(self, number, name):
            self.set_font("DejaVu", "B", 16)
            self.set_text_color(33, 66, 133)
            self.cell(0, 10, f"{number}. {name}", ln=True)

        def section_title(self, title):
            self.set_font("DejaVu", "B", 12)
            self.set_text_color(33, 66, 133)
            self.ln(5)
            self.cell(0, 10, title, ln=True)

        def add_table(self, data):
            self.set_font("DejaVu", "B", 10)
            if logo == "LogoAAA":
                self.set_fill_color(255, 153, 0)
                tableborder = 1
                self.set_text_color(0)
                self.cell(80, 8, "Field", border=tableborder, fill=True)
                self.cell(110, 8, "Value", border=tableborder, ln=True, fill=True)
            elif logo == "LogoBBB": 
                self.set_fill_color(255, 255, 255)
                tableborder = 0
                self.set_text_color(0)
                self.cell(80, 8, " ", border=tableborder, fill=True)
                self.cell(110, 8, " ", border=tableborder, ln=True, fill=True)

            self.set_font("DejaVu", "", 10)
            for field, value in data:
                self.cell(80, 8, field, border=tableborder)
                self.cell(110, 8, value, border=tableborder, ln=True)
    
    
    try:
        
        ###LOG
        aLogEntry.objects.create(
                user=user,
                message=f"at {now()} {user} accessed Load  "
            )
        print(f"at {now()} {User} accessed Download Report")
        ###LOG

        aCompany = UserCompany.objects.get(user=user)
        companyname=aCompany.company.nameCompanies
        firstletter = companyname[0]
        project = APP_Project.objects.get(id=project_id)
        machines = modelcalc.objects.filter(project=project)
        
        
        print(aCompany.id)
        print(project.id)
    
        print(f"Company {aCompany.id}")
    
    
        # Create a Word document
        doc = Document()

        pdf = PDF()

        # Add header and footer with page numbers
        add_header_footer(doc)

        pdf.alias_nb_pages()  # Important for "of {nb}" to work
        pdf.add_page()

        pdf.header_footer(project.name,  project.client_name, project.capacity)

        for _ in range(4):  # Adjust this number based on how centered you want it
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Centered content
        lines = [
            "Project Name: ",
            project.name,
            "Client Name: ",
            project.client_name,
            "Capacity: ",
            project.capacity,
        ]

        for line in lines:
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.size = Pt(25) 
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing after if needed
        para = doc.add_paragraph("\n")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()     
        doc.add_paragraph("\n")

        # Add machine details
        for index, machine in enumerate(machines, start=1):  # Add numbering
            machine_name = machine.oSec00Field03
            machine_name == f"{machine_name}_{firstletter}"

            try:
                themachines = AddMachine.objects.get(keyValue=machine_name)
            except AddMachine.DoesNotExist:
                print(f"Skipping machine '{machine_name}' for company ID {aCompany.id}: not found in AddMachine.")
                continue  # Skip this machine and continue with the rest
            
            themachinename = themachines.nameMachine
            section_titles = ["Inputs", "Outputs"]

            # Add machine title with font size 14 and numbering
            machine_title = doc.add_paragraph(f"{index}. {themachinename}", style="Heading3")
            machine_title.runs[0].font.size = Pt(14)
            
            pdf.alias_nb_pages()  # Important for "of {nb}" to work
            pdf.add_page()
            pdf.colored_header(index, themachinename)

            for i in range(1, 3):  # Loop from Sec01 to Sec02
                section_name = f"Sec{i:02d}"
                if aCompany.company.nameCompanies == "AAAA":
                    section_data = [("Field", "Value")]
                elif aCompany.company.nameCompanies == "BBBB":
                    section_data = [(" ", " ")]
                pdf_section_data = []

                for j in range(1, 31, 2):  # Step by 2 to avoid duplication
                    key = getattr(machine, f"o{section_name}Field{j:02d}", "").strip()
                    value = getattr(machine, f"o{section_name}Field{j+1:02d}", "").strip()

                    if key and value and key.lower() != "oooo" and value.lower() != "oooo":
                        section_data.append((key, value))
                        pdf_section_data.append((key, value))

                if len(section_data) > 1:  # If the section has valid data, create a table
                    if aCompany.company.nameCompanies == "AAAA":
                        section_title = section_titles[i-1] if i-1 < len(section_titles) else f"Section {i}"
                        doc.add_paragraph(f"{section_name}: {section_title}", style="Heading3")  # Only one title now

                        pdf.section_title(f"{section_name}: {section_title}")

                    add_table(doc, section_data)  # Removed redundant title

                    pdf.add_table(pdf_section_data)

            doc.add_page_break()     

        # Define the folder path
        company_slug = slugify(project.company.nameCompanies)
        company_name = company_slug.upper()
        project_slug = slugify(project.name)
        folder_name = slugify(f"{project_id}_{company_slug}_{project_slug}")

        project_folder = os.path.join(settings.BASE_DIR, 'static', 'aReports', company_slug, folder_name)
        os.makedirs(project_folder, exist_ok=True)  # Create folder if it doesn't exist

        # Save the file to that path
        file_path = os.path.join(project_folder, f"{project_slug}_Calculation_report.docx")
        doc.save(file_path)
        
        """ doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0) """

        file_name = f"{project_slug}_Calculation_report.docx"
        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        folder_exist, folder_data = check_folder_exists(service, "aReports")
        if folder_exist == True :
            folder_id = folder_data['id']
            company_folder_exist, company_folder_data = check_folder_exists(service, company_name, folder_id)
            if company_folder_exist == True:
                company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
                project_folder_exist, project_folder_data = check_folder_exists(service, folder_name, company_folder_id)
                if project_folder_exist == True:
                    project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                    upload_files(service,file_path, file_name, mime_type, project_folder_id)
                    """ upload_files_directly(service, doc_buffer, file_name, mime_type, project_folder_id) """
                else:
                    create_folder(service, folder_name, company_folder_id)
                    project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                    upload_files(service,file_path, file_name, mime_type, project_folder_id)
                    """ upload_files_directly(service, doc_buffer, file_name, mime_type, project_folder_id) """
            else:
                create_folder(service, company_name, folder_id)
                company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
                create_folder(service, folder_name, company_folder_id)
                project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                upload_files(service,file_path, file_name, mime_type, project_folder_id)
                """ upload_files_directly(service, doc_buffer, file_name, mime_type, project_folder_id) """
        else:
            create_folder(service, "aReports")
            folder_id = get_folder_id_by_name(service, "aReports")
            create_folder(service, company_name, folder_id)
            company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
            create_folder(service, folder_name, company_folder_id)
            project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
            upload_files(service,file_path, file_name, mime_type, project_folder_id)

        pdf_file_path = os.path.join(project_folder, f"{project_slug}_Calculation_report.pdf")
        pdf.output(pdf_file_path)

        """ pdf_bytes = pdf.output(dest='S').encode('latin1')  # Get PDF as bytes
        pdf_buffer = BytesIO(pdf_bytes)
        pdf_buffer.seek(0) """

        file_name = f"{project_slug}_Calculation_report.pdf"
        mime_type = 'application/pdf'

        folder_exist, folder_data = check_folder_exists(service, "aReports")
        if folder_exist == True :
            folder_id = folder_data['id']
            company_folder_exist, company_folder_data = check_folder_exists(service, company_name, folder_id)
            if company_folder_exist == True:
                company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
                project_folder_exist, project_folder_data = check_folder_exists(service, folder_name, company_folder_id)
                if project_folder_exist == True:
                    project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                    upload_files(service,pdf_file_path, file_name, mime_type, project_folder_id)
                    """ upload_files_directly(service, pdf_buffer, file_name, mime_type, project_folder_id) """
                else:
                    create_folder(service, folder_name, company_folder_id)
                    project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                    upload_files(service,pdf_file_path, file_name, mime_type, project_folder_id)
                    """ upload_files_directly(service, pdf_buffer, file_name, mime_type, project_folder_id) """
            else:
                create_folder(service, company_name, folder_id)
                company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
                create_folder(service, folder_name, company_folder_id)
                project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                upload_files(service,pdf_file_path, file_name, mime_type, project_folder_id)
                """ upload_files_directly(service, pdf_buffer, file_name, mime_type, project_folder_id) """
        else:
            create_folder(service, "aReports")
            folder_id = get_folder_id_by_name(service, "aReports")
            create_folder(service, company_name, folder_id)
            company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
            create_folder(service, folder_name, company_folder_id)
            project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
            upload_files(service,pdf_file_path, file_name, mime_type, project_folder_id)

        return HttpResponse(status=204)

        

    except APP_Project.DoesNotExist:
        return HttpResponse("Project not found", status=404)

        
def save_all_pdf_report(user, project_id, logo):
    
    # Define a custom class for the PDF layout
    class PDF(FPDF):
        def __init__(self):
            super().__init__()
            self.set_auto_page_break(auto=True, margin=15)

            # Add the TrueType Unicode font (DejaVuSans)
            self.add_font('DejaVu', '', 'static/aApp1/fonts/DejaVu/DejaVuSans.ttf', uni=True)  # Regular
            self.add_font('DejaVu', 'B', 'static/aApp1/fonts/DejaVu/DejaVuSans-Bold.ttf', uni=True)  # Bold
            self.add_font('DejaVu', 'I', 'static/aApp1/fonts/DejaVu/DejaVuSerif-Italic.ttf', uni=True)  # Italic
            self.set_font('DejaVu', '', 12)

        def header_page(self, projectname, clientname, capacity):
            self.set_y(50)
            
            page_height = self.h
            top_margin = self.t_margin
            bottom_margin = self.b_margin

            
            lines = [
                ("Project Name: ", "", 14, (0, 0, 0)),
                (projectname, "", 14, (0, 0, 0)),
                ("Client Name: ", "", 14, (0, 0, 0)),
                (clientname, "", 14, (0, 0, 0)),
                ("Capacity: ", "", 14, (0, 0, 0)),
                (capacity, "", 14, (0, 0, 0)),
            ]

            
            line_height = 16
            total_content_height = line_height * len(lines)

            
            y_position = (page_height - total_content_height) / 2
            self.set_y(y_position)

            
            for text, style, size, color in lines:
                self.set_font("DejaVu", style, size)
                self.set_text_color(*color)
                self.cell(0, line_height, text, ln=True, align="C")

        def header(self):
            # Logo
            page_width = self.w - 2 * self.l_margin
            
            try:
                self.image(f"static/aLogo/{logo}.png", x=self.l_margin, y=10, w=page_width)  

                # Move cursor below image to avoid overlapping next content
                self.set_y(10 + 50)
            except Exception as e:
                print(f"Error loading logo: {e}")


        def footer(self):
            self.set_y(-15)  # Position 15 mm from bottom
            self.set_font("DejaVu", "I", 8)
            self.set_text_color(128)

            # Add "Page X of Y"
            self.cell(0, 10, f"Page {self.page_no()} of {{nb}}", align='C')

        def colored_header(self, number, name):
            self.set_font("DejaVu", "B", 16)
            self.set_text_color(33, 66, 133)
            
            text = f"{number}. {name}"
            text_width = self.get_string_width(text) + 2  # Optional padding
            text_height = 10  # Height for one line of text

            # Page dimensions
            page_width = self.w
            page_height = self.h

            # Center position
            x = (page_width - text_width) / 2
            y = (page_height - text_height) / 2

            self.set_xy(x, y)
            self.cell(text_width, text_height, text, border=0, align='C')

        def section_title(self, title):
            self.set_font("DejaVu", "B", 12)
            self.set_text_color(33, 66, 133)
            self.ln(5)
            self.cell(0, 10, title, ln=True)

        def add_table(self, data):
            self.set_font("DejaVu", "B", 10)
            if logo == "LogoAAA":
                self.set_fill_color(255, 153, 0)
                tableborder = 1
                self.set_text_color(0)
                self.cell(80, 8, "Field", border=tableborder, fill=True)
                self.cell(110, 8, "Value", border=tableborder, ln=True, fill=True)
            elif logo == "LogoBBB": 
                self.set_fill_color(255, 255, 255)  
                tableborder = 0
                self.set_text_color(0)
                self.cell(80, 8, " ", border=tableborder, fill=True)
                self.cell(110, 8, " ", border=tableborder, ln=True, fill=True)

            self.set_font("DejaVu", "", 10)
            for field, value in data:
                self.cell(80, 8, field, border=tableborder)
                self.cell(110, 8, value, border=tableborder, ln=True)
    
    
    try:
        
        ###LOG
        aLogEntry.objects.create(
                user=user,
                message=f"at {now()} {user} accessed Load  "
            )
        print(f"at {now()} {User} accessed Save Report")
        ###LOG

        aCompany = UserCompany.objects.get(user=user)
        company_id = aCompany.company
        project = APP_Project.objects.get(id=project_id)
        machines = Machine.objects.filter(project=project)
        paragraph_text = project.cover_page_text or "No description provided."

        # Define the folder path
        company_slug = slugify(project.company.nameCompanies)
        company_name = company_slug.upper()
        project_slug = slugify(project.name)
        folder_name = slugify(f"{project_id}_{company_slug}_{project_slug}")

        project_folder = os.path.join(settings.BASE_DIR, 'static', 'aReports', company_slug, folder_name)
        os.makedirs(project_folder, exist_ok=True)  # Create folder if it doesn't exist

        
        pdf_file_path = os.path.join(project_folder, f"all_{project_slug}_report.pdf")

        print(f"Company ID: {aCompany.id}")
        print(f"Project ID: {project.id}")
            
        output = PdfWriter()

        # Create a Pdf document
        pdf = PDF()

        pdf.alias_nb_pages()  # Important for "of {nb}" to work
        pdf.add_page()
        pdf.header_page(project.name,  project.client_name, project.capacity)

        pdf.alias_nb_pages()  # Important for "of {nb}" to work
        pdf.add_page()
        pdf.set_font("DejaVu", "", 12)
        pdf.set_text_color(0)
        pdf.multi_cell(0, 10, paragraph_text, align='L')

        pdf.output(pdf_file_path)

        mainreader = PdfReader(pdf_file_path)

        for page in mainreader.pages:
            output.add_page(page)
        
        # clean up
        os.remove(pdf_file_path)

        # Add machine details
        for index, machine in enumerate(machines, start=1):  # Add numbering
            machine_name = machine.oSec00Field03
            machine_DB = machine.oSec00Field03
            
            try:
                themachines = AddMachine.objects.get(nameDB=machine_DB, company=company_id)
                
            except AddMachine.DoesNotExist:
                print(f"Skipping machine '{machine_name}' for company ID {company_id}: not found in AddMachine.")
                continue  # Skip this machine and continue with the rest
            
            themachinename = themachines.nameMachine
            file_name = themachines.nameDXF

            section_titles = [" ", " ", " ", " ", " ", " ", " ", " ", " ", " "]
            if machine_name == "DataSheetNS":
                machine_name = "Manual Screen" 
                sheet_key = "NS"
                section_titles = ["General Data", "Design Data", "Material Data", "Channel Data", " ", " ", " ", " ", " ", " "]

            elif machine_name == "DataSheetMS":
                machine_name = "Mechanical Screen"
                sheet_key = "MS" 
                section_titles = ["General Data", "Design Data", "Gearmotor Data", "Control panel Data", "Material Data", "Other Data", " ", " ", " ", " "]

            elif machine_name == "DataSheetBC":
                machine_name = "Belt Conveyor"
                sheet_key = "BC"
                section_titles = ["General Data", "Design Data", "Gearbox Data", "Motor Data", "Material Data", " ", " ", " ", " ", " "]

            elif machine_name == "DataSheetCO":
                machine_name = "Container"
                sheet_key = "CO"
                section_titles = ["General Data", "Design Data", "Material Data", " ", " ", " ", " ", " ", " ", " "]

            elif machine_name == "DataSheetGR":
                machine_name = "Gritremoval"
                sheet_key = "GR"
                section_titles = ["General Data", "Design Data", "Walkway, Handrail, Wheel Data", "Scrapper Data", "Gearmotor Data", "Scrapper Data", "Drive unit", "Control panel Data", "Material Data ", " "]



            sheet_key = themachines.keyValue[0:-2]
            # Add machine name 
            pdf = PDF()
            pdf.alias_nb_pages()  
            pdf.add_page()
            pdf.colored_header(index, themachinename)

            if machine.oSec01Field01 not in ["oooo", None , ""]:
                pdf.alias_nb_pages()  
                pdf.add_page()

            for i in range(1, 11):  # Loop from Sec01 to Sec10
                section_name = f"Sec{i:02d}"
                pdf_section_data = []

                for j in range(1, 21, 2):  # Step by 2 to avoid duplication
                    key = getattr(machine, f"o{section_name}Field{j:02d}", "").strip()
                    value = getattr(machine, f"o{section_name}Field{j+1:02d}", "").strip()

                    if key and value and key.lower() != "oooo" and value.lower() != "oooo":
                        pdf_section_data.append((key, value))

                if len(pdf_section_data) > 1:  # If the section has valid data, create a table
                    if aCompany.company.nameCompanies == "AAAA":
                        section_title = section_titles[i-1] if i-1 < len(section_titles) else f"Section {i}"
                        pdf.section_title(f"{section_name}: {section_title}")
                    pdf.add_table(pdf_section_data)

            
            pdf.output(pdf_file_path)

            mainreader = PdfReader(pdf_file_path)

            for page in mainreader.pages:
                output.add_page(page)

            # clean up
            os.remove(pdf_file_path)

            appendix = None
            if os.path.exists(f"static/aReports/{company_slug.upper()}/{folder_name}/{file_name}.pdf") :
                appendix = PdfReader(f"static/aReports/{company_slug.upper()}/{folder_name}/{file_name}.pdf")
            elif os.path.exists(f"static/aReports/{company_slug.upper()}/{folder_name}/{sheet_key}_new.pdf") :
                appendix = PdfReader(f"static/aReports/{company_slug.upper()}/{folder_name}/{sheet_key}_new.pdf")

            if appendix != None:
                for page in appendix.pages:
                    output.add_page(page)


        if company_slug == "aaaa":
            for i in range(1 , 6):
                emptyappendix = None
                cost_file_name = f"cost{i}_pdf.pdf"
                if os.path.exists(f"static/aReports/{company_slug.upper()}/{folder_name}/{cost_file_name}") :
                    emptyappendix = PdfReader(f"static/aReports/{company_slug.upper()}/{folder_name}/{cost_file_name}")

                if emptyappendix != None:
                    for page in emptyappendix.pages:
                        output.add_page(page)

        # Save the new PDF
        with open(pdf_file_path, "wb") as f:
            output.write(f)

        file_name = f"all_{project_slug}_report.pdf"
        
        mime_type = 'application/pdf'

        folder_exist, folder_data = check_folder_exists(service, "aReports")
        if folder_exist == True :
            folder_id = folder_data['id']
            company_folder_exist, company_folder_data = check_folder_exists(service, company_name, folder_id)
            if company_folder_exist == True:
                company_folder_id = company_folder_data['id']
                project_folder_exist, project_folder_data = check_folder_exists(service, folder_name, company_folder_id)

                if project_folder_exist == True:
                    project_folder_id = project_folder_data['id']
                    upload_files(service,pdf_file_path, file_name, mime_type, project_folder_id)

                else:
                    create_folder(service, folder_name, company_folder_id)
                    project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                    upload_files(service,pdf_file_path, file_name, mime_type, project_folder_id)

            else:
                create_folder(service, company_name, folder_id)
                company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
                create_folder(service, folder_name, company_folder_id)
                project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
                upload_files(service,pdf_file_path, file_name, mime_type, project_folder_id)
        
        else:
            create_folder(service, "aReports")
            folder_id = get_folder_id_by_name(service, "aReports")
            create_folder(service, company_name, folder_id)
            company_folder_id = get_folder_id_by_name(service, company_name, folder_id)
            create_folder(service, folder_name, company_folder_id)
            project_folder_id = get_folder_id_by_name(service, folder_name, company_folder_id)
            upload_files(service,pdf_file_path, file_name, mime_type, project_folder_id)

        return HttpResponse(status=204)
        

        

    except APP_Project.DoesNotExist:
        return HttpResponse("Project not found", status=404)

